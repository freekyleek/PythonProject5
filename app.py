import re
from flask import Flask, render_template, redirect, url_for, request, flash, send_from_directory, abort, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, logout_user, login_required, UserMixin, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from flask_cors import CORS
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
import os, json, re, datetime, unicodedata, mimetypes
from openpyxl import Workbook
import io
from docx import Document
import json
import datetime

app = Flask(__name__)



# === Avatar upload directory (global) ===
try:
    PROFILE_DIR
except NameError:
    try:
        PROFILE_DIR = os.path.join(app.root_path, 'static', 'avatars')
    except Exception:
        PROFILE_DIR = os.path.join(os.path.dirname(__file__), 'static', 'avatars')
try:
    os.makedirs(PROFILE_DIR, exist_ok=True)
except Exception:
    pass

# --- Real-time status (Socket.IO) ---
try:
    from flask_socketio import SocketIO, emit
    socketio = SocketIO(app, async_mode='threading', cors_allowed_origins="*")
except Exception:
    socketio = None

# ===== Reset password token helpers =====
def _reset_serializer():
    salt = app.config.get('RESET_TOKEN_SALT', 'reset-password')
    return URLSafeTimedSerializer(app.config['SECRET_KEY'], salt=salt)

def generate_reset_token(username: str) -> str:
    return _reset_serializer().dumps({'u': username})

def verify_reset_token(token: str, max_age: int = None):
    if max_age is None:
        max_age = int(app.config.get('RESET_TOKEN_MAX_AGE', 3600))  # default 1h
    try:
        data = _reset_serializer().loads(token, max_age=max_age)
        return data.get('u')
    except SignatureExpired:
        return None
    except BadSignature:
        return None

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'super-secret-key-change-me')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///site.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# --- Gmail konfiguracija po zahtjevu korisnika ---
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = 'webtest806@gmail.com'
app.config['MAIL_PASSWORD'] = 'gsqb frrp bsmh uxyj'  # Gmail App Password (na zahtjev korisnika)
app.config['MAIL_DEFAULT_SENDER'] = 'webtest806@gmail.com'

db = SQLAlchemy(app)
CORS(app)
mail = Mail(app)

login_manager = LoginManager(app)
login_manager.login_view = "login"
login_manager.login_message_category = "info"

# -------------------- MODELI --------------------
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(30), nullable=False, default="user")

    def set_password(self, password: str):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password: str) -> bool:
        return check_password_hash(self.password_hash, password)

    @property
    def is_superadmin(self):
        return self.role == "superadmin"

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# -------------------- ROLE & PERMISSIONS (Jinja helpers) --------------------
import unicodedata

def _norm(s: str) -> str:
    try:
        s = unicodedata.normalize('NFKD', s or '').encode('ascii','ignore').decode('ascii')
    except Exception:
        s = str(s or '')
    return s.strip().lower()

def current_role():
    try:
        r = getattr(current_user, 'role', '') or ''
        return _norm(r)
    except Exception:
        return ''

def has_role(*roles) -> bool:
    r = current_role()
    allowed = {_norm(x) for x in roles}
    return r in allowed

@app.context_processor
def inject_role_helpers():
    # Expose helpers to Jinja
    return {'has_role': has_role, 'current_role': current_role}

# -------------------- JSON pomoćne funkcije --------------------
STATIC_DIR = os.path.join(app.root_path, 'static')

# === Signature upload directory (PNG only) ===
try:
    SIGNATURE_DIR
except NameError:
    try:
        SIGNATURE_DIR = os.path.join(app.root_path, 'static', 'signatures')
    except Exception:
        SIGNATURE_DIR = os.path.join(os.path.dirname(__file__), 'static', 'signatures')
try:
    os.makedirs(SIGNATURE_DIR, exist_ok=True)
except Exception:
    pass

# -------------------- OTPISANI UREĐAJI JSON --------------------
OTPISANI_UREDJAJI_JSON_PATH = os.path.join(STATIC_DIR, 'otpisani.uredjaji.JSON')
def read_otpisani_uredjaji():
    return _read_json(OTPISANI_UREDJAJI_JSON_PATH)
def write_otpisani_uredjaji(data):
    _write_json(OTPISANI_UREDJAJI_JSON_PATH, data)

OPERATERI_JSON_PATH = os.path.join(STATIC_DIR, 'operateri.json')
KLJENTI_JSON_PATH = os.path.join(STATIC_DIR, 'klijenti.json')
NALOZI_JSON_PATH = os.path.join(STATIC_DIR, 'nalozi.json')
# Usluge & Pay opcije konfiguracije (dinamički prikaz u RN malom zapisniku)
USLUGE_JSON_PATH = os.path.join(STATIC_DIR, 'usluge.JSON')
PAY_OPCIJE_JSON_PATH = os.path.join(STATIC_DIR, 'pay.opcije.JSON')
NACIN_ISPORUKE_JSON_PATH = os.path.join(STATIC_DIR, 'nacin.isporuke.json')
def read_nacin_isporuke(): return _read_json(NACIN_ISPORUKE_JSON_PATH)
def read_usluge_config(): return _read_json(USLUGE_JSON_PATH)
def read_pay_opcije(): return _read_json(PAY_OPCIJE_JSON_PATH)


def _read_json(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return []

def _write_json(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def read_operateri(): return _read_json(OPERATERI_JSON_PATH)
def write_operateri(data): _write_json(OPERATERI_JSON_PATH, data)
def read_klijenti(): return _read_json(KLJENTI_JSON_PATH)
def write_klijenti(data): _write_json(KLJENTI_JSON_PATH, data)
def read_nalozi(): return _read_json(NALOZI_JSON_PATH)
def write_nalozi(data): _write_json(NALOZI_JSON_PATH, data)
# -------------------- TEHNIČARI / ZADUŽENI UREĐAJI JSON --------------------
TEHNICARI_JSON_PATH = os.path.join(STATIC_DIR, 'tehnicari.json')
ZADUZENI_UREDJAJI_JSON_PATH = os.path.join(STATIC_DIR, 'zaduzeni.uredjaji.json')
# -------------------- ZADUŽENI SIM JSON --------------------
ZADUZENI_SIM_JSON_PATH = os.path.join(STATIC_DIR, 'zaduzeni.sim.json')
# -------------------- AKTIVNI UREĐAJI (KOD KLIJENTA) --------------------
AKTIVNI_UREDJAJI_JSON_PATH = os.path.join(STATIC_DIR, 'aktivni_uredjaji.JSON')

# -------------------- NAMJENA UREĐAJA JSON --------------------
NAMJENA_UREDJAJA_JSON_PATH = os.path.join(STATIC_DIR, 'namjena.uredjaja.JSON')
def read_namjena_uredjaja():
    try:
        data = _read_json(NAMJENA_UREDJAJA_JSON_PATH)
    except Exception:
        data = []
    mapping = {}
    try:
        if isinstance(data, dict):
            for k, v in data.items():
                if k:
                    mapping[str(k)] = str(v)
        elif isinstance(data, list):
            for row in data:
                try:
                    s = str((row.get('serijski') or row.get('sn') or '')).strip()
                    n = str((row.get('namjena') or '')).strip()
                    if s and n:
                        mapping[s] = n
                except Exception:
                    pass
    except Exception:
        pass


    """
    # --- PROVJERA: Ako su svi uređaji privremeno isključeni, klijent postaje neaktivan ---
    try:
        any_active = any((u.get('status_color') == 'green') for u in uredjaji_klijent or [])
        if not any_active:
            klijent_status = 'neaktivan'
        else:
            klijent_status = 'aktivan'
    except Exception:
        klijent_status = 'aktivan'
    """
    return mapping

def write_namjena_uredjaja(mapping):
    try:
        _write_json(NAMJENA_UREDJAJA_JSON_PATH, mapping)
    except Exception:
        pass


def read_aktivni_uredjaji(): return _read_json(AKTIVNI_UREDJAJI_JSON_PATH)
def write_aktivni_uredjaji(data): _write_json(AKTIVNI_UREDJAJI_JSON_PATH, data)




# -------------------- AKTIVNI SIM (KOD KLIJENTA) --------------------
AKTIVNI_SIM_JSON_PATH = os.path.join(STATIC_DIR, 'aktivni_sim.JSON')
def read_aktivni_sim(): return _read_json(AKTIVNI_SIM_JSON_PATH)
def write_aktivni_sim(data): _write_json(AKTIVNI_SIM_JSON_PATH, data)
# -------------------- SERVIS UREĐAJI (uređaji uklonjeni s klijenta zbog servisa) --------------------
SERVIS_UREDJAJI_JSON_PATH = os.path.join(STATIC_DIR, 'servis.uredjaji.JSON')

KUPLJENI_UREDJAJI_JSON_PATH = os.path.join(STATIC_DIR, 'kupljeni.uredjaji.JSON')

def read_servis_uredjaji(): return _read_json(SERVIS_UREDJAJI_JSON_PATH)


def get_namjena_by_serijski(serial):
    try:
        base = read_uredjaji()
    except Exception:
        base = []
    for d in base:
        if str(d.get('serijski')) == str(serial):
            return d.get('namjena', '')
    return ''

def read_kupljeni_uredjaji(): return _read_json(KUPLJENI_UREDJAJI_JSON_PATH)

def sync_kupljeni_from_aktivni():
    """
    Iz aktivni_uredjaji.JSON prebaci (bez brisanja iz izvora) SVE stavke čija je namjena točno 'Kupnja'
    u kupljeni.uredjaji.JSON. One s 'Najam' se ne prebacuju. 'Namjena' se čita iz uredjaji.JSON po serijskom broju.
    """
    try:
        aktivni = read_aktivni_uredjaji()
    except Exception:
        aktivni = []
    try:
        base = read_uredjaji()
    except Exception:
        base = []
    # mapa SN -> namjena
    namjena_by_sn = {}
    for row in base:
        try:
            s = str(row.get('serijski') or '').strip()
            if s:
                namjena_by_sn[s] = str(row.get('namjena','')).strip()
        except Exception:
            pass
    kupljeni = []
    seen = set()
    for it in aktivni or []:
        try:
            sn = str(it.get('serijski') or '').strip()
            if not sn or sn in seen:
                continue
            nm = namjena_by_sn.get(sn, '')
            if nm == 'Kupnja':  # samo Kupnja
                # kopiraj sve dostupne podatke + namjena
                row = dict(it)
                row['namjena'] = 'Kupnja'
                kupljeni.append(row)
                seen.add(sn)
        except Exception:
            pass
    # zapiši kupljene
    try:
        write_kupljeni_uredjaji(kupljeni)
    except Exception:
        pass
    return kupljeni

def write_servis_uredjaji(data): _write_json(SERVIS_UREDJAJI_JSON_PATH, data)


def read_zaduzene_sim(): return _read_json(ZADUZENI_SIM_JSON_PATH)
def write_zaduzene_sim(data): _write_json(ZADUZENI_SIM_JSON_PATH, data)


def read_tehnicari(): return _read_json(TEHNICARI_JSON_PATH)
def write_tehnicari(data): _write_json(TEHNICARI_JSON_PATH, data)
# -------------------- KLIJENTI JSON --------------------
KLIJENTI_JSON_PATH = os.path.join(STATIC_DIR, 'klijenti.json')
def read_klijenti(): return _read_json(KLIJENTI_JSON_PATH)
def write_klijenti(data): _write_json(KLIJENTI_JSON_PATH, data)

# ---- Helpers for client status broadcasting ----
# ---- Compute client activity (devices/SIMs) ----
def _compute_client_active_by_name(name: str) -> bool:
    cname = str(name or '')

    def _norm_sn(v):
        import re as _re
        return _re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    # === Active DEVICES for this client (excluding those temporarily disabled) ===
    try:
        akt_dev = read_aktivni_uredjaji() or []
    except Exception:
        akt_dev = []
    # Privremeno isključeni uređaji set
    try:
        _priv_dev = read_privremeno_iskljuceni() or []
    except Exception:
        _priv_dev = []
    priv_dev_set = {_norm_sn(x.get('serijski')) for x in _priv_dev if x.get('serijski')}

    dev_active = False
    for d in akt_dev:
        try:
            if not d.get('active', True):
                continue
            if str(d.get('client','')) != cname:
                continue
            # Exclude if device is marked as temporarily disabled
            if _norm_sn(d.get('serijski')) in priv_dev_set:
                continue
            dev_active = True
            break
        except Exception:
            pass

    # === Active SIMs for this client (excluding those temporarily disabled) ===
    try:
        akt_sim = read_aktivni_sim() or []
    except Exception:
        akt_sim = []
    try:
        _priv_sim = read_privremeno_iskljuceni_sim() or []
    except Exception:
        _priv_sim = []
    priv_sim_set = {_norm_sn(x.get('serijski')) for x in _priv_sim if x.get('serijski')}

    sim_active = False
    for s in akt_sim:
        try:
            if not s.get('active', True):
                continue
            if str(s.get('client','')) != cname:
                continue
            if _norm_sn(s.get('serijski')) in priv_sim_set:
                # temporarily disabled -> treat as inactive
                continue
            sim_active = True
            break
        except Exception:
            pass

    # Client is active if at least one device OR one SIM is active (and not temporarily disabled)
    return bool(dev_active or sim_active)

def _compute_all_clients_statuses():
    try:
        kl = read_klijenti() or []
    except Exception:
        kl = []
    out = []
    for k in kl:
        try:
            nm = str(k.get('name') or k.get('naziv') or '')
            out.append({'name': nm, 'active': _compute_client_active_by_name(nm)})
        except Exception:
            pass
    return out

def _broadcast_client_status(name: str, active: bool):
    try:
        if 'socketio' in globals() and socketio:
            socketio.emit('client_status', {'name': name, 'active': bool(active)}, namespace='/status', broadcast=True)
    except Exception:
        pass


# -------------------- ZADUŽENI NALOZI JSON --------------------
ZADUZENI_NALOZI_JSON_PATH = os.path.join(STATIC_DIR, 'zaduzeni.nalozi.json')
def read_zaduzene_nalozi(): return _read_json(ZADUZENI_NALOZI_JSON_PATH)
def write_zaduzene_nalozi(data): _write_json(ZADUZENI_NALOZI_JSON_PATH, data)
# -------------------- OTVORENE INSTALACIJE JSON --------------------
OTVORENE_INSTALACIJE_JSON_PATH = os.path.join(STATIC_DIR, 'otvorene_instalacije.JSON')
def read_otvorene_instalacije(): return _read_json(OTVORENE_INSTALACIJE_JSON_PATH)
def write_otvorene_instalacije(data): _write_json(OTVORENE_INSTALACIJE_JSON_PATH, data)

# -------------------- OTVORENE DEINSTALACIJE JSON --------------------
OTVORENE_DEINSTALACIJE_JSON_PATH = os.path.join(STATIC_DIR, 'otvorene_deinstalacije.JSON')
def read_otvorene_deinstalacije(): return _read_json(OTVORENE_DEINSTALACIJE_JSON_PATH)
def write_otvorene_deinstalacije(data): _write_json(OTVORENE_DEINSTALACIJE_JSON_PATH, data)

# -------------------- OTVORENI SERVISI JSON --------------------
OTVORENI_SERVISI_JSON_PATH = os.path.join(STATIC_DIR, 'otvoreni_servisi.JSON')

# -------------------- SERVISNI NALOGI JSON (svi servisni nalozi) --------------------
SERVISNI_NALOG_JSON_PATH = os.path.join(STATIC_DIR, 'servisni.nalog.JSON')
def read_servisni_nalozi(): return _read_json(SERVISNI_NALOG_JSON_PATH)
def write_servisni_nalozi(data): _write_json(SERVISNI_NALOG_JSON_PATH, data)

def read_otvoreni_servisi(): return _read_json(OTVORENI_SERVISI_JSON_PATH)
def write_otvoreni_servisi(data): _write_json(OTVORENI_SERVISI_JSON_PATH, data)


# -------------------- ZAKLJUČENI NALOZI JSON --------------------
ZAKLJUCENI_NALOZI_JSON_PATH = os.path.join(STATIC_DIR, 'zakljuceni.nalozi.JSON')
def read_zakljuceni_nalozi(): return _read_json(ZAKLJUCENI_NALOZI_JSON_PATH)
def write_zakljuceni_nalozi(data): _write_json(ZAKLJUCENI_NALOZI_JSON_PATH, data)


def read_zaduzene_uredjaje(): return _read_json(ZADUZENI_UREDJAJI_JSON_PATH)
def write_zaduzene_uredjaje(data): _write_json(ZADUZENI_UREDJAJI_JSON_PATH, data)

def refresh_tehnicari():
    ops = read_operateri()
    tehn = [o for o in ops if str(o.get('role','')).strip().lower() in ('tehničar','tehnicar','tehnicar/serviser','tehničar/serviser')]
    write_tehnicari(tehn)
    return tehn


# -------------------- UREĐAJI JSON --------------------
NAZIVI_UREDJaja_JSON_PATH = os.path.join(STATIC_DIR, 'naziv_uredjaja.json')
UREDJAJI_JSON_PATH = os.path.join(STATIC_DIR, 'uredjaji.JSON')

def read_nazivi_uredjaja(): return _read_json(NAZIVI_UREDJaja_JSON_PATH)
def write_nazivi_uredjaja(data): _write_json(NAZIVI_UREDJaja_JSON_PATH, data)
def read_uredjaji(): return _read_json(UREDJAJI_JSON_PATH)
def write_uredjaji(data): _write_json(UREDJAJI_JSON_PATH, data)

# -------------------- PRIVREMENO ISKLJUČENI UREĐAJI --------------------
PRIVREMENO_ISKLJ_JSON_PATH = os.path.join(STATIC_DIR, 'privremeno.iskljuceni.uredjaji.JSON')
def read_privremeno_iskljuceni():
    data = _read_json(PRIVREMENO_ISKLJ_JSON_PATH)
    # format: list of {'serijski': 'SN', 'since': iso}
    if isinstance(data, list):
        return data
    return []

def write_privremeno_iskljuceni(data):
    _write_json(PRIVREMENO_ISKLJ_JSON_PATH, data)

# -------------------- SIM JSON --------------------
PROVIDER_JSON_PATH = os.path.join(STATIC_DIR, 'provider.json')
SIM_JSON_PATH = os.path.join(STATIC_DIR, 'sim.json')

def read_provider(): return _read_json(PROVIDER_JSON_PATH)
def write_provider(data): _write_json(PROVIDER_JSON_PATH, data)
def read_sim(): return _read_json(SIM_JSON_PATH)
def write_sim(data): _write_json(SIM_JSON_PATH, data)

# -------------------- PRIVREMENO ISKLJUČENI SIM JSON --------------------
PRIVREMENO_ISKLJUCENI_SIM_JSON_PATH = os.path.join(STATIC_DIR, 'privremeno.iskljuceni.sim.JSON')
def read_privremeno_iskljuceni_sim():
    try:
        data = _read_json(PRIVREMENO_ISKLJUCENI_SIM_JSON_PATH)
    except Exception:
        data = []
    # Normalize to list of dicts with 'serijski'
    out = []
    if isinstance(data, list):
        out = data
    elif isinstance(data, dict):
        # allow dict form {serijski: {...}}
        for k, v in data.items():
            row = dict(v or {})
            row['serijski'] = row.get('serijski') or k
            out.append(row)
    else:
        out = []
    return out

def write_privremeno_iskljuceni_sim(data):
    _write_json(PRIVREMENO_ISKLJUCENI_SIM_JSON_PATH, data)

# -------------------- SIM RUTE --------------------
@app.route('/sim')
@login_required
def sim():
    items = read_sim()

    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    # Aktivni kod klijenta (for status + fallback client name)
    try:
        aktivni_sim = read_aktivni_sim()
    except Exception:
        aktivni_sim = []
    active_serials = {_norm_sn(s.get('serijski')) for s in aktivni_sim if s.get('active', True)}
    active_client_by_sn = {}

    # Privremeno isključeni SIM-ovi set (za status prikaz)
    try:
        _priv_sim = read_privremeno_iskljuceni_sim()
    except Exception:
        _priv_sim = []
    priv_set = { _norm_sn(x.get('serijski')) for x in _priv_sim if x.get('serijski') }
    for s in aktivni_sim:
        try:
            if s.get('active', True) and s.get('client'):
                active_client_by_sn[_norm_sn(s.get('serijski'))] = str(s.get('client'))
        except Exception:
            pass

    # Zaduženi kod tehničara -> ŽUTO i popuniti "assigned_to"
    try:
        zaduzeni_sim = read_zaduzene_sim()
    except Exception:
        zaduzeni_sim = []
    tech_by_serial = {_norm_sn(z.get('serijski')): str(z.get('assigned_to') or '') for z in zaduzeni_sim}

    # KLIJENT MAPA iz klijenti.json ("sn_SIM")
    try:
        klijenti = read_klijenti()
    except Exception:
        klijenti = []
    client_by_serial = {}
    for k in klijenti:
        cname = str(k.get('name') or '').strip()
        try:
            for sn in str(k.get('sn_SIM') or '').split(','):
                snn = _norm_sn(sn)
                if snn:
                    client_by_serial[snn] = cname
        except Exception:
            pass

    # Merge fallback from aktivni_* where klijenti.json doesn't have it
    for snn, cname in active_client_by_sn.items():
        client_by_serial.setdefault(snn, cname)

    for it in items:
        raw_sn = it.get('serijski','')
        sn = _norm_sn(raw_sn)
        assigned_to = tech_by_serial.get(sn, '')
        is_assigned = bool(assigned_to)

        # NEW status logic:
        # - GREEN if active at client (aktivni or klijenti)
        # - RED if NOT assigned to technician AND NOT active
        # - YELLOW otherwise (free or assigned to technician)
        if sn in priv_set:
            it['status_color'] = 'orange'
        elif sn in active_serials or sn in client_by_serial:
            it['status_color'] = 'green'
        elif (not is_assigned) and (sn not in active_serials) and (sn not in client_by_serial):
            it['status_color'] = 'red'
        else:
            it['status_color'] = 'yellow'

        it['assigned_to'] = assigned_to
        it['client_name'] = client_by_serial.get(sn, '')

    items.sort(key=lambda x: (x.get('provider',''), x.get('serijski','')))
    # --- Brojači za prikaz na /sim ---
    aktivni_count_sim = sum(1 for _it in items if _it.get('status_color') == 'green')
    neaktivni_count_sim = sum(1 for _it in items if _it.get('status_color') == 'red')
    zaduzeni_count_sim = sum(1 for _it in items if str(_it.get('assigned_to') or '').strip())
    privremeni_count_sim = sum(1 for _it in items if _it.get('status_color') == 'orange')
    
    # --- Server-side filtering by query param `q` (serijski broj or klijent) ---
    try:
        q = (request.args.get('q') or '').strip()
    except Exception:
        q = ''
    def __norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()
    if q:
        _q = __norm_sn(q)
        view_items = []
        for it in items:
            try:
                ser = __norm_sn(it.get('serijski'))
                cli = __norm_sn(it.get('client_name'))
                if (_q in ser) or (_q in cli):
                    view_items.append(it)
            except Exception:
                pass
    else:
        view_items = items
    return render_template('sim.html', title="SIM", username=current_user.username, sims=view_items, aktivni_count=aktivni_count_sim, neaktivni_count=neaktivni_count_sim, zaduzeni_count=zaduzeni_count_sim, privremeni_count=privremeni_count_sim)

@app.route('/dodaj-sim', methods=['GET','POST'])
@login_required
def dodaj_sim():
    providers = read_provider()
    if request.method == 'POST':
        provider = request.form.get('provider','').strip()
        serijski = request.form.get('serijski','').strip()
        if not provider:
            flash("Provider je obavezan.", "danger")
            return render_template('dodaj_sim.html', title="Dodaj SIM", username=current_user.username, providers=providers)
        if not re.fullmatch(r'[A-Za-z0-9\-_/\.]+', serijski or ''):
            flash("Serijski broj smije sadržavati samo slova i brojeve.", "danger")
            return render_template('dodaj_sim.html', title="Dodaj SIM", username=current_user.username, providers=providers, provider_sel=provider, serijski=serijski)
        lst = read_sim()
        if any(u.get('serijski','').lower() == serijski.lower() for u in lst):
            flash("SIM s istim serijskim brojem već postoji.", "warning")
            return render_template('dodaj_sim.html', title="Dodaj SIM", username=current_user.username, providers=providers, provider_sel=provider, serijski=serijski)
        lst.append({'provider': provider, 'serijski': serijski, 'created_at': datetime.datetime.now().isoformat()})
        write_sim(lst)
        flash("SIM je spremljen.", "success")
        return redirect(url_for('sim'))
    return render_template('dodaj_sim.html', title="Dodaj SIM", username=current_user.username, providers=providers)

@app.route('/api/sim/<serijski>/delete', methods=['POST'])
@login_required
def api_delete_sim(serijski):
    if not current_user.is_superadmin:
        return jsonify({'ok': False, 'error': 'Brisanje dozvoljeno samo superadminu.'}), 403
    lst = read_sim()
    new_lst = [u for u in lst if u.get('serijski') != serijski]
    if len(new_lst) == len(lst):
        return jsonify({'ok': False, 'error': 'SIM nije pronađen.'}), 404
    write_sim(new_lst)
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True})


# -------------------- UREĐAJI RUTE --------------------



@app.route('/api/sim/<path:serijski>/status', methods=['POST'], endpoint='api_set_sim_status')
@login_required
def api_set_sim_status(serijski):
    try:
        payload = request.get_json(silent=True, force=True) or {}
    except Exception:
        payload = {}
    new_status = str(payload.get('status') or '').strip().lower()

    if new_status not in ('aktivan', 'privremeno isključen', 'privremeno iskljucen'):
        return jsonify({'ok': False, 'error': 'Neispravan status.'}), 400

    norm_status = 'privremeno isključen' if 'privremeno' in new_status else 'aktivan'

    # Update main SIM record's status field (but do not remove from sim.json)
    sims = read_sim() or []
    found = False
    for row in sims:
        try:
            if str(row.get('serijski') or '') == str(serijski):
                if row.get('status') != norm_status:
                    row['status'] = norm_status
                found = True
                break
        except Exception:
            continue
    if found:
        try:
            write_sim(sims)
        except Exception:
            pass

    # Maintain privremeno.iskljuceni.sim.JSON mirror:
    priv_list = read_privremeno_iskljuceni_sim() or []
    now_iso = datetime.datetime.now().isoformat()
    if norm_status == 'privremeno isključen':
        if not any(str(r.get('serijski') or '') == str(serijski) for r in priv_list):
            priv_list.append({'serijski': str(serijski), 'since': now_iso})
            write_privremeno_iskljuceni_sim(priv_list)
    else:
        new_list = [r for r in priv_list if str(r.get('serijski') or '') != str(serijski)]
        if len(new_list) != len(priv_list):
            write_privremeno_iskljuceni_sim(new_list)
    # Try to infer client of this SIM and broadcast recomputed client activity
    try:
        akt_sim = read_aktivni_sim() or []
    except Exception:
        akt_sim = []
    owner = ''
    for s in akt_sim:
        try:
            if str(s.get('serijski')) == str(serijski):
                owner = str(s.get('client') or '')
                break
        except Exception:
            pass
    if owner:
        try:
            _broadcast_client_status(owner, _compute_client_active_by_name(owner))
        except Exception:
            pass
    
    return jsonify({'ok': True, 'status': norm_status}), 200


@app.route('/uredjaji')
@login_required
def uredjaji():
    # Sinkronizacija: osiguraj da su svi kreirani/poznati uređaji upisani u uredjaji.json
    base = read_uredjaji()
    try:
        seen = {str(d.get('serijski')) for d in base}
    except Exception:
        seen = set()
    # iz aktivnih kod klijenata (ako postoji helper)
    if 'read_aktivni_uredjaji' in globals():
        try:
            for d in read_aktivni_uredjaji():
                sid = str(d.get('serijski') or '')
                if sid and sid not in seen:
                    base.append({'model': d.get('model',''), 'serijski': sid, 'created_at': datetime.datetime.now().isoformat()})
                    seen.add(sid)
        except Exception:
            pass
    # iz zaduženih na tehničare (ako postoji helper)
    if 'read_zaduzene_uredjaje' in globals():
        try:
            for d in read_zaduzene_uredjaje():
                sid = str(d.get('serijski') or '')
                if sid and sid not in seen:
                    base.append({'model': d.get('model',''), 'serijski': sid, 'created_at': datetime.datetime.now().isoformat()})
                    seen.add(sid)
        except Exception:
            pass
    # zapiši natrag ako je dopunjeno
    try:
        write_uredjaji(base)
    except Exception:
        pass

    items = list(base)

    # --- Namjena lookup strictly from static/uredjaji.JSON ---
    try:
        _namjena_by_sn = {str(d.get('serijski')): str(d.get('namjena','')) for d in base}
    except Exception:
        _namjena_by_sn = {}

    # --- Status & kolone (isti princip kao /sim) ---
    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    # Aktivni kod klijenta (status GREEN + fallback client naziv)
    try:
        aktivni = read_aktivni_uredjaji()
    except Exception:
        aktivni = []
    active_serials = {_norm_sn(d.get('serijski')) for d in aktivni if d.get('active', True)}
    active_client_by_sn = {}
    for d in aktivni:
        try:
            if d.get('active', True) and d.get('client'):
                active_client_by_sn[_norm_sn(d.get('serijski'))] = str(d.get('client'))
        except Exception:
            pass

    # Zaduženi kod tehničara (status YELLOW + popuni 'assigned_to')
    try:
        zaduzeni = read_zaduzene_uredjaje()
    except Exception:
        zaduzeni = []
    tech_by_serial = {_norm_sn(z.get('serijski')): str(z.get('assigned_to') or '') for z in zaduzeni}

    # KLIJENT MAPA iz klijenti.json ("sn_uredjaja") — fallback ako nema u aktivnim
    try:
        klijenti = read_klijenti()
    except Exception:
        klijenti = []
    client_by_serial = {}
    for k in klijenti:
        cname = str(k.get('name') or '').strip()
        try:
            for sn in str(k.get('sn_uredjaja') or '').split(','):
                snn = _norm_sn(sn)
                if snn:
                    client_by_serial[snn] = cname
        except Exception:
            pass

    # Nadopuni iz aktivnih evidencija gdje nema u klijenti.json
    for snn, cname in active_client_by_sn.items():
        client_by_serial.setdefault(snn, cname)
    # Privremeno isključeni uređaji (orange status)
    try:
        _priv = read_privremeno_iskljuceni()
    except Exception:
        _priv = []
    def _norm_sn(v):
        import re as _re
        return _re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()
    priv_set = {_norm_sn(p.get('serijski')) for p in _priv}


    # Dodaj status i kolone za prikaz
    for it in items:
        raw_sn = it.get('serijski','')
        sn = _norm_sn(raw_sn)
        assigned_to = tech_by_serial.get(sn, '')
        is_assigned = bool(assigned_to)
        # Orange override: if device is in privremeno list
        if sn in priv_set:
            it['status_color'] = 'orange'
            it['assigned_to'] = tech_by_serial.get(sn, '')
            it['client_name'] = client_by_serial.get(sn, '')
            # skip standard status calculation
            continue

        # Pravila boja:
        # - GREEN ako je kod klijenta (aktivni ili klijenti.json)
        # - RED ako NIJE zadužen i NIJE kod klijenta
        # - YELLOW u ostalim slučajevima (npr. zadužen kod tehničara)
        if sn in active_serials or sn in client_by_serial:
            it['status_color'] = 'green'
        elif (not is_assigned) and (sn not in active_serials) and (sn not in client_by_serial):
            it['status_color'] = 'red'
        else:
            it['status_color'] = 'yellow'

        it['assigned_to'] = assigned_to
        it['client_name'] = client_by_serial.get(sn, '')
                # Namjena uređaja (isključivo iz static/uredjaji.JSON)
        try:
            it['namjena'] = _namjena_by_sn.get(it.get('serijski',''), '')
        except Exception:
            it['namjena'] = ''


    try:
        items.sort(key=lambda x: (x.get('model',''), x.get('serijski','')))
    except Exception:
        pass

    # --- Brojači za prikaz na /uredjaji ---
    aktivni_count = sum(1 for _it in items if _it.get('status_color') == 'green')
    neaktivni_count = sum(1 for _it in items if _it.get('status_color') == 'red')
    zaduzeni_count = sum(1 for _it in items if str(_it.get('assigned_to') or '').strip())
    privremeni_count = sum(1 for _it in items if _it.get('status_color') == 'orange')
    return render_template('uredjaji.html', title="Uređaji", username=current_user.username, uredjaji=items,
                           aktivni_count=aktivni_count, zaduzeni_count=zaduzeni_count,
                           privremeni_count=privremeni_count, neaktivni_count=neaktivni_count)


# -------------------- EXPORT UREĐAJA (XLSX) --------------------
@app.route('/uredjaji/export.xlsx')
@login_required
def export_uredjaji_xlsx():
    """
    Preuzimanje XLSX tablice uređaja sa kolonama: Model, Serijski broj, Klijent.
    Generira se ista lista kao na stranici /uredjaji.
    """
    base = read_uredjaji()
    items = list(base)

    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    try:
        aktivni = read_aktivni_uredjaji()
    except Exception:
        aktivni = []
    active_serials = {_norm_sn(d.get('serijski')) for d in aktivni if d.get('active', True)}
    active_client_by_sn = {}
    for d in aktivni:
        try:
            if d.get('active', True) and d.get('client'):
                active_client_by_sn[_norm_sn(d.get('serijski'))] = str(d.get('client'))
        except Exception:
            pass

    try:
        zaduzeni = read_zaduzene_uredjaje()
    except Exception:
        zaduzeni = []
    tech_by_serial = {_norm_sn(z.get('serijski')): str(z.get('assigned_to') or '') for z in zaduzeni}

    try:
        klijenti = read_klijenti()
    except Exception:
        klijenti = []
    client_by_serial = {}
    for k in klijenti:
        cname = str(k.get('name') or '').strip()
        try:
            for sn in str(k.get('sn_uredjaja') or '').split(','):
                snn = _norm_sn(sn)
                if snn:
                    client_by_serial[snn] = cname
        except Exception:
            pass

    for snn, cname in active_client_by_sn.items():
        client_by_serial.setdefault(snn, cname)

    for it in items:
        raw_sn = it.get('serijski','')
        sn = _norm_sn(raw_sn)
        assigned_to = tech_by_serial.get(sn, '')
        is_assigned = bool(assigned_to)
        if sn in active_serials or sn in client_by_serial:
            it['status_color'] = 'green'
        elif (not is_assigned) and (sn not in active_serials) and (sn not in client_by_serial):
            it['status_color'] = 'red'
        else:
            it['status_color'] = 'yellow'
        it['assigned_to'] = assigned_to
        it['client_name'] = client_by_serial.get(sn, '')

    try:
        items.sort(key=lambda x: (x.get('model',''), x.get('serijski','')))
    except Exception:
        pass

    wb = Workbook()
    ws = wb.active
    ws.title = "Uređaji"
    ws.append(["Model", "Serijski broj", "Namjena uređaja"])
    for it in items:
        ws.append([str(it.get('model', '') or ''), str(it.get('serijski', '') or '')])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="uredjaji.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# -------------------- IMPORT UREĐAJA (XLSX) --------------------
@app.route('/uredjaji/import', methods=['POST'])
@login_required
def import_uredjaji_xlsx():
        # Dozvoljen upload samo superadmin/admin/serviser
    if not has_role('superadmin','admin','serviser'):
        flash('Nemate ovlasti za uvoz uređaja.', 'danger')
        return redirect(url_for('uredjaji'))

# Upload samo dva stupca: Model i Serijski broj.
    file = request.files.get('file')
    if not file or not file.filename.lower().endswith('.xlsx'):
        flash("Odaberite .xlsx datoteku.", "danger")
        return redirect(url_for('uredjaji'))
    try:
        from openpyxl import load_workbook
    except Exception:
        flash("Nedostaje openpyxl modul.", "danger")
        return redirect(url_for('uredjaji'))
    try:
        wb = load_workbook(file, read_only=True, data_only=True)
        ws = wb.active
    except Exception as e:
        flash(f"Ne mogu pročitati XLSX: {e}", "danger")
        return redirect(url_for('uredjaji'))

    # Zaglavlje (očekuje se 'Model' i 'Serijski broj')
    try:
        header_cells = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        header = [str(x or '').strip() for x in header_cells]
    except Exception:
        header = []

    def _find_idx(name):
        name = str(name).strip().lower()
        for i, h in enumerate(header):
            if str(h).strip().lower() == name:
                return i
        return -1

    idx_model = _find_idx('model')
    idx_sn = _find_idx('serijski broj')
    idx_namjena = _find_idx('namjena uređaja')
    if idx_namjena == -1:
        # fallback: allow 'namjena' as header too
        idx_namjena = _find_idx('namjena')
    if idx_model == -1 or idx_sn == -1 or idx_namjena == -1:
        flash("Zaglavlje mora imati točno stupce: 'Model', 'Serijski broj' i 'Namjena uređaja'.", "danger")
        return redirect(url_for('uredjaji'))
    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    # Skupljanje postojećih SN-ova iz svih evidencija
    existing = set()
    try:
        base = read_uredjaji()
        existing |= {_norm_sn(x.get('serijski')) for x in base}
    except Exception:
        base = []
    try:
        if 'read_aktivni_uredjaji' in globals():
            existing |= {_norm_sn(x.get('serijski')) for x in (read_aktivni_uredjaji() or [])}
    except Exception:
        pass
    try:
        if 'read_zaduzene_uredjaje' in globals():
            existing |= {_norm_sn(x.get('serijski')) for x in (read_zaduzene_uredjaje() or [])}
    except Exception:
        pass
    try:
        if 'read_servis_uredjaji' in globals():
            existing |= {_norm_sn(x.get('serijski')) for x in (read_servis_uredjaji() or [])}
    except Exception:
        pass


    added = 0
    skipped_invalid = 0
    skipped_dup_in_file = 0
    skipped_existing = 0
    seen_in_file = set()
    new_rows = []

    for r in ws.iter_rows(min_row=2, values_only=True):
        model = str((r[idx_model] if idx_model < len(r) else '') or '').strip()
        sn_raw = str((r[idx_sn] if idx_sn < len(r) else '') or '').strip()
        namjena_raw = str((r[idx_namjena] if idx_namjena < len(r) else '') or '').strip()
        if not sn_raw:
            skipped_invalid += 1
            continue
        if not namjena_raw:
            skipped_invalid += 1
            continue
        namjena_norm = namjena_raw.strip().lower()
        if namjena_norm in ('kupnja','kupnja.','kupnja,'):
            namjena_norm = 'Kupnja'
        elif namjena_norm in ('najam','najam.','najam,'):
            namjena_norm = 'Najam'
        else:
            # try capitalize first letter
            nm_try = namjena_raw.capitalize()
            if nm_try in ('Kupnja','Najam'):
                namjena_norm = nm_try
            else:
                skipped_invalid += 1
                continue
        if not re.fullmatch(r'[A-Za-z0-9\-_/\.]+', sn_raw):
            skipped_invalid += 1
            continue
        sn_norm = _norm_sn(sn_raw)
        if sn_norm in seen_in_file:
            skipped_dup_in_file += 1
            continue
        seen_in_file.add(sn_norm)
        if sn_norm in existing:
            skipped_existing += 1
            continue
        new_rows.append({'model': model, 'serijski': sn_raw, 'namjena': namjena_norm, 'created_at': datetime.datetime.now().isoformat()})
        existing.add(sn_norm)
        added += 1

    if new_rows:
        base.extend(new_rows)
        try:
            write_uredjaji(base)
        except Exception as e:
            flash(f"Greška pri spremanju: {e}", "danger")
            return redirect(url_for('uredjaji'))

    msg = f"Uvoz dovršen. Dodano: {added}, već postoji: {skipped_existing}, duplikat u datoteci: {skipped_dup_in_file}, neispravno: {skipped_invalid}."
    flash(msg, "success" if added else "warning")
    return redirect(url_for('uredjaji'))

# --- BEGIN servis link redirects ---
@app.route('/servis')
@login_required
def _servis_redirect():
    # Preusmjeri stari/izbornik link na novu rutu
    return redirect(url_for('servis_uredjaja'))

@app.route('/servis_uredjaja')
@login_required
def _servis_redirect_alt():
    # Alternativna varijanta (podcrtano) -> preusmjeri
    return redirect(url_for('servis_uredjaja'))
# --- END servis link redirects ---

@app.route('/servis-uredjaja')
@login_required
def servis_uredjaja():
    """Lista uređaja na servisu i otpisanih uređaja."""
    try:
        uredjaji = read_servis_uredjaji() or []
    except Exception:
        uredjaji = []
    try:
        otpisani = read_otpisani_uredjaji() or []
    except Exception:
        otpisani = []
    # stabilan prikaz
    try:
        uredjaji.sort(key=lambda x: (str(x.get('client','')), str(x.get('model','')), str(x.get('serijski',''))))
        otpisani.sort(key=lambda x: (str(x.get('client','')), str(x.get('model','')), str(x.get('serijski',''))))
    except Exception:
        pass
    return render_template('servis_uredjaja.html', uredjaji=uredjaji, otpisani=otpisani)


@app.route('/kupljeni.uredjaji')
@login_required
def kupljeni_uredjaji():
    # osvježi kupljeni.uredjaji.JSON na temelju aktivnih uređaja s namjenom 'Kupnja'
    items = sync_kupljeni_from_aktivni() or []
    # Dodaj statusnu kolonu prema atributima (active/client/assigned_to)
    def _status(it):
        try:
            if it.get('active') is True:
                return 'Aktivan'
            if str(it.get('assigned_to') or '').strip():
                return 'Zadužen tehničaru'
            if str(it.get('client') or '').strip():
                return 'Kod klijenta'
        except Exception:
            pass
        return 'Neaktivan'
    for it in items:
        it['Status uređaja'] = _status(it)

    # Pripremi dinamičke kolone: sve ključeve iz aktivni_uredjaji zapisa + "Status uređaja" na kraju
    columns = []
    if items:
        for k in list(items[0].keys()):
            if k not in columns and k != 'Status uređaja':
                columns.append(k)
        for it in items:
            for k in it.keys():
                if k not in columns and k != 'Status uređaja':
                    columns.append(k)
        columns.append('Status uređaja')

    # stabilan prikaz: sort po klijentu, modelu, serijskom
    try:
        items.sort(key=lambda x: (str(x.get('client','')), str(x.get('model','')), str(x.get('serijski',''))))
    except Exception:
        pass
    return render_template('kupljeni.uredjaji.html', title='Kupljeni uređaji', items=items, columns=columns, username=getattr(current_user, 'username', None))


@app.route('/dodaj_uredjaj', methods=['GET','POST'])

@app.route('/dodaj-uredjaj', methods=['GET','POST'])
@login_required
def dodaj_uredjaj():
    if not has_role('superadmin','admin','serviser'):
        flash('Nemate ovlasti za ovu radnju.', 'danger')
        return redirect(url_for('index'))

    models = read_nazivi_uredjaja()
    if request.method == 'POST':
        model = request.form.get('model','').strip()
        serijski = request.form.get('serijski','').strip()
        namjena = request.form.get('namjena','').strip()
        if not namjena:
            flash("Namjena uređaja je obavezna.", "danger")
            return render_template('dodaj_uredjaj.html', title="Dodaj uređaj", username=getattr(current_user, 'username', None), models=models, namjene=(list(read_namjena_uredjaja().values()) if isinstance(read_namjena_uredjaja(), dict) else list(read_namjena_uredjaja() or [])) or ['Kupnja','Najam'])

        if not model:
            flash("Model uređaja je obavezan.", "danger")
            return render_template('dodaj_uredjaj.html', title="Dodaj uređaj", username=getattr(current_user, 'username', None), models=models, namjene=(list(read_namjena_uredjaja().values()) if isinstance(read_namjena_uredjaja(), dict) else list(read_namjena_uredjaja() or [])) or ['Kupnja','Najam'])
        if not re.fullmatch(r'[A-Za-z0-9\-_/\.]+', serijski or ''):
            flash("Serijski broj smije sadržavati samo slova i brojeve (dozvoljeni su - _ / .).", "danger")
            return render_template('dodaj_uredjaj.html', title="Dodaj uređaj", username=getattr(current_user, 'username', None), models=models, namjene=(list(read_namjena_uredjaja().values()) if isinstance(read_namjena_uredjaja(), dict) else list(read_namjena_uredjaja() or [])) or ['Kupnja','Najam'])
        lst = read_uredjaji()
        if any(u.get('serijski','').lower() == serijski.lower() for u in lst):
            flash("Uređaj s istim serijskim brojem već postoji.", "warning")
            return render_template('dodaj_uredjaj.html', title="Dodaj uređaj", username=getattr(current_user, 'username', None), models=models, namjene=(list(read_namjena_uredjaja().values()) if isinstance(read_namjena_uredjaja(), dict) else list(read_namjena_uredjaja() or [])) or ['Kupnja','Najam'])
        lst.append({'model': model, 'serijski': serijski, 'namjena': namjena, 'created_at': datetime.datetime.now().isoformat()})
        write_uredjaji(lst)
        flash("Uređaj je spremljen.", "success")
        try:
            append_device_log(serijski, 'kreiran', details={'model': model, 'namjena': namjena})
        except Exception:
            pass
        return redirect(url_for('uredjaji'))
    return render_template('dodaj_uredjaj.html', title="Dodaj uređaj", username=getattr(current_user, 'username', None), models=models, namjene=(list(read_namjena_uredjaja().values()) if isinstance(read_namjena_uredjaja(), dict) else list(read_namjena_uredjaja() or [])) or ['Kupnja','Najam'])

@app.route('/api/uredjaj/<path:serijski>/delete', methods=['POST'])
@login_required
def api_delete_uredjaj(serijski):
    """Potpuno uklanja uređaj iz svih JSON evidencija.
    Uklanja iz:
      - uredjaji.JSON (glavna baza)
      - aktivni_uredjaji.JSON (kod klijenta)
      - zaduzeni.uredjaji.json (kod tehničara)
      - servis.uredjaji.JSON (uređaji na servisu)
      - privremeno.iskljuceni.uredjaji.JSON (lista privremeno isključenih)
      - klijenti.json (polja sn_uredjaja i model_uredjaja – čisti SN i rekalkulira modele)
      - otpisani.uredjaji.JSON (ako postoji isti SN)
    Napomena: samo superadmin smije brisati.
    """
    if not current_user.is_superadmin:
        return jsonify({'ok': False, 'error': 'Brisanje dozvoljeno samo superadminu.'}), 403

    sn_target = str(serijski)

    # Helper: normalize SN for comparison across sources
    def _norm_sn(v):
        import re as _re
        return _re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    target_norm = _norm_sn(sn_target)
    removed_any = False

    # 1) Glavna baza uredjaji.JSON
    try:
        base = read_uredjaji() or []
    except Exception:
        base = []
    new_base = [u for u in base if _norm_sn(u.get('serijski')) != target_norm]
    if len(new_base) != len(base):
        write_uredjaji(new_base)
        removed_any = True

    # 2) Aktivni kod klijenta
    try:
        akt = read_aktivni_uredjaji() or []
    except Exception:
        akt = []
    new_akt = [d for d in akt if _norm_sn(d.get('serijski')) != target_norm]
    if len(new_akt) != len(akt):
        write_aktivni_uredjaji(new_akt)
        removed_any = True

    # 3) Zaduženi kod tehničara
    try:
        zad = read_zaduzene_uredjaje() or []
    except Exception:
        zad = []
    new_zad = [z for z in zad if _norm_sn(z.get('serijski')) != target_norm]
    if len(new_zad) != len(zad):
        write_zaduzene_uredjaje(new_zad)
        removed_any = True

    # 4) Servis uređaji
    try:
        srv = read_servis_uredjaji() or []
    except Exception:
        srv = []
    new_srv = [s for s in srv if _norm_sn(s.get('serijski')) != target_norm]
    if len(new_srv) != len(srv):
        write_servis_uredjaji(new_srv)
        removed_any = True

    # 5) Privremeno isključeni uređaji
    try:
        prv = read_privremeno_iskljuceni() or []
    except Exception:
        prv = []
    new_prv = [p for p in prv if _norm_sn(p.get('serijski')) != target_norm]
    if len(new_prv) != len(prv):
        write_privremeno_iskljuceni(new_prv)
        removed_any = True

    # 6) Otpisani uređaji (ako se tamo eventualno nalazi)
    try:
        otp = read_otpisani_uredjaji() or []
    except Exception:
        otp = []
    new_otp = [o for o in otp if _norm_sn(o.get('serijski')) != target_norm]
    if len(new_otp) != len(otp):
        write_otpisani_uredjaji(new_otp)
        removed_any = True

    # 7) Ukloni iz klijenti.json (SN + rekalkulacija modela za preostale SN-ove)
    try:
        kl = read_klijenti() or []
    except Exception:
        kl = []
    changed_clients = False
    for c in kl:
        try:
            sn_list = [s.strip() for s in str(c.get('sn_uredjaja') or '').split(',') if s.strip()]
            sn_list_norm = [_norm_sn(s) for s in sn_list]
            if target_norm in sn_list_norm:
                sn_list = [s for s in sn_list if _norm_sn(s) != target_norm]
                c['sn_uredjaja'] = ", ".join(sn_list)

                # rekalkuliraj modele iz aktivnih + glavne baze (nakon uklanjanja)
                model_map = {}
                try:
                    for d in (new_akt or []):
                        if str(d.get('client') or '') == str(c.get('name') or ''):
                            model_map[str(d.get('serijski'))] = d.get('model','')
                except Exception:
                    pass
                try:
                    for d in (new_base or []):
                        model_map.setdefault(str(d.get('serijski')), d.get('model',''))
                except Exception:
                    pass

                models = []
                for s in sn_list:
                    m = model_map.get(s, '')
                    if m and m not in models:
                        models.append(m)
                c['model_uredjaja'] = ", ".join(models)
                changed_clients = True
        except Exception:
            pass
    if changed_clients:
        write_klijenti(kl)

    # Lagani log (ako helper postoji)
    try:
        append_device_log(sn_target, 'obrisan', details={'by': getattr(current_user, 'username', None)})
    except Exception:
        pass

    if not removed_any and not changed_clients:
        return jsonify({'ok': False, 'error': 'Uređaj nije pronađen.'}), 404

    # Reizračun statusa klijenta kojem je uređaj možda pripadao
    try:
        for c in kl or []:
            nm = str(c.get('name') or '')
            if nm:
                _broadcast_client_status(nm, _compute_client_active_by_name(nm))
    except Exception:
        pass

    return jsonify({'ok': True}), 200




# -------------------- NALOZI – brisanje (samo superadmin) --------------------
@app.route('/api/nalog/<path:rn>/delete', methods=['POST'])
@login_required
def api_delete_nalog(rn):
    if not current_user.is_superadmin:
        return jsonify({'ok': False, 'error': 'Brisanje dozvoljeno samo superadminu.'}), 403
    lst = read_nalozi()
    new_lst = [n for n in lst if n.get('rn') != rn]
    if len(new_lst) == len(lst):
        return jsonify({'ok': False, 'error': 'Nalog nije pronađen.'}), 404
    write_nalozi(new_lst)
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True})


# -------------------- SEED SUPERADMIN --------------------
def ensure_superadmin():
    db.create_all()
    sa = User.query.filter_by(username="hlisac").first()
    if not sa:
        sa = User(username="hlisac", role="superadmin")
        sa.set_password("123")
        db.session.add(sa)
        db.session.commit()
        print("[seed] Superadmin kreiran: hlisac/123")
    else:
        changed = False
        if sa.role != "superadmin":
            sa.role = "superadmin"; changed = True
        if not sa.check_password("123"):
            sa.set_password("123"); changed = True
        if changed:
            db.session.commit()
            print("[seed] Superadmin ažuriran: hlisac/123")

# -------------------- UTIL --------------------
# -------------------- PODSJETNICI (e-mail scheduler) --------------------
REMINDERS_JSON_PATH = os.path.join(STATIC_DIR, 'reminders.json')

def read_reminders():
    try:
        data = _read_json(REMINDERS_JSON_PATH)
        if isinstance(data, list):
            return data
        return []
    except Exception:
        return []

def write_reminders(data):
    try:
        _write_json(REMINDERS_JSON_PATH, data or [])
    except Exception:
        pass

def _find_operator_email(username: str) -> str:
    """Pokušaj dohvatiti e-mail operatera iz static/operateri.json prema username-u."""
    try:
        ops = read_operateri() or []
    except Exception:
        ops = []
    uname = str(username or '').strip().lower()
    for o in ops:
        try:
            if str(o.get('username','')).strip().lower() == uname:
                em = str(o.get('email') or o.get('mail') or '').strip()
                if em:
                    return em
        except Exception:
            pass
    return ''

def _send_mail_now(to_email: str, subject: str, body: str) -> bool:
    try:
        with app.app_context():
            msg = Message(subject=subject, recipients=[to_email])
            msg.body = body
            mail.send(msg)
        return True
    except Exception as e:
        try:
            print('[reminder] Slanje e-maila nije uspjelo:', e)
        except Exception:
            pass
        return False

def _reminder_loop():
    import time
    while True:
        try:
            now = datetime.datetime.now()
            items = read_reminders()
            changed = False
            for it in items:
                try:
                    if it.get('sent'):
                        continue
                    # when je ISO bez timezonea iz <input type="datetime-local">
                    dt_str = str(it.get('when') or '')
                    if not dt_str:
                        continue
                    # podrži varijante "YYYY-MM-DDTHH:MM" i s sekundama
                    try:
                        send_at = datetime.datetime.fromisoformat(dt_str)
                    except Exception:
                        # fallback: strip everything after minutes
                        send_at = datetime.datetime.strptime(dt_str[:16], '%Y-%m-%dT%H:%M')
                    if now >= send_at:
                        ok = _send_mail_now(it.get('to'), it.get('subject','Podsjetnik'), it.get('body',''))
                        it['sent'] = bool(ok)
                        it['sent_at'] = datetime.datetime.now().isoformat()
                        changed = True
                except Exception:
                    pass
            if changed:
                write_reminders(items)
        except Exception:
            pass
        time.sleep(30)  # provjera svakih 30s

# Start background thread once
try:
    import threading
    if not globals().get('_REMINDER_THREAD_STARTED'):
        t = threading.Thread(target=_reminder_loop, name='ReminderLoop', daemon=True)
        t.start()
        _REMINDER_THREAD_STARTED = True
except Exception as e:
    try:
        print('[reminder] Background thread start failed:', e)
    except Exception:
        pass

@app.route('/api/reminders', methods=['POST'])
@login_required
def api_create_reminder():
    try:
        payload = request.get_json(silent=True, force=True) or {}
    except Exception:
        payload = {}
    when = str(payload.get('when') or '').strip()
    client = str(payload.get('client') or '').strip() or 'Nepoznat klijent'
    if not when:
        return jsonify({'ok': False, 'error': 'Nedostaje datum/vrijeme.'}), 400
    # Operator e-mail prema prijavljenom korisniku
    to_email = _find_operator_email(getattr(current_user, 'username', None))
    if not to_email:
        return jsonify({'ok': False, 'error': 'Nije pronađen e-mail operatera za slanje podsjetnika.'}), 400
    body = f"Bok, šaljem podsjetnik da se kontaktira ili izvrši zadatak za klijenta {client}!"
    item = {
        'id': f"r-{int(datetime.datetime.now().timestamp()*1000)}",
        'to': to_email,
        'when': when,
        'subject': 'Podsjetnik',
        'body': body,
        'created_by': getattr(current_user, 'username', None),
        'created_at': datetime.datetime.now().isoformat(),
        'sent': False
    }
    items = read_reminders()
    items.append(item)
    write_reminders(items)
    return jsonify({'ok': True, 'created': item}), 200

def slugify(value: str) -> str:
    value = unicodedata.normalize('NFKD', value).encode('ascii', 'ignore').decode('ascii')
    value = re.sub(r'[^a-zA-Z0-9_-]+', '-', value).strip('-').lower()
    return value or 'klijent'

from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement



def set_table_uniform_borders(tbl, size=12, color='000000', space=0):
    """Postavi vidljive linije (borders) oko SVIH ćelija tablice.
    size je u 1/8 pt (npr. 12 = 1.5pt), color = '000000' crna.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    def _set_border(tc, tag):
        tcPr = tc._tc.get_or_add_tcPr()
        # pronađi ili kreiraj <w:tcBorders>
        borders = None
        for ch in tcPr.iterchildren():
            if ch.tag == qn('w:tcBorders'):
                borders = ch
                break
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tcPr.append(borders)
        el = None
        for ch in borders.iterchildren():
            if ch.tag == qn(f'w:{tag}'):
                el = ch
                break
        if el is None:
            el = OxmlElement(f'w:{tag}')
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(int(size)))
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), str(int(space)))

    for row in tbl.rows:
        for cell in row.cells:
            _set_border(cell, 'top')
            _set_border(cell, 'left')
            _set_border(cell, 'bottom')
            _set_border(cell, 'right')
def set_tbl_cell_margins(tbl, top=80, left=120, bottom=80, right=120):
    """Postavi unutarnje margine (padding) u twips za sve ćelije tablice.
    Default: top/bottom 80 (~5.6pt), left/right 120 (~8.4pt).
    """
    tbl_pr = tbl._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tbl_pr)

    # w:tblCellMar container
    tbl_cell_mar = None
    for el in tbl_pr.iterchildren():
        if el.tag == qn('w:tblCellMar'):
            tbl_cell_mar = el
            break
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(tbl_cell_mar)

    def _set(side, val):
        el = None
        for c in tbl_cell_mar.iterchildren():
            if c.tag == qn(f'w:{side}'):
                el = c
                break
        if el is None:
            el = OxmlElement(f'w:{side}')
            tbl_cell_mar.append(el)
        el.set(qn('w:w'), str(int(val)))
        el.set(qn('w:type'), 'dxa')

    _set('top', top)
    _set('left', left)
    _set('bottom', bottom)
    _set('right', right)

def remove_placeholder_paragraphs(doc: Document, tokens):
    """Ukloni paragrafe IZVAN tablica koji sadrže bilo koji od navedenih placeholdera.
    Time sprječavamo 'prazne popise' ispod tablice.
    """
    for p in list(doc.paragraphs):
        try:
            if any(tok in p.text for tok in tokens):
                p._element.getparent().remove(p._element)
        except Exception:
            # ako ne uspije, samo preskoči
            pass

def drop_empty_numbered_paragraphs(doc: Document):
    """Ukloni prazne list-item paragrafe (bullets/numbering) izvan tablica."""
    for p in list(doc.paragraphs):
        try:
            txt = (p.text or '').strip()
            # Detektiraj numerirani/bullet paragraf (ima numPr u pPr)
            pPr = getattr(p._p, "pPr", None)
            has_num = bool(pPr is not None and getattr(pPr, "numPr", None) is not None)
            if has_num and txt == "":
                p._element.getparent().remove(p._element)
        except Exception:
            pass


def find_table_with_header(doc: Document, header_cells):
    for tbl in doc.tables:
        if len(tbl.rows) > 0:
            headers = [c.text.strip() for c in tbl.rows[0].cells]
            if all(any(h.lower() in cell.lower() for cell in headers) for h in header_cells):
                return tbl
    return None

def replace_text_in_doc(doc: Document, mapping: dict):
    # Operiraj na cijelom paragrafu/cell tekstu da placeholderi rade i kad su razlomljeni u runovima
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for k, v in mapping.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)




def process_hitna_in_doc(doc: Document, is_checked: bool, emergency_label: str):
    """
    U rn_templateu linija može biti u odvojenim paragrafima, ponekad i u tabličnim ćelijama.
    Ne diramo cell.text izravno (to flattena format), već prolazimo kroz pojedine paragrafe.
    Pravila:
      - is_checked: "* {{hitna}}" ili "{{hitna}}" ili red sa SAMO "*" -> "☑ " + emergency_label
      - not is_checked: paragrafe koji sadrže "{{hitna}}" ili su samo "*" -> ukloniti tekst (postaviti na empty)
    """
    import re
    def fix_line(text: str) -> str:
        if is_checked:
            # zamjena placeholdera
            t = re.sub(r"\*\s*\{\{hitna\}\}", "☑ " + emergency_label, text)
            t = t.replace("{{hitna}}", "☑ " + emergency_label)
            # ako je linija samo zvjezdica -> checkbox + label
            if re.fullmatch(r"[\ \t]*\*[\ \t]*", t):
                t = "☑ " + emergency_label
            return t
        else:
            # izbaci red s hitna placeholderom i redove koji su samo zvjezdica
            if "{{hitna}}" in text or re.fullmatch(r"[\ \t]*\*[\ \t]*", text):
                return ""
            return text

    # Paragrafi izvan tablica
    for p in list(doc.paragraphs):
        new_t = fix_line(p.text)
        if new_t != p.text:
            p.text = new_t

    # Paragrafi unutar tablica
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    new_t = fix_line(p.text)
                    if new_t != p.text:
                        p.text = new_t


def next_rn_for_year(nalozi, year):
    # Broji samo instalacije u tekućoj godini i vrati sljedeći redni broj
    nums = []
    for n in nalozi:
        if n.get('type') == 'instalacija':
            try:
                dt = datetime.datetime.fromisoformat(n.get('created_at','')[:19])
            except Exception:
                continue
            if dt.year == year and 'rn' in n:
                # rn format 0001/2025
                m = re.match(r'(\d{4})/(\d{4})', n['rn'])
                if m and int(m.group(2)) == year:
                    nums.append(int(m.group(1)))
    return (max(nums) + 1) if nums else 1


def next_rn_for_year_deinst(nalozi, year):
    """Broji samo deinstalacije u tekućoj godini i vraća sljedeći redni broj (0001/YEAR)."""
    nums = []
    for n in nalozi:
        if (n.get('type') or '').lower() == 'deinstalacija':
            try:
                dt = datetime.datetime.fromisoformat(str(n.get('created_at',''))[:19])
            except Exception:
                continue
            if dt.year == year and 'rn' in n:
                m = re.match(r'(\d{4})/(\d{4})', str(n['rn']))
                if m and int(m.group(2)) == year:
                    try:
                        nums.append(int(m.group(1)))
                    except Exception:
                        pass
    return (max(nums) + 1) if nums else 1

def next_rn_continuous_deinst(nalozi):
    """
    Kontinuirani redni broj za deinstalacije (bez resetiranja po godinama).
    Vraća sljedeći broj (1-based) prema najvećem do sada zabilježenom u nalozi.json za tip 'deinstalacija'.
    """
    nums = []
    for n in nalozi or []:
        if (str(n.get('type') or '').lower().strip() == 'deinstalacija') and n.get('rn'):
            m = re.match(r'(\d{4})/(\d{4})', str(n['rn']))
            if m:
                try:
                    nums.append(int(m.group(1)))
                except Exception:
                    pass
    return (max(nums) + 1) if nums else 1







def next_rn_continuous_servis(nalozi):
    '''
    Kontinuirani redni broj za servisne naloge (bez resetiranja po godinama).
    Vraća sljedeći broj (1-based) prema najvećem do sada zabilježenom u nalozi.json za tip 'servis'.
    '''
    nums = []
    for n in nalozi or []:
        if (str(n.get('type') or '').lower().strip() == 'servis') and n.get('rn'):
            m = re.match(r'(\d{4})/(\d{4})', str(n['rn']))
            if m:
                try:
                    nums.append(int(m.group(1)))
                except Exception:
                    pass
    return (max(nums) + 1) if nums else 1

def next_rn_for_year_servis(nalozi, year):
    """Broji samo servisne naloge u tekućoj godini i vraća sljedeći redni broj (0001/YEAR)."""
    nums = []
    for n in nalozi:
        if (n.get('type') or '').lower() == 'servis':
            try:
                dt = datetime.datetime.fromisoformat(str(n.get('created_at',''))[:19])
            except Exception:
                continue
            if dt.year == year and 'rn' in n:
                m = re.match(r'(\d{4})/(\d{4})', str(n['rn']))
                if m and int(m.group(2)) == year:
                    try:
                        nums.append(int(m.group(1)))
                    except Exception:
                        pass
    return (max(nums) + 1) if nums else 1
def generate_deinstalacija_docx(klijent, selected_devices=None, selected_sims=None, napomena: str = "", save_in_zapisnici=False, skip_nalozi=False):
    # Template path for deinstalacija
    template_path = os.path.join(STATIC_DIR, 'datoteke', 'rn_template_deinstalacija.docx')
    if not os.path.exists(template_path):
        abort(500, description="Nedostaje rn_template_deinstalacija.docx u static/datoteke")

    doc = Document(template_path)
    # PRE-CLEAN: ukloni moguće placeholder paragrafe izvan tablica da ne ostane 'prazan popis'
    try:
        remove_placeholder_paragraphs(doc, ['{{Uređaj}}', '{{serijski_broj_uredjaja}}', '{{serijski_broj_sim}}', '{{SIM kartice}}'])
    except Exception:
        pass

    korisnik = str(klijent.get('name',''))
    oib = str(klijent.get('oib',''))
    adresa = str(klijent.get('shipping') or klijent.get('headquarters') or '')
    kontakt = str(klijent.get('phone',''))

    # RN za tekuću godinu (odvojeni slijed za deinstalacije)
    nalozi = read_nalozi()
    year = datetime.datetime.now().year
    broj = next_rn_continuous_deinst(nalozi)
    rn_str = f"{broj:04d}/{year}"

    # Placeholder mapping
    mapping = {
        '{{RN}}': rn_str,
                '{{0001/2025}}': rn_str,
        '{{KORISNIK}}': korisnik,
        '{{ADRESA_ISPORUKE}}': adresa,
        '{{OIB}}': oib,
        '{{KONTAKT}}': kontakt,
    }



    # --- DODANO: Napomena + placeholderi za uređaje/SIM ---
    napomena_text = str(napomena or "").strip()
    if napomena_text:
        mapping['{{Napomena}}'] = napomena_text

    # Pokušaj izvući mapiranje SN->Model iz klijenta i aktivnih evidencija (točnije odabire sa /deinstalacija/)
    _dev_model_map = {}
    try:
        sn_list_tmp = [s.strip() for s in str(klijent.get('sn_uredjaja','')).split(',') if s.strip()]
    except Exception:
        sn_list_tmp = []
    try:
        model_list_tmp = [s.strip() for s in str(klijent.get('model_uredjaja','')).split(',') if s.strip()]
    except Exception:
        model_list_tmp = []

    for i, snv in enumerate(sn_list_tmp):
        if i < len(model_list_tmp):
            _dev_model_map[snv] = model_list_tmp[i]

    # Upotpuni mapu i iz aktivnih uređaja kod klijenta (ako postoji helper)
    try:
        _akt = read_aktivni_uredjaji()
    except Exception:
        _akt = []
    for d in _akt or []:
        snv = str(d.get('serijski') or '').strip()
        mdl = str(d.get('model') or '').strip()
        if snv and mdl:
            _dev_model_map.setdefault(snv, mdl)

    # Odredi prvi odabrani uređaj i njegov model (za singularni placeholder u headeru predloška)
    _first_dev_sn = (selected_devices[0] if selected_devices else "")
    _first_dev_model = _dev_model_map.get(_first_dev_sn, "")

    # Popuni singularne placeholdere: Uređaj = MODEL odabranog uređaja, Serijski = SN
    mapping.setdefault('{{Uređaj}}', _first_dev_model or '')
    mapping.setdefault('{{serijski_broj_uredjaja}}', _first_dev_sn or '')

    # Spriječi dupliranje SIM informacija ispod tablice – singularne SIM placeholdere praznimo,
    # a SIM-ovi se isključivo upisuju u retke tablice niže.
    mapping['{{SIM kartice}}'] = ''
    mapping['{{serijski_broj_sim}}'] = ''

    replace_text_in_doc(doc, mapping)

    # Odredi redove za tablicu iz klijenta; filtriraj ako su selektirani
    sn_list = [s.strip() for s in str(klijent.get('sn_uredjaja','')).split(',') if s.strip()]
    model_list = [s.strip() for s in str(klijent.get('model_uredjaja','')).split(',') if s.strip()]
    sim_list = [s.strip() for s in str(klijent.get('sn_SIM','')).split(',') if s.strip()]

    rows = []
    if selected_devices:
        sn_filter = set(selected_devices)
        for sn in sn_list:
            if sn in sn_filter:
                md = model_list[sn_list.index(sn)] if sn in sn_list and sn_list.index(sn) < len(model_list) else ''
                sims = sim_list[sn_list.index(sn)] if sn in sn_list and sn_list.index(sn) < len(sim_list) else ''
                rows.append((sn, md, sims))
    else:
        max_len = max(len(sn_list), len(model_list), len(sim_list) if sim_list else 0)
        for i in range(max_len):
            sn = sn_list[i] if i < len(sn_list) else ''
            md = model_list[i] if i < len(model_list) else ''
            sims = sim_list[i] if i < len(sim_list) else ''
            rows.append((sn, md, sims))



    
    # Nađi tablicu (pokušaj po zaglavlju 'Naziv opreme' ili 5-col RN), inače prva tablica
    tbl = None
    for headers in (
        ['Rb', 'Naziv', 'TID', 'Serijski', 'OTP'],
        ['Rb.', 'Naziv Opreme', 'TID', 'Serijski broj', 'OTP'],
        ['Naziv Opreme', 'Serijski'],
        ['Naziv', 'opreme'],
        ['Naziv opreme'],
        ['Model','SN','SIM']
    ):
        tbl = find_table_with_header(doc, headers)
        if tbl:
            break
    if not tbl and doc.tables:
        tbl = doc.tables[0]
    if not tbl:
        # Ako u predlošku nema očekivane tablice, kreiraj novu s 5 stupaca
        tbl = doc.add_table(rows=1, cols=5)
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = 'Rb.', 'Naziv Opreme', 'TID', 'Serijski broj', 'OTP'

    if tbl:
        # Odredi indekse kolona po nazivu (case-insensitive, djelomično poklapanje)
        header_cells = [c.text.strip() for c in tbl.rows[0].cells]
        def _col_idx(name):
            name_l = name.lower()
            for i, c in enumerate(header_cells):
                if name_l in c.lower():
                    return i
            return None

        idx_rb  = _col_idx('rb') or _col_idx('rb.') or 0
        idx_naz = _col_idx('naziv opreme') or _col_idx('naziv') or 1
        idx_tid = _col_idx('tid') or 2
        idx_sn  = _col_idx('serijski') or _col_idx('serijski broj') or 3
        idx_otp = _col_idx('otp') or 4

        # Očisti sve retke osim headera (ukloni placeholder redove iz predloška)
        while len(tbl.rows) > 1:
            r = tbl.rows[-1]
            tbl._tbl.remove(r._tr)

        # Izgradi listu uređaja iz odabira
        device_pairs = []
        if selected_devices:
            for sn in selected_devices:
                md = _dev_model_map.get(sn, '')
                device_pairs.append((md, sn))
        else:
            # Fallback (ako selected_devices nije predan)
            sn_list = [s.strip() for s in str(klijent.get('sn_uredjaja','')).split(',') if s.strip()]
            model_list = [s.strip() for s in str(klijent.get('model_uredjaja','')).split(',') if s.strip()]
            for i, sn in enumerate(sn_list):
                md = model_list[i] if i < len(model_list) else ''
                device_pairs.append((md, sn))

        # Uređaji: svaki u novi red (Rb., Naziv Opreme, TID, Serijski broj, OTP)
        rb = 1
        for md, sn in device_pairs:
            cells = tbl.add_row().cells
            if idx_rb  is not None and idx_rb  < len(cells): cells[idx_rb].text  = str(rb)
            if idx_naz is not None and idx_naz < len(cells): cells[idx_naz].text = md or 'Uređaj'
            if idx_tid is not None and idx_tid < len(cells): cells[idx_tid].text = ''  # dopuni ako postoji TID
            if idx_sn  is not None and idx_sn  < len(cells): cells[idx_sn].text  = sn or ''
            if idx_otp is not None and idx_otp < len(cells): cells[idx_otp].text = ''  # dopuni ako postoji OTP
            rb += 1

        # SIM kartice: svaka u novi red; Naziv Opreme='SIM kartica', Serijski broj=broj SIM-a; TID/OTP prazno
        if selected_sims:
            for sim_sn in selected_sims:
                cells = tbl.add_row().cells
                if idx_rb  is not None and idx_rb  < len(cells): cells[idx_rb].text  = str(rb)
                if idx_naz is not None and idx_naz < len(cells): cells[idx_naz].text = 'SIM kartica'
                if idx_tid is not None and idx_tid < len(cells): cells[idx_tid].text = ''
                if idx_sn  is not None and idx_sn  < len(cells): cells[idx_sn].text  = sim_sn or ''
                if idx_otp is not None and idx_otp < len(cells): cells[idx_otp].text = ''
                rb += 1
        # Postavi margine ćelija (padding) za bolju čitljivost
        try:
            set_tbl_cell_margins(tbl, top=120, left=160, bottom=120, right=160)
            set_table_uniform_borders(tbl, size=8, color='000000', space=0)
        except Exception:
            pass
    # Nakon popunjavanja tablice počisti prazne numerirane paragrafe ispod
    try:
        drop_empty_numbered_paragraphs(doc)
    except Exception:
        pass
    # Spremi dokument
    client_slug = slugify(korisnik)
    if save_in_zapisnici:
        out_dir = os.path.join(STATIC_DIR, 'zapisnici', 'Deinstalacija')
        os.makedirs(out_dir, exist_ok=True)
        filename = f"DEINST RN {rn_str.replace('/', '-') } {korisnik},{adresa}.docx"
        out_path = os.path.join(out_dir, filename)
        file_rel = f"zapisnici/Deinstalacija/{filename}"
    else:
        out_dir = os.path.join(STATIC_DIR, 'nalozi', client_slug)
        os.makedirs(out_dir, exist_ok=True)
        filename = f"DEINST RN {rn_str.replace('/', '-') } {korisnik},{adresa}.docx"
        out_path = os.path.join(out_dir, filename)
        file_rel = f"nalozi/{client_slug}/{filename}"
    doc.save(out_path)

    # Upis naloga u nalozi.json i otvorene_deinstalacije.JSON




    order = {
        'type': 'deinstalacija',
        'client': korisnik,
        'file': file_rel,
        'created_at': datetime.datetime.now().isoformat(),
        'rn': rn_str,
        'status': 'nezadužen',


    'deinstalirani sn uredjaja': ", ".join([str(x) for x in selected_devices]) if selected_devices else "",

    'deinstalirani sn SIM-a': ", ".join([str(x) for x in selected_sims]) if selected_sims else ""

    }
    try:
        n = read_nalozi()
    except Exception:
        n = []
    if not skip_nalozi:
        n.append(dict(order))
        write_nalozi(n)
    else:
        n2 = [x for x in n if str(x.get('rn')) != str(rn_str)]
        if len(n2) != len(n):
            write_nalozi(n2)

    try:
        od = read_otvorene_deinstalacije()
    except Exception:
        od = []
    od.append(dict(order))
    write_otvorene_deinstalacije(od)

    try:
        for _sn in (selected_devices or []):
            append_device_log(_sn, 'deinstaliran', details={'rn': rn_str, 'client': korisnik})
    except Exception:
        pass
    
    
    return rn_str, out_path



def generate_servis_docx(klijent, selected_devices=None, selected_sims=None, napomena: str = "", save_in_zapisnici=False, skip_nalozi=False):
    """Kreira SERVIS RN DOCX koristeći template rn_template_servis.docx.
    ZAHTJEV: U tablici umjesto {{uređaj}} upisati NAZIV uređaja (model) iz boxa "Uređaji kod klijenta" (stranica /servis/),
    a umjesto {{SN_UREDAJA}} upisati serijski broj. Ako je označeno više uređaja, kreirati novi red za svaki.
    U dokument upisivati ISKLJUČIVO označene uređaje; ako nema označenih uređaja, ne kreirati red u tablici.
    """
    template_path = os.path.join(STATIC_DIR, 'datoteke', 'rn_template_servis.docx')
    if not os.path.exists(template_path):
        abort(500, description='Nedostaje rn_template_servis.docx u static/datoteke')

    doc = Document(template_path)

    korisnik = str(klijent.get('name',''))
    oib = str(klijent.get('oib',''))
    adresa = str(klijent.get('shipping') or klijent.get('headquarters') or '')
    kontakt = str(klijent.get('phone',''))

    # RN kontinuirani za servis
    nalozi = read_nalozi()
    year = datetime.datetime.now().year
    broj = next_rn_continuous_servis(nalozi)
    rn_str = f"{broj:04d}/{year}"

    # Placeholder mapping
    mapping = {
        '{{RN}}': rn_str,
        '{{0001/2025}}': rn_str,
        '{{KORISNIK}}': korisnik,
        '{{ADRESA_ISPORUKE}}': adresa,
        '{{OIB}}': oib,
        '{{KONTAKT}}': kontakt,
        # Osiguraj da eventualni preostali placeholderi u tekstu izvan tablice ne ostanu vidljivi
        '{{uređaj}}': '',
        '{{SN_UREDAJA}}': '',
        '{{TID}}': '',
    }
    _napomena_txt = str(napomena or '').strip()
    if _napomena_txt:
        mapping['{{NAPOMENA}}'] = _napomena_txt
        mapping['[[NAPOMENA]]'] = _napomena_txt
        mapping['{{Napomena}}'] = _napomena_txt
    replace_text_in_doc(doc, mapping)

    # --- Mapiranje SN -> Model iz podataka klijenta i aktivnih uređaja kod klijenta ---
    _dev_model_map = {}
    try:
        sn_list_tmp = [s.strip() for s in str(klijent.get('sn_uredjaja','')).split(',') if s.strip()]
    except Exception:
        sn_list_tmp = []
    try:
        model_list_tmp = [s.strip() for s in str(klijent.get('model_uredjaja','')).split(',') if s.strip()]
    except Exception:
        model_list_tmp = []
    for i, snv in enumerate(sn_list_tmp):
        if i < len(model_list_tmp):
            _dev_model_map[snv] = model_list_tmp[i]

    # Dopuni iz aktivnih evidencija (filtrirano na ovog klijenta kada je moguće)
    try:
        _akt = read_aktivni_uredjaji()
    except Exception:
        _akt = []
    for d in _akt or []:
        try:
            belongs = (str(d.get('client','')) == korisnik) or (oib and str(d.get('oib','')) == oib)
        except Exception:
            belongs = False
        if belongs:
            snv = str(d.get('serijski') or '').strip()
            mdl = str(d.get('model') or '').strip()
            if snv and mdl:
                _dev_model_map.setdefault(snv, mdl)

    # --- Priprema redaka: isključivo označeni uređaji ---
    device_pairs = []  # (model, sn)
    if selected_devices:
        for sn in selected_devices:
            sn = str(sn).strip()
            if not sn:
                continue
            md = _dev_model_map.get(sn, 'Uređaj')
            device_pairs.append((md, sn))

    # --- Lociraj tablicu (5-stupaca: Rb., Naziv Opreme, TID, Serijski broj, OTP) ---
    tbl = None
    for headers in (
        ['Rb.', 'Naziv Opreme', 'TID', 'Serijski broj', 'OTP'],
        ['Rb', 'Naziv Opreme', 'TID', 'Serijski broj', 'OTP'],
        ['Naziv Opreme', 'TID', 'Serijski broj'],
    ):
        tbl = find_table_with_header(doc, headers)
        if tbl:
            break
    if not tbl and doc.tables:
        tbl = doc.tables[0]
    if not tbl:
        # Kreiraj novu tablicu s očekivanim headerima
        tbl = doc.add_table(rows=1, cols=5)
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = 'Rb.', 'Naziv Opreme', 'TID', 'Serijski broj', 'OTP'

    # --- Očisti sve postojeće retke osim headera (uklanja i placeholder {{uređaj}} / {{SN_UREDAJA}}) ---
    while len(tbl.rows) > 1:
        r = tbl.rows[-1]
        tbl._tbl.remove(r._tr)

    # --- Ako NEMA označenih uređaja: ne kreiraj nikakav red ---
    if device_pairs:
        # Odredi indekse kolona po nazivu (case-insensitive, djelomično poklapanje)
        header_cells = [c.text.strip() for c in tbl.rows[0].cells]
        def _col_idx(name):
            name_l = name.lower()
            for i, c in enumerate(header_cells):
                if name_l in c.lower():
                    return i
            return None
        idx_rb  = _col_idx('rb') or _col_idx('rb.') or 0
        idx_naz = _col_idx('naziv opreme') or _col_idx('naziv') or 1
        idx_tid = _col_idx('tid') or 2
        idx_sn  = _col_idx('serijski') or _col_idx('serijski broj') or 3
        idx_otp = _col_idx('otp') or 4

        # Upis označenih uređaja
        rb = 1
        for md, sn in device_pairs:
            cells = tbl.add_row().cells
            if idx_rb  is not None and idx_rb  < len(cells): cells[idx_rb].text  = str(rb)
            if idx_naz is not None and idx_naz < len(cells): cells[idx_naz].text = str(md or 'Uređaj')
            if idx_tid is not None and idx_tid < len(cells): cells[idx_tid].text = ''  # TID nije specificiran u zahtjevu
            if idx_sn  is not None and idx_sn  < len(cells): cells[idx_sn].text  = str(sn or '')
            if idx_otp is not None and idx_otp < len(cells): cells[idx_otp].text = ''
            rb += 1

        # (Opcionalno) SIM kartice — ako su označene, dodaj redove "SIM kartica" s njihovim SN
        if selected_sims:
            for sim_sn in selected_sims:
                cells = tbl.add_row().cells
                if idx_rb  is not None and idx_rb  < len(cells): cells[idx_rb].text  = str(rb)
                if idx_naz is not None and idx_naz < len(cells): cells[idx_naz].text = 'SIM kartica'
                if idx_tid is not None and idx_tid < len(cells): cells[idx_tid].text = ''
                if idx_sn  is not None and idx_sn  < len(cells): cells[idx_sn].text  = str(sim_sn or '')
                if idx_otp is not None and idx_otp < len(cells): cells[idx_otp].text = ''
                rb += 1

        # Uljepšaj tablicu (nije kritično, ali konzistentno s deinstalacijama)
        try:
            set_tbl_cell_margins(tbl, top=120, left=160, bottom=120, right=160)
            set_table_uniform_borders(tbl, size=8, color='000000', space=0)
        except Exception:
            pass

    # Spremi dokument
    client_slug = slugify(korisnik)
    if save_in_zapisnici:
        out_dir = os.path.join(STATIC_DIR, 'zapisnici', 'Servis')
        os.makedirs(out_dir, exist_ok=True)
        filename = f"SERVIS RN {rn_str.replace('/', '-') } {korisnik},{adresa}.docx"
        out_path = os.path.join(out_dir, filename)
        file_rel = f"zapisnici/Servis/{filename}"
    else:
        out_dir = os.path.join(STATIC_DIR, 'nalozi', client_slug)
        os.makedirs(out_dir, exist_ok=True)
        filename = f"SERVIS RN {rn_str.replace('/', '-') } {korisnik},{adresa}.docx"
        out_path = os.path.join(out_dir, filename)
        file_rel = f"nalozi/{client_slug}/{filename}"
    doc.save(out_path)

    # Zapis naloga (postojeća logika)
    order = {
        'type': 'servis',
        'client': korisnik,
        'file': file_rel,
        'created_at': datetime.datetime.now().isoformat(),
        'rn': rn_str,
        'status': 'nezadužen',
        'servisirani sn uredjaja': ", ".join([str(x) for x in selected_devices]) if selected_devices else "",
        'servisirani sn SIM-a': ", ".join([str(x) for x in selected_sims]) if selected_sims else ""
    }
    try:
        base = read_nalozi()
    except Exception:
        base = []
    if not skip_nalozi:
        base.append(dict(order))
        write_nalozi(base)
    else:
        base2 = [x for x in base if str(x.get('rn')) != str(rn_str)]
        if len(base2) != len(base):
            write_nalozi(base2)

    # upiši u otvoreni_servisi.JSON
    try:
        osv = read_otvoreni_servisi()
    except Exception:
        osv = []
    osv.append(dict(order))
    write_otvoreni_servisi(osv)

    # upiši u servisni.nalog.JSON (svi servisni nalozi)
    try:
        snj = read_servisni_nalozi()
    except Exception:
        snj = []
    snj.append(dict(order))
    write_servisni_nalozi(snj)

    try:
        for _sn in (selected_devices or []):
            append_device_log(_sn, 'servis evidentiran', details={'rn': rn_str, 'client': korisnik})
    except Exception:
        pass

    return rn_str, out_path


@app.route('/zaduzi-uredjaj', methods=['GET','POST'])
@app.route('/zaduzi_uredjaj', methods=['GET','POST'])
def zaduzi_uredjaj():
    if not has_role('superadmin','admin','voditelj','podrška','podrska'):
        flash('Nemate ovlasti za ovu radnju.', 'danger')
        return redirect(url_for('index'))

    tehnicari = refresh_tehnicari()

    # --- DODANO: Podrška operateri (status "Podrška") za /zaduzi_nalog ---
    try:
        _ops_all = read_operateri()
    except Exception:
        _ops_all = []
    podrska_ops = []
    for _o in _ops_all:
        try:
            st = str((_o.get('status') or _o.get('role') or '')).strip().lower()
            if st in ('podrška','podrska'):
                podrska_ops.append(_o)
        except Exception:
            pass

















    uredjaji = read_uredjaji()

    # === Prikaži u padajućem izborniku SAMO 'Neaktivne' uređaje ===
    # Neaktivni = nisu kod klijenta (aktivni), nisu evidentirani kod klijenta (klijenti.json) i nisu zaduženi kod tehničara.
    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()
    try:
        aktivni = read_aktivni_uredjaji()
    except Exception:
        aktivni = []
    active_serials = {_norm_sn(d.get('serijski')) for d in aktivni if d.get('active', True)}
    try:
        zaduzeni = read_zaduzene_uredjaje()
    except Exception:
        zaduzeni = []
    assigned_serials = {_norm_sn(z.get('serijski')) for z in zaduzeni}
    try:
        klijenti = read_klijenti()
    except Exception:
        klijenti = []
    client_serials = set()
    for k in klijenti:
        try:
            for sn in str(k.get('sn_uredjaja') or '').split(','):
                snn = _norm_sn(sn)
                if snn:
                    client_serials.add(snn)
        except Exception:
            pass
    # Isključi i privremeno isključene uređaje
    try:
        priv = read_privremeno_iskljuceni()
    except Exception:
        priv = []
    priv_set = {_norm_sn(p.get('serijski')) for p in (priv or [])}
    neaktivni = [u for u in uredjaji
                 if _norm_sn(u.get('serijski')) not in active_serials
                 and _norm_sn(u.get('serijski')) not in client_serials
                 and _norm_sn(u.get('serijski')) not in assigned_serials
                 and _norm_sn(u.get('serijski')) not in priv_set]
    if request.method == 'POST':
        tech_username = request.form.get('tehnicar','').strip()
        serijski = request.form.get('uredjaj','').strip()
        if not tech_username or not serijski:
            flash("Odaberi tehničara i uređaj.", "danger")
            return render_template('zaduzi.uredjaj.html', title="Zaduži uređaj", username=current_user.username, tehnicari=tehnicari, uredjaji=neaktivni, podrska_ops=podrska_ops)
        # find device in available list
        lst = read_uredjaji()
        dev = next((d for d in lst if d.get('serijski') == serijski), None)
        if not dev:
            flash("Uređaj nije pronađen ili je već zadužen.", "warning")
            return render_template('zaduzi.uredjaj.html', title="Zaduži uređaj", username=current_user.username, tehnicari=tehnicari, uredjaji=neaktivni, podrska_ops=podrska_ops)
        # move to zaduzeni
        zlist = read_zaduzene_uredjaje()
        if any(z.get('serijski','').lower()==serijski.lower() for z in zlist):
            flash("Uređaj je već zadužen.", "warning")
            return redirect(url_for('zaduzi_uredjaj'))
        # (PROMJENA) Ne uklanjamo iz uredjaji.json kako se 'namjena' ne bi izgubila
        # lst = [d for d in lst if d.get('serijski') != serijski]
        # write_uredjaji(lst)
        zitem = {
            'model': dev.get('model',''),
            'serijski': dev.get('serijski',''),
            'assigned_to': tech_username,
            'assigned_at': datetime.datetime.now().isoformat()
        }
        zlist.append(zitem)
        write_zaduzene_uredjaje(zlist)
        try:
            append_device_log(serijski, 'zadužen tehničaru', who=tech_username)
        except Exception:
            pass
        flash(f"Uređaj {serijski} zadužen za {tech_username}.", "success")
        return redirect(url_for('uredjaji'))
    return render_template('zaduzi.uredjaj.html', title="Zaduži uređaj", username=current_user.username, tehnicari=tehnicari, uredjaji=neaktivni, podrska_ops=podrska_ops)

@app.route('/api/razduzi/<serijski>', methods=['POST'])
@login_required
def api_razduzi(serijski):
    zlist = read_zaduzene_uredjaje()
    item = next((z for z in zlist if z.get('serijski') == serijski), None)
    if not item:
        return jsonify({'ok': False, 'error': 'Uređaj nije pronađen u zaduženima.'}), 404
    # Dozvoli razduživanje superadminu ili samom tehničaru
    if not (current_user.is_superadmin or current_user.username == item.get('assigned_to')):
        return jsonify({'ok': False, 'error': 'Nedovoljna prava za razduživanje.'}), 403
    # remove from zaduzeni and return to available
    zlist = [z for z in zlist if z.get('serijski') != serijski]
    write_zaduzene_uredjaje(zlist)
    # (PROMJENA) Ne dodajemo natrag u uredjaji.json jer uređaj nikad nije uklonjen; time 'namjena' ostaje netaknuta
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True})



# -------------------- ZADUŽI / RAZDUŽI SIM --------------------
@app.route('/zaduzi-sim', methods=['GET','POST'])
@login_required
def zaduzi_sim():
    tehnicari = refresh_tehnicari() if 'refresh_tehnicari' in globals() else read_tehnicari()
    # --- DODANO: Podrška operateri (status 'Podrška') za /zaduzi-sim ---
    try:
        _ops_all = read_operateri()
    except Exception:
        _ops_all = []
    podrska_ops = []
    for _o in _ops_all:
        try:
            st = str((_o.get('status') or _o.get('role') or '')).strip().lower()
            if st in ('podrška','podrska'):
                podrska_ops.append(_o)
        except Exception:
            pass
    # --- DODANO: Podrška operateri (status 'Podrška') za /zaduzi_sim ---
    try:
        _ops_all = read_operateri()
    except Exception:
        _ops_all = []
    podrska_ops = []
    for _o in _ops_all:
        try:
            st = str((_o.get('status') or _o.get('role') or '')).strip().lower()
            if st in ('podrška','podrska'):
                podrska_ops.append(_o)
        except Exception:
            pass
    # === Prikaži u padajućem izborniku SAMO 'Neaktivne' SIM-ove ===
    sims_all = read_sim()
    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()
    try:
        aktivni = read_aktivni_sim()
    except Exception:
        aktivni = []
    active_serials = {_norm_sn(s.get('serijski')) for s in aktivni if s.get('active', True)}
    try:
        zaduzeni = read_zaduzene_sim()
    except Exception:
        zaduzeni = []
    assigned_serials = {_norm_sn(z.get('serijski')) for z in zaduzeni}
    try:
        klijenti = read_klijenti()
    except Exception:
        klijenti = []
    client_serials = set()
    for k in klijenti:
        try:
            for sn in str(k.get('sn_SIM') or '').split(','):
                snn = _norm_sn(sn)
                if snn:
                    client_serials.add(snn)
        except Exception:
            pass
    neaktivni_sims = [s for s in sims_all
                      if _norm_sn(s.get('serijski')) not in active_serials
                      and _norm_sn(s.get('serijski')) not in client_serials
                      and _norm_sn(s.get('serijski')) not in assigned_serials]
    if request.method == 'POST':
        tech_username = request.form.get('tehnicar','').strip()
        serijski = request.form.get('sim','').strip()
        if not tech_username or not serijski:
            flash("Odaberi tehničara i SIM.", "danger")
            return render_template('zaduzi.sim.html', title="Zaduži SIM", username=current_user.username, tehnicari=tehnicari, sims=neaktivni_sims, podrska_ops=podrska_ops)
        # potvrdi da SIM postoji u sim.json (ali GA NE BRIŠEMO više)
        lst = read_sim()
        sim_item = next((s for s in lst if s.get('serijski') == serijski), None)
        if not sim_item:
            flash("SIM nije pronađen u skladištu.", "warning")
            return render_template('zaduzi.sim.html', title="Zaduži SIM", username=current_user.username, tehnicari=tehnicari, sims=neaktivni_sims, podrska_ops=podrska_ops)
        # upiši/azuriraj u zaduzeni.sim.json
        zlist = read_zaduzene_sim()
        existing = next((z for z in zlist if z.get('serijski') == serijski), None)
        if existing:
            existing['assigned_to'] = tech_username
            existing['assigned_at'] = datetime.datetime.now().isoformat()
            existing['provider'] = sim_item.get('provider','')
        else:
            zlist.append({
                'provider': sim_item.get('provider',''),
                'serijski': sim_item.get('serijski',''),
                'assigned_to': tech_username,
                'assigned_at': datetime.datetime.now().isoformat()
            })
        write_zaduzene_sim(zlist)
        flash(f"SIM {serijski} zadužen za {tech_username}.", "success")
        return redirect(url_for('sim'))
    return render_template('zaduzi.sim.html', title="Zaduži SIM", username=current_user.username, tehnicari=tehnicari, sims=neaktivni_sims, podrska_ops=podrska_ops)
@app.route('/api/razduzi-sim/<serijski>', methods=['POST'])
@login_required
def api_razduzi_sim(serijski):
    zlist = read_zaduzene_sim()
    item = next((z for z in zlist if z.get('serijski') == serijski), None)
    if not item:
        return jsonify({'ok': False, 'error': 'SIM nije pronađen u zaduženima.'}), 404
    if not (current_user.is_superadmin or current_user.username == item.get('assigned_to')):
        return jsonify({'ok': False, 'error': 'Nedovoljna prava za razduživanje.'}), 403
    # samo ukloni iz zaduženih; ne diramo sim.json (SIM ostaje vidljiv)
    zlist = [z for z in zlist if z.get('serijski') != serijski]
    write_zaduzene_sim(zlist)
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True})

@app.route('/zaduzi-nalog', methods=['GET','POST'])

@app.route('/zaduzi_nalog', methods=['GET','POST'])
@login_required
def zaduzi_nalog():

    tehnicari = refresh_tehnicari() if 'refresh_tehnicari' in globals() else read_tehnicari()
    # --- DODANO: Podrška operateri (status 'Podrška') za /zaduzi_nalog ---
    try:
        _ops_all = read_operateri()
    except Exception:
        _ops_all = []
    podrska_ops = []
    for _o in _ops_all:
        try:
            st = str((_o.get('status') or _o.get('role') or '')).strip().lower()
            if st in ('podrška','podrska'):
                podrska_ops.append(_o)
        except Exception:
            pass

    # Učitaj otvorene instalacije i deinstalacije
    try:
        inst = read_otvorene_instalacije()
    except Exception:
        inst = []
    try:
        deinst = read_otvorene_deinstalacije()
    except Exception:
        deinst = []
    # Fallback: include any legacy/singular filename variants if present
    try:
        _fallbacks = [
            os.path.join(STATIC_DIR, 'otvorena_deinstalacije.JSON'),
            os.path.join(STATIC_DIR, 'otvorene_deinstalacije.json'),
            os.path.join(STATIC_DIR, 'otvorena_deinstalacije.json'),
        ]
        existing = set()
        try:
            existing = {(str(x.get('rn') or x.get('RN') or x.get('id') or x.get('broj') or x.get('number') or '')).strip() for x in deinst}
        except Exception:
            existing = set()
        for _p in _fallbacks:
            try:
                if os.path.exists(_p):
                    extra = _read_json(_p) or []
                    for x in extra:
                        key = (str(x.get('rn') or x.get('RN') or x.get('id') or x.get('broj') or x.get('number') or '')).strip()
                        if key and key not in existing:
                            deinst.append(x)
                            existing.add(key)
            except Exception:
                pass
    except Exception:
        pass



    # Učitaj otvorene servise
    try:
        serv = read_otvoreni_servisi()
    except Exception:
        serv = []
    # Fallback: dopuštene varijante naziva datoteke
    try:
        _serv_fallbacks = [
            os.path.join(STATIC_DIR, 'otvoreni_servisi.json'),
            os.path.join(STATIC_DIR, 'otvoreni_servis.JSON'),
            os.path.join(STATIC_DIR, 'otvoreni_servis.json'),
        ]
        existing_serv = set()
        try:
            existing_serv = {(str(x.get('rn') or x.get('RN') or x.get('id') or x.get('broj') or x.get('number') or '')).strip() for x in serv}
        except Exception:
            existing_serv = set()
        for _p in _serv_fallbacks:
            try:
                if os.path.exists(_p):
                    extra = _read_json(_p) or []
                    for x in extra:
                        key = (str(x.get('rn') or x.get('RN') or x.get('id') or x.get('broj') or x.get('number') or '')).strip()
                        if key and key not in existing_serv:
                            serv.append(x)
                            existing_serv.add(key)
            except Exception:
                pass
    except Exception:
        pass
# Pripremi nalozi: INSTALACIJE prvo, zatim DEINSTALACIJE, s prefiksom za prikaz
    nalozi = []
    for _src_type, _lst in (('instalacija', inst), ('deinstalacija', deinst), ('servis', serv)):
        for _it in _lst:
            _d = dict(_it)
            _d['type'] = _src_type
            rn_plain = str(_d.get('rn') or _d.get('RN') or _d.get('broj') or _d.get('id') or _d.get('number') or '')
            prefix = 'INSTALL' if _src_type == 'instalacija' else ('DEINSTALL' if _src_type == 'deinstalacija' else 'SERVIS')
            _d['rn_plain'] = rn_plain
            _d['rn'] = f"{prefix} {rn_plain}" if rn_plain else prefix
            nalozi.append(_d)


            # Ukloni iz padajućeg izbornika RN-ove koji su već zaduženi ili zaključeni
    # Kriterij: (tip, čisti RN). Ako za postojeći zapis ne možemo pouzdano odrediti tip, NE filtriramo po tom zapisu.
    def _infer_type(item):
        # 1) eksplicitni 'type'
        t = str((item.get('type') or '')).lower().strip()
        if t in ('instalacija','deinstalacija','servis'):
            return t
        # 2) RN prefiks
        r = str(item.get('rn') or '')
        if r.startswith('INSTALL'):
            return 'instalacija'
        if r.startswith('DEINSTALL'):
            return 'deinstalacija'
        if r.startswith('SERVICE') or r.startswith('SERVIS'):
            return 'servis'
        # 3) heuristike po ključevima
        keys_low = ' '.join(item.keys()).lower() if isinstance(item, dict) else ''
        if 'deinstal' in keys_low:
            return 'deinstalacija'
        if 'servis' in keys_low or 'service' in keys_low:
            return 'servis'
        # 4) nepoznato -> ne filtriramo na temelju ovog zapisa
        return 'unknown'

    def _pure_rn(item):
        r = str(item.get('rn') or '')
        rp = str(item.get('rn_plain') or '')
        m_rn = re.search(r'(\d{4}/\d{4})', r)
        pure = m_rn.group(1) if m_rn else (rp or r)
        return pure

    try:
        _assigned = read_zaduzene_nalozi()
    except Exception:
        _assigned = []
    try:
        _closed = read_zakljuceni_nalozi()
    except Exception:
        _closed = []

    assigned_pairs = {( _infer_type(z), str(_pure_rn(z)) ) for z in _assigned if _infer_type(z) != 'unknown'}
    closed_pairs   = {( _infer_type(z), str(_pure_rn(z)) ) for z in _closed   if _infer_type(z) != 'unknown'}

    nalozi = [x for x in nalozi if (_infer_type(x), _pure_rn(x)) not in assigned_pairs and (_infer_type(x), _pure_rn(x)) not in closed_pairs]
    # Deduplicate within the same type by pure RN (keep first occurrence)
    _seen_pairs = set()
    _nalozi_dedup = []
    for _x in nalozi:
        _pair = (_infer_type(_x), _pure_rn(_x))
        if _pair not in _seen_pairs:
            _nalozi_dedup.append(_x)
            _seen_pairs.add(_pair)
    nalozi = _nalozi_dedup

    # Sigurnosni fallback: ako je nakon filtriranja lista prazna, a postoje otvorene instalacije,
    # ponovo napuni instalacije kako padajući izbornik ne bi ostao prazan (ne dira deinst/servise)
    try:
        if (not nalozi) and inst:
            for _it in inst:
                _d = dict(_it)
                _d['type'] = 'instalacija'
                rn_plain = str(_d.get('rn') or _d.get('RN') or _d.get('broj') or _d.get('id') or _d.get('number') or '')
                _d['rn_plain'] = rn_plain
                _d['rn'] = f"INSTALL {rn_plain}" if rn_plain else 'INSTALL'
                nalozi.append(_d)
    except Exception:
        pass



    if request.method == 'POST':
        tech_username = request.form.get('tehnicar','').strip()
        rn_input = request.form.get('nalog','').strip()

        # Determine selected type from rn_input prefix, then extract pure RN (e.g. 0001/2025)
        _val = rn_input.strip()
        _sel_type = 'unknown'
        if _val.upper().startswith('INSTALL'):
            _sel_type = 'instalacija'
        elif _val.upper().startswith('DEINSTALL'):
            _sel_type = 'deinstalacija'
        elif _val.upper().startswith('SERVICE') or _val.upper().startswith('SERVIS'):
            _sel_type = 'servis'

        m_rn = re.search(r'(\d{4}/\d{4})', rn_input or '')
        rn = m_rn.group(1) if m_rn else rn_input

        if not tech_username or not rn:
            flash("Odaberi tehničara i nalog.", "danger")
            return render_template('zaduzi.nalog.html', title="Zaduži nalog", username=current_user.username, tehnicari=tehnicari, nalozi=nalozi, inst_list=inst, dein_list=deinst, serv_list=serv, instalacije=inst, deinstalacije=deinst, servisi=serv, icons={'instalacija':'⚙️','deinstalacija':'♻️','servis':'🛠️'}, podrska_ops=podrska_ops)

        # Find the order by BOTH (type, rn) to avoid collisions when RN numbers are the same across types
        try:
            lst = read_nalozi()
        except Exception:
            lst = []

        def _tmatch(v):
            t = str((v.get('type') or '')).lower().strip()
            return t in ('instalacija','deinstalacija','servis') and ((_sel_type == 'unknown') or (t == _sel_type))

        nalog = next((n for n in lst if str(n.get('rn')) == rn and _tmatch(n)), None)
        if not nalog:
            # As a fallback, try to find in open lists (inst/deinst/serv) by type + rn
            _fallback_pool = []
            for _src_type, _lst in (('instalacija', inst), ('deinstalacija', deinst), ('servis', serv)):
                for _it in _lst:
                    _d = dict(_it)
                    _d['type'] = _src_type
                    _fallback_pool.append(_d)
            nalog = next((n for n in _fallback_pool if str(n.get('rn') or n.get('RN') or '') == rn and str((n.get('type') or '')).lower().strip() == _sel_type), None)

        if not nalog:
            flash("Nalog nije pronađen ili je već zadužen.", "warning")
            return render_template('zaduzi.nalog.html', title="Zaduži nalog", username=current_user.username, tehnicari=tehnicari, nalozi=nalozi, inst_list=inst, dein_list=deinst, podrska_ops=podrska_ops)

        # Check if already assigned for the same (type, rn)
        zlist = read_zaduzene_nalozi()
        def _infer_type_str(item):
            t = str((item.get('type') or '')).lower().strip()
            if t in ('instalacija','deinstalacija','servis'):
                return t
            r = str(item.get('rn') or '')
            if r.startswith('INSTALL'):
                return 'instalacija'
            if r.startswith('DEINSTALL'):
                return 'deinstalacija'
            if r.startswith('SERVICE') or r.startswith('SERVIS'):
                return 'servis'
            return 'unknown'

        already = any((str(z.get('rn')) == rn) and (_infer_type_str(z) == _sel_type) for z in zlist)
        if already:
            flash("Nalog je već zadužen.", "warning")
            return redirect(url_for('zaduzi_nalog'))

        # Assign
        zitem = dict(nalog)
        zitem['rn'] = rn
        zitem['type'] = _sel_type if _sel_type != 'unknown' else (zitem.get('type') or '')
        zitem['assigned_to'] = tech_username
        zitem['assigned_at'] = datetime.datetime.now().isoformat()
        zlist.append(zitem)
        write_zaduzene_nalozi(zlist)

        # === AUTO E-MAIL tehničaru na zaduženje naloga (po zahtjevu) ===
        try:
            # Pronađi e-mail operatera (tehničara) po username-u
            to_email = _find_operator_email(tech_username)
            if to_email:
                # Odredi prefiks i podatke za naslov
                _type = (zitem.get('type') or '').strip().lower()
                _pref = 'INSTALL'
                if _type == 'deinstalacija':
                    _pref = 'DEINSTALL'
                elif _type == 'servis':
                    _pref = 'SERVIS'
                _rn = str(rn)
                _client = str(zitem.get('client') or zitem.get('klijent') or zitem.get('customer') or '')
                # Pokušaj dohvatiti adresu iz klijenti.json
                _addr = ''
                try:
                    for _k in (read_klijenti() or []):
                        nm = str(_k.get('name') or _k.get('naziv') or '')
                        if _client and nm == _client:
                            _addr = str(_k.get('shipping') or _k.get('headquarters') or '')
                            break
                except Exception:
                    pass
                # Kreiraj subject i body
                _subject = f"{_pref} {_rn} {_client}, {_addr}".strip().rstrip(',')
                _body = f'Na Vas je zadužen radni nalog "{_subject}". Možete preuzeti nalog.'
                # Pripremi attachment: word dokument naloga, ako postoji
                _attach_path = None
                try:
                    _file_rel = str(zitem.get('file') or '')
                    if _file_rel:
                        # Ako je relativno unutar static/, složi apsolutnu putanju
                        _candidate = os.path.join(STATIC_DIR, _file_rel)
                        if os.path.isfile(_candidate):
                            _attach_path = _candidate
                        elif os.path.isfile(_file_rel):
                            _attach_path = _file_rel
                except Exception:
                    _attach_path = None
                from flask_mail import Message as _Msg
                _msg = _Msg(subject=_subject, recipients=[to_email])
                _msg.body = _body
                if _attach_path:
                    import mimetypes as _mt
                    _ctype, _enc = _mt.guess_type(_attach_path)
                    _ctype = _ctype or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    with open(_attach_path, 'rb') as _fp:
                        _msg.attach(filename=os.path.basename(_attach_path), content_type=_ctype, data=_fp.read())
                mail.send(_msg)
            else:
                try:
                    flash('Upozorenje: e-mail adresa tehničara nije pronađena.', 'warning')
                except Exception:
                    pass
        except Exception as _err_mail_assign:
            try:
                flash(f'Upozorenje: slanje e-maila tehničaru nije uspjelo ({_err_mail_assign}).', 'warning')
            except Exception:
                pass

        pref = 'DEINSTALL' if zitem.get('type') == 'deinstalacija' else ('SERVIS' if zitem.get('type') == 'servis' else 'INSTALL')
        flash(f"Nalog RN {pref} {rn} zadužen za {tech_username}.", "success")
        return redirect(url_for('zaduzeni_nalozi'))



    return render_template('zaduzi.nalog.html', title="Zaduži nalog", username=current_user.username, tehnicari=tehnicari, nalozi=nalozi, inst_list=inst, dein_list=deinst, podrska_ops=podrska_ops)
# -------------------- ZAKLJUČI NALOG --------------------
@app.route('/zakljuci-nalog/<path:rn>', methods=['GET','POST'])

@app.route('/zaduzeni.nalog/<path:rn>', methods=['GET','POST'])
@login_required
def zakljuci_nalog(rn):
    # find nalog in assigned list
    zn = read_zaduzene_nalozi()
    nalog = next((n for n in zn if str(n.get('rn')) == str(rn)), None)
    if not nalog:
        abort(404)
    tech = (nalog.get('assigned_to') or nalog.get('assignedTo') or nalog.get('tehnicar') or '').strip()
    # load client
    klijenti = read_klijenti()
    cli = None
    # try match by id or name
    cid = nalog.get('client_id') or nalog.get('clientId') or nalog.get('oib')
    cname = nalog.get('client') or nalog.get('klijent') or nalog.get('customer')
    if cid:
        cli = next((k for k in klijenti if str(k.get('id') or k.get('oib')) == str(cid)), None)
    if not cli and cname:
        cli = next((k for k in klijenti if str(k.get('name') or k.get('naziv') or k.get('client')) == str(cname)), None)

    ## DEINSTALACIJA-SPECIAL-VIEW (posebna stranica i tok za deinstalacije)
    nalog_type = (nalog.get('type') or '').lower().strip()
    if nalog_type == 'deinstalacija':
        rn_str = str(nalog.get('rn') or rn)
        # Priprema podataka iz klijenti.json (SN-ovi) + mapiranje modela/providera
        client_name = (cli.get('name') if cli else (nalog.get('client') or ''))
        # SN liste iz klijenti.json (zarezom odvojene)
        try:
            sn_devices = [s.strip() for s in (cli.get('sn_uredjaja') or '').split(',') if s.strip()] if cli else []
        except Exception:
            sn_devices = []
        try:
            sn_sims = [s.strip() for s in (cli.get('sn_SIM') or '').split(',') if s.strip()] if cli else []
        except Exception:
            sn_sims = []
        # Mapiranja
        # Fallback: ako nema podataka u klijenti.json, uzmi iz samog naloga (npr. "deinstalirani sn uredjaja", "deinstalirani sn SIM-a")
        if not sn_devices:
            try:
                sn_devices = [s.strip() for s in str(nalog.get('deinstalirani sn uredjaja','')).split(',') if s.strip()]
            except Exception:
                sn_devices = []
        if not sn_sims:
            try:
                sn_sims = [s.strip() for s in str(nalog.get('deinstalirani sn SIM-a','')).split(',') if s.strip()]
            except Exception:
                sn_sims = []

        # Mapiranja iz aktivnih (kod klijenta)
        try:
            akt_dev = read_aktivni_uredjaji()
        except Exception:
            akt_dev = []
        model_by_sn = {str(d.get('serijski')): d.get('model','') for d in akt_dev}
        try:
            akt_sim = read_aktivni_sim()
        except Exception:
            akt_sim = []
        provider_by_sn = {str(s.get('serijski')): s.get('provider','') for s in akt_sim}

        if request.method == 'GET':
            dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in sn_devices]
            sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sn_sims]
            return render_template('zakljuci.deinstalaciju.html',
                                   title='Zaključi deinstalaciju',
                                   username=current_user.username,
                                   nalog=nalog, klijent=cli,
                                   uredjaji_klijent=dev_rows,
                                   sims_klijent=sim_rows)

        # POST – očekujemo checkboxe 'return_device' i 'return_sim'
        if request.method == 'POST' and request.form.get('deinst_action') == 'zakljuci':
            selected_dev = {s.strip() for s in request.form.getlist('return_device') if s.strip()}
            selected_sim = {s.strip() for s in request.form.getlist('return_sim') if s.strip()}


            # === ZAPISNIK (obavezno) — spremi sliku u static/zapisnici/Deinstalacija ===
            file = request.files.get('zapisnik_img')
            if not file or not file.filename:
                flash('Obavezno je priložiti sliku zapisnika (JPG/PNG).', 'danger')
                dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in sn_devices]
                sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sn_sims]
                return render_template('zakljuci.deinstalaciju.html',
                                       title='Zaključi deinstalaciju',
                                       username=current_user.username,
                                       nalog=nalog, klijent=cli,
                                       uredjaji_klijent=dev_rows,
                                       sims_klijent=sim_rows)
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in ('.jpg', '.jpeg', '.png'):
                flash('Dozvoljeni formati slike: JPG/PNG.', 'danger')
                dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in sn_devices]
                sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sn_sims]
                return render_template('zakljuci.deinstalaciju.html',
                                       title='Zaključi deinstalaciju',
                                       username=current_user.username,
                                       nalog=nalog, klijent=cli,
                                       uredjaji_klijent=dev_rows,
                                       sims_klijent=sim_rows)
            ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_rn = re.sub(r'[^A-Za-z0-9_-]+', '_', str(rn_str))
            fname = f"{safe_rn}_{ts}{ext}"
            _dest_dir = os.path.join(STATIC_DIR, 'zapisnici', 'Deinstalacija')
            os.makedirs(_dest_dir, exist_ok=True)
            dest = os.path.join(_dest_dir, fname)
            file.save(dest)
            _saved_rel = os.path.join('static','zapisnici','Deinstalacija', fname)

            # === AUTO E-MAIL (DEINSTALACIJA): po zahtjevu korisnika ===
            try:
                _operater_username = getattr(current_user, 'username', '') or (nalog.get('assigned_to') or '')
                _rn_mail = str(rn_str)
                # Sastavi tijelo poruke
                _lines = []
                _lines.append(f'Operater {_operater_username} je zaključio deinstalacijski nalog sa sljedećim stavkama:')
                _lines.append('')
                # Uređaji preuzeti s korisnika (ako su označeni)
                if selected_dev:
                    _lines.append('Preuzeti uređaj:')
                    for _sn in sorted(selected_dev):
                        _mdl = str(model_by_sn.get(_sn, '') or '')
                        _lines.append(f'{_mdl} - SN - {_sn}')
                    _lines.append('')
                # SIM kartice preuzete s korisnika (ako su označene)
                if selected_sim:
                    _lines.append('Preuzeti SIM:')
                    for _sn in sorted(selected_sim):
                        _prov = str(provider_by_sn.get(_sn, '') or '')
                        _lines.append(f'{_prov} - SN - {_sn}')
                    _lines.append('')
                _lines.append('Ovo je automatska poruka, molim Vas da ne odgovarate na nju. Vaš Billy Pos d.o.o.')
                _body = "\n".join(_lines)

                from flask_mail import Message as FMMessage
                _subject = f"Zaključen deinstalacijski nalog {_rn_mail}"
                msg = FMMessage(subject=_subject, recipients=['webtest806@gmail.com'])
                msg.body = _body

                import mimetypes as _mt
                _mtype, _enc = _mt.guess_type(dest)
                _mtype = _mtype or 'application/octet-stream'
                with open(dest, 'rb') as _fp:
                    msg.attach(filename=os.path.basename(dest), content_type=_mtype, data=_fp.read())
                mail.send(msg)
            except Exception as _mail_err:
                try:
                    flash(f'Upozorenje: e-mail obavijest (deinstalacija) nije poslana ({_mail_err}).', 'warning')
                except Exception:
                    pass
            # Ažuriraj klijenti.json: ukloni odabrane SN iz sn_uredjaja i sn_SIM; rekonstruiraj model_uredjaja
            try:
                kl = read_klijenti()
            except Exception:
                kl = []
            idx = None
            for i,c in enumerate(kl):
                if str(c.get('name')) == str(client_name):
                    idx = i; break
            if idx is not None:
                c = kl[idx]
                # uređaji
                try:
                    exist_dev = [s.strip() for s in (c.get('sn_uredjaja') or '').split(',') if s.strip()]
                except Exception:
                    exist_dev = []
                remain_dev = [s for s in exist_dev if s not in selected_dev]
                c['sn_uredjaja'] = ", ".join(remain_dev)
                # modeli rekonstruiraj na temelju preostalih SN iz aktivnih
                models_remain = []
                for s in remain_dev:
                    m = model_by_sn.get(s, '')
                    if m: models_remain.append(m)
                c['model_uredjaja'] = ", ".join(sorted(set(models_remain)))
                # SIM
                try:
                    exist_sim = [s.strip() for s in (c.get('sn_SIM') or '').split(',') if s.strip()]
                except Exception:
                    exist_sim = []
                remain_sim = [s for s in exist_sim if s not in selected_sim]
                c['sn_SIM'] = ", ".join(remain_sim)
                kl[idx] = c
                write_klijenti(kl)

            # DEAKTIVIRAJ IZ AKTIVNIH EVIDENCIJA (uređaji i SIM-ovi kod klijenta)
            dev_set = set(selected_dev or [])
            sim_set = set(selected_sim or [])

            try:
                akt = read_aktivni_uredjaji()
            except Exception:
                akt = []
            try:
                changed = False
                now_iso = datetime.datetime.now().isoformat()
                for d in akt:
                    try:
                        if str(d.get('serijski')) in dev_set:
                            d['active'] = False
                            d['unassigned_at'] = now_iso
                            changed = True
                    except Exception:
                        pass
                if changed:
                    write_aktivni_uredjaji(akt)
            except Exception:
                pass

            try:
                akt_s = read_aktivni_sim()
            except Exception:
                akt_s = []
            try:
                akt_s = read_aktivni_sim()
            except Exception:
                akt_s = []
            try:
                changed_s = False
                now_iso = datetime.datetime.now().isoformat()
                for s in akt_s:
                    try:
                        if str(s.get('serijski')) in sim_set:
                            s['active'] = False
                            s['unassigned_at'] = datetime.datetime.now().isoformat()
                            changed_s = True
                    except Exception:
                        pass
                if changed_s:
                    write_aktivni_sim(akt_s)
            except Exception:
                pass

            # Vraćanje odabranih uređaja nakon deinstalacije:


            # - ako je operater Tehničar: NE vraćamo u uredjaji.JSON, već zadužujemo tehničaru (zaduzeni.uredjaji.json)


            # - ako je operater Podrška (ili nepoznat): vraćamo u uredjaji.JSON (kao dosad)


            _operater_username = (tech or getattr(current_user, 'username', '') or '').strip()


            def _norm_status(v):


                import unicodedata as _u


                try:


                    return _u.normalize('NFKD', str(v or '')).encode('ascii','ignore').decode('ascii').strip().lower()


                except Exception:


                    return str(v or '').strip().lower()


            _status = ''


            try:


                for _op in (read_operateri() or []):


                    un = _norm_status(_op.get('username') or _op.get('user') or _op.get('name'))


                    if un and un == _norm_status(_operater_username):


                        _status = _norm_status(_op.get('status') or _op.get('role'))


                        break


            except Exception:


                _status = ''


            _is_tehnicar = _status.startswith('tehnicar') or _status.startswith('tehni')


            _is_podrska  = _status.startswith('podrska') or _status.startswith('podrška')


            if selected_dev:


                if _is_tehnicar:


                    zlist_dev = read_zaduzene_uredjaje()


                    existing_sn = {str(z.get('serijski')) for z in zlist_dev}


                    _now_iso = datetime.datetime.now().isoformat()


                    for s in selected_dev:


                        if str(s) not in existing_sn:


                            zlist_dev.append({


                                'model': model_by_sn.get(s, ''),


                                'serijski': str(s),


                                'assigned_to': _operater_username,


                                'assigned_at': _now_iso


                            })


                            existing_sn.add(str(s))


                    write_zaduzene_uredjaje(zlist_dev)


                else:


                    pool = read_uredjaji()


                    seen = {str(d.get('serijski')) for d in pool}


                    for s in selected_dev:


                        if s not in seen:


                            pool.append({'model': model_by_sn.get(s,''), 'serijski': s, 'created_at': datetime.datetime.now().isoformat()})


                            seen.add(s)


                    write_uredjaji(pool)

            
            # Vraćanje odabranih SIM-ova ovisno o statusu operatera
            # - Tehničar: zadužujemo tehničaru (zaduzeni.sim.json) — postojeće ponašanje
            # - Podrška: vraćamo u skladište (sim.json) kao neaktivan
            if selected_sim:
                if '_is_tehnicar' in locals() and _is_tehnicar:
                    zsim = read_zaduzene_sim()
                    existing_sim = {str(z.get('serijski')) for z in zsim}
                    now_iso = datetime.datetime.now().isoformat()
                    for s in selected_sim:
                        if s not in existing_sim:
                            zsim.append({
                                'provider': provider_by_sn.get(s, ''),
                                'serijski': s,
                                'assigned_to': tech,
                                'assigned_at': now_iso
                            })
                            existing_sim.add(s)
                    write_zaduzene_sim(zsim)
                elif '_is_podrska' in locals() and _is_podrska:
                    # Vrati u sim.json (skladište) ako ne postoji
                    pool_sim = read_sim()
                    seen_sim = {str(d.get('serijski')) for d in pool_sim}
                    now = datetime.datetime.now().isoformat()
                    for sn in selected_sim:
                        if sn not in seen_sim:
                            pool_sim.append({'provider': provider_by_sn.get(sn,''), 'serijski': sn, 'created_at': now})
                            seen_sim.add(sn)
                    write_sim(pool_sim)
                else:
                    # Default: ponašaj se kao za tehničara (zaduži ga), da se ne mijenja postojeća logika za ostale uloge
                    zsim = read_zaduzene_sim()
                    existing_sim = {str(z.get('serijski')) for z in zsim}
                    now_iso = datetime.datetime.now().isoformat()
                    for s in selected_sim:
                        if s not in existing_sim:
                            zsim.append({
                                'provider': provider_by_sn.get(s, ''),
                                'serijski': s,
                                'assigned_to': tech,
                                'assigned_at': now_iso
                            })
                            existing_sim.add(s)
                    write_zaduzene_sim(zsim)
            
            
# Premjesti nalog iz zaduženih -> zakljucene.deinstalacije.JSON; u nalozi.json označi kao zaključen
            try:
                zlist = read_zaduzene_nalozi()
            except Exception:
                zlist = []
            rn_str = str(nalog.get('rn'))
            zlist = [z for z in zlist if str(z.get('rn')) != rn_str]
            write_zaduzene_nalozi(zlist)

            # === AUTO E-MAIL tehničaru na zaduženje naloga (po zahtjevu) ===
            try:
                # Pronađi e-mail operatera (tehničara) po username-u
                to_email = _find_operator_email(tech_username)
                if to_email:
                    # Odredi prefiks i podatke za naslov
                    _type = (zitem.get('type') or '').strip().lower()
                    _pref = 'INSTALL'
                    if _type == 'deinstalacija':
                        _pref = 'DEINSTALL'
                    elif _type == 'servis':
                        _pref = 'SERVIS'
                    _rn = str(rn)
                    _client = str(zitem.get('client') or zitem.get('klijent') or zitem.get('customer') or '')
                    # Pokušaj dohvatiti adresu iz klijenti.json
                    _addr = ''
                    try:
                        for _k in (read_klijenti() or []):
                            nm = str(_k.get('name') or _k.get('naziv') or '')
                            if _client and nm == _client:
                                _addr = str(_k.get('shipping') or _k.get('headquarters') or '')
                                break
                    except Exception:
                        pass
                    # Kreiraj subject i body
                    _subject = f"{_pref} {_rn} {_client}, {_addr}".strip().rstrip(',')
                    _body = f'Na Vas je zadužen radni nalog "{_subject}". Možete preuzeti nalog.'
                    # Pripremi attachment: word dokument naloga, ako postoji
                    _attach_path = None
                    try:
                        _file_rel = str(zitem.get('file') or '')
                        if _file_rel:
                            # Ako je relativno unutar static/, složi apsolutnu putanju
                            _candidate = os.path.join(STATIC_DIR, _file_rel)
                            if os.path.isfile(_candidate):
                                _attach_path = _candidate
                            elif os.path.isfile(_file_rel):
                                _attach_path = _file_rel
                    except Exception:
                        _attach_path = None
                    from flask_mail import Message as _Msg
                    _msg = _Msg(subject=_subject, recipients=[to_email])
                    _msg.body = _body
                    if _attach_path:
                        import mimetypes as _mt
                        _ctype, _enc = _mt.guess_type(_attach_path)
                        _ctype = _ctype or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        with open(_attach_path, 'rb') as _fp:
                            _msg.attach(filename=os.path.basename(_attach_path), content_type=_ctype, data=_fp.read())
                    from smtplib import SMTPServerDisconnected as _SMTPDisc
                    try:
                        with mail.connect() as _conn:
                            _conn.send(_msg)
                    except _SMTPDisc:
                        # Jedan ponovni pokušaj na neočekivano zatvaranje konekcije
                        with mail.connect() as _conn:
                            _conn.send(_msg)
                else:
                    try:
                        flash('Upozorenje: e-mail adresa tehničara nije pronađena.', 'warning')
                    except Exception:
                        pass
            except Exception as _err_mail_assign:
                try:
                    flash(f'Upozorenje: slanje e-maila tehničaru nije uspjelo ({_err_mail_assign}).', 'warning')
                except Exception:
                    pass

            try:
                zk = read_zakljuceni_nalozi()
            except Exception:
                zk = []
            closed = dict(nalog)
            # Ne spremati Zahtjev u zatvoreni zapis
            try:
                if 'zahtjev_image' in closed:
                    closed['zahtjev_image'] = ''
            except Exception:
                pass
            closed['status'] = 'zaključen'
            closed['closed_at'] = datetime.datetime.now().isoformat()
            closed['zapisnik_image'] = _saved_rel
            zk.append(closed)
            write_zakljuceni_nalozi(zk)

            # upiši status i u globalne nalozi.json (radi povijesti i /nalozi prikaza)
            try:
                base = read_nalozi()
                found = False
                for it in base:
                    if str(it.get('rn')) == rn_str:
                        it['status'] = 'zaključen'
                        it['closed_at'] = closed['closed_at']
                        it['zapisnik_image'] = _saved_rel
                        found = True
                        break
                if not found:
                    base.append(closed)
                write_nalozi(base)
            except Exception:
                pass

            flash(f"Deinstalacija RN {rn_str} zaključena.", "success")
            tech = (nalog.get('assigned_to') or nalog.get('assignedTo') or nalog.get('tehnicar') or '').strip()
            if tech:
                return redirect(url_for('operater_profil', username=tech))
            return redirect(url_for('zaduzeni_nalozi'))

    ## SERVIS-SPECIAL-VIEW (posebna stranica i tok za servis)
    if nalog_type == 'servis':
        rn_str = str(nalog.get('rn') or rn)
        client_name = (cli.get('name') if cli else (nalog.get('client') or ''))
        client_oib = (cli.get('oib') if cli else (nalog.get('client_id') or nalog.get('clientId') or nalog.get('oib') or ''))
        # SN odabrani prilikom kreiranja naloga
        try:
            sn_kvar_devices = [s.strip() for s in str(nalog.get('servisirani sn uredjaja','')).split(',') if s.strip()]
        except Exception:
            sn_kvar_devices = []
        try:
            sn_kvar_sims = [s.strip() for s in str(nalog.get('servisirani sn SIM-a','')).split(',') if s.strip()]
        except Exception:
            sn_kvar_sims = []
        # Fallback: ako nema selekcije u nalogu, uzmi aktivne kod klijenta
        try:
            akt = read_aktivni_uredjaji()
        except Exception:
            akt = []
        try:
            akt_sim = read_aktivni_sim()
        except Exception:
            akt_sim = []
        # Filtriraj na one koji su kod ovog klijenta
        def _belongs_to_client(item):
            try:
                return (str(item.get('client','')) == str(client_name)) or (client_oib and str(item.get('oib','')) == str(client_oib))
            except Exception:
                return False
        client_devs = [d for d in akt if d.get('active', True) and _belongs_to_client(d)]
        client_sims = [s for s in akt_sim if s.get('active', True) and _belongs_to_client(s)]
        model_by_sn = {str(d.get('serijski')): d.get('model','') for d in client_devs}
        provider_by_sn = {str(s.get('serijski')): s.get('provider','') for s in client_sims}
        # Intersekcija s onima kod klijenta (sigurnost)
        sn_kvar_devices = [s for s in sn_kvar_devices if s in model_by_sn]
        sn_kvar_sims = [s for s in sn_kvar_sims if s in provider_by_sn]

        if request.method == 'GET':
            kvar_dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in sn_kvar_devices] or []
            kvar_sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sn_kvar_sims] or []
            # učitaj i zamjenske opcije (zaduženo na tehničara)
            zdev = read_zaduzene_uredjaje() if 'read_zaduzene_uredjaje' in globals() else []
            zsim = read_zaduzene_sim() if 'read_zaduzene_sim' in globals() else []
            dev_opts = [d for d in zdev if (str(d.get('assigned_to') or '') == str(tech))]
            sim_opts = [s for s in zsim if (str(s.get('assigned_to') or '') == str(tech))]
            # uniq by serial
            seen=set(); dev_opts_u=[]
            for d in dev_opts:
                s = d.get('serijski');
                if s and s not in seen: dev_opts_u.append(d); seen.add(s)
            seen=set(); sim_opts_u=[]
            for s in sim_opts:
                ss = s.get('serijski');
                if ss and ss not in seen: sim_opts_u.append(s); seen.add(ss)
            return render_template('zakljuci.servis.html',
                                   title='Zaključi servisni nalog',
                                   username=current_user.username,
                                   nalog=nalog, klijent=cli,
                                   kvar_uredjaji=kvar_dev_rows,
                                   kvar_sims=kvar_sim_rows,
                                   zamjena_uredjaji=dev_opts_u,
                                   zamjena_sims=sim_opts_u)

        # POST
        if request.method == 'POST' and request.form.get('servis_action') == 'zakljuci':
            # SN-ovi koji idu na servis (iz više padajućih izbornika s istim imenom)
            selected_kvar_dev = [s.strip() for s in request.form.getlist('kvar_uredjaj') if s.strip()]
            selected_kvar_sim = [s.strip() for s in request.form.getlist('kvar_sim') if s.strip()]
            # Zamjenske koje uvodimo kod klijenta
            rep_devs = [s.strip() for s in request.form.getlist('uredjaj') if s and s.strip()]
            rep_sims = [s.strip() for s in request.form.getlist('sim') if s and s.strip()]

            # ZapIsnik slika obavezna
            file = request.files.get('zapisnik_img')
            if not file or not file.filename:
                flash('Obavezno je priložiti sliku zapisnika (JPG/PNG).', 'danger')
                # regen view data
                kvar_dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in sn_kvar_devices] or []
                kvar_sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sn_kvar_sims] or []
                # učitaj i zamjenske opcije
                zdev = read_zaduzene_uredjaje() if 'read_zaduzene_uredjaje' in globals() else []
                zsim = read_zaduzene_sim() if 'read_zaduzene_sim' in globals() else []
                dev_opts_u = [d for d in zdev if (str(d.get('assigned_to') or '') == str(tech))]
                # Isključi privremeno isključene iz zamjene
                try:
                    priv = read_privremeno_iskljuceni()
                except Exception:
                    priv = []
                _pns = {str(x.get('serijski')) for x in (priv or [])}
                dev_opts_u = [d for d in dev_opts_u if str(d.get('serijski')) not in _pns]
                sim_opts_u = [s for s in zsim if (str(s.get('assigned_to') or '') == str(tech))]
                sim_opts_u = [s for s in sim_opts_u if str(s.get('serijski')) not in _pns]
                return render_template('zakljuci.servis.html', title='Zaključi servisni nalog', username=current_user.username, nalog=nalog, klijent=cli, kvar_uredjaji=kvar_dev_rows, kvar_sims=kvar_sim_rows, zamjena_uredjaji=dev_opts_u, zamjena_sims=sim_opts_u)
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in ('.jpg','.jpeg','.png'):
                flash('Dozvoljeni formati slike: JPG/PNG.', 'danger')
                kvar_dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in sn_kvar_devices] or []
                kvar_sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sn_kvar_sims] or []
                zdev = read_zaduzene_uredjaje() if 'read_zaduzene_uredjaje' in globals() else []
                zsim = read_zaduzene_sim() if 'read_zaduzene_sim' in globals() else []
                dev_opts_u = [d for d in zdev if (str(d.get('assigned_to') or '') == str(tech))]
                sim_opts_u = [s for s in zsim if (str(s.get('assigned_to') or '') == str(tech))]
                sim_opts_u = [s for s in sim_opts_u if str(s.get('serijski')) not in _pns]
                return render_template('zakljuci.servis.html', title='Zaključi servisni nalog', username=current_user.username, nalog=nalog, klijent=cli, kvar_uredjaji=kvar_dev_rows, kvar_sims=kvar_sim_rows, zamjena_uredjaji=dev_opts_u, zamjena_sims=sim_opts_u)

            # spremi sliku u zapisnici/Servis
            ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_rn = re.sub(r'[^A-Za-z0-9_-]+', '_', str(rn_str))
            fname = f"{safe_rn}_{ts}{ext}"
            _dest_dir = os.path.join(STATIC_DIR, 'zapisnici', 'Servis')
            os.makedirs(_dest_dir, exist_ok=True)
            dest = os.path.join(_dest_dir, fname)
            file.save(dest)
            _saved_rel = os.path.join('static','zapisnici','Servis', fname)
            # === AUTO E-MAIL: po zahtjevu korisnika
            try:
                # Operater i RN
                _operater_username = getattr(current_user, 'username', '') or (nalog.get('assigned_to') or '')
                _rn_mail = str(rn_str)

                # Mapiranja za modele/providere iz zaduženih (zamjenski) i klijenta (zamijenjeni)
                try:
                    _zdev_all = read_zaduzene_uredjaje() if 'read_zaduzene_uredjaje' in globals() else []
                except Exception:
                    _zdev_all = []
                _rep_dev_model_by_sn = {str(d.get('serijski')): str(d.get('model') or '') for d in (_zdev_all or [])}

                try:
                    _zsim_all = read_zaduzene_sim() if 'read_zaduzene_sim' in globals() else []
                except Exception:
                    _zsim_all = []
                _rep_sim_provider_by_sn = {str(s.get('serijski')): str(s.get('provider') or '') for s in (_zsim_all or [])}

                # Sastavi tekst poruke
                _lines = []
                _lines.append(f'Operater {_operater_username} je zaključio servisni nalog sa sljedećim stavkama:')
                _lines.append('')
                # Zamijenjeni uređaji (bili kod korisnika -> idu na servis)
                if selected_kvar_dev:
                    _lines.append('Zamjenjeni uređaj:')
                    for _sn in selected_kvar_dev:
                        _mdl = str(model_by_sn.get(_sn, '') or '')
                        _lines.append(f'{_mdl} - SN - {_sn}')
                    _lines.append('')
                # Zamjenski uređaji (ostavljeni kod korisnika)
                if rep_devs:
                    _lines.append('Zamjenski uređaj:')
                    for _sn in rep_devs:
                        _mdl = str(_rep_dev_model_by_sn.get(_sn, '') or '')
                        _lines.append(f'{_mdl} - SN - {_sn}')
                    _lines.append('')
                # Zamijenjeni SIM-ovi
                if selected_kvar_sim:
                    _lines.append('Zamjenjeni SIM:')
                    for _sn in selected_kvar_sim:
                        _prov = str(provider_by_sn.get(_sn, '') or '')
                        _lines.append(f'{_prov} - SN - {_sn}')
                    _lines.append('')
                # Zamjenski SIM-ovi
                if rep_sims:
                    _lines.append('Zamjenski SIM:')
                    for _sn in rep_sims:
                        _prov = str(_rep_sim_provider_by_sn.get(_sn, '') or '')
                        _lines.append(f'{_prov} - SN - {_sn}')
                    _lines.append('')
                _lines.append('')
                _lines.append('Ovo je automatska poruka, molim Vas da ne odgovarate na nju. Vaš Billy Pos d.o.o.')
                _body = "\n".join(_lines)

                from flask_mail import Message as FMMessage
                import mimetypes
                _subject = f"Zaključen servisni nalog { _rn_mail }"
                msg = FMMessage(subject=_subject, recipients=['webtest806@gmail.com'])
                msg.body = _body
                _mtype, _enc = mimetypes.guess_type(dest)
                _mtype = _mtype or 'application/octet-stream'
                with open(dest, 'rb') as _fp:
                    msg.attach(filename=os.path.basename(dest), content_type=_mtype, data=_fp.read())
                mail.send(msg)
            except Exception as _mail_err:
                try:
                    flash(f'Upozorenje: e-mail obavijest nije poslana ({_mail_err}).', 'warning')
                except Exception:
                    pass
            # Ažuriraj servisni DOCX napomenom (ako placeholder postoji), inače dodaj paragraf
            try:
                file_rel = nalog.get('file') or ''
                doc_path = os.path.join(STATIC_DIR, file_rel) if not os.path.isabs(file_rel) else file_rel
                if os.path.exists(doc_path):
                    from docx import Document as _Doc
                    _doc = _Doc(doc_path)
                    napomena_txt = (request.form.get('napomena') or '').strip()
                    replaced = False
                    for p in _doc.paragraphs:
                        if '{{NAPOMENA}}' in p.text or '[[NAPOMENA]]' in p.text:
                            for run in p.runs:
                                run.text = run.text.replace('{{NAPOMENA}}', napomena_txt).replace('[[NAPOMENA]]', napomena_txt)
                            replaced = True
                    if not replaced:
                        _doc.add_paragraph(f"Napomena: {napomena_txt}")
                    _doc.save(doc_path)
            except Exception as e:
                try:
                    flash(f'Napomena nije zapisana u servisni zapisnik: {e}', 'warning')
                except Exception:
                    pass

            # Ukloni KVAR uređaj(e) i SIM(ove) s klijenta (klijenti.json + aktivne evidencije)
            try:
                kl = read_klijenti()
            except Exception:
                kl = []
            idx = None
            for i,c in enumerate(kl):
                if str(c.get('name')) == str(client_name):
                    idx = i; break
            if idx is not None:
                c = kl[idx]
                # sn_uredjaja
                try:
                    exist_dev = [s.strip() for s in (c.get('sn_uredjaja') or '').split(',') if s.strip()]
                except Exception:
                    exist_dev = []
                remain_dev = [s for s in exist_dev if s not in set(selected_kvar_dev)]
                c['sn_uredjaja'] = ", ".join(remain_dev)
                # model_uredjaja rekonstruiraj iz aktivnih preostalih
                models_remain = []
                for s in remain_dev:
                    m = model_by_sn.get(s, '')
                    if m: models_remain.append(m)
                c['model_uredjaja'] = ", ".join(sorted(set(models_remain)))
                # SIM
                try:
                    exist_sim = [s.strip() for s in (c.get('sn_SIM') or '').split(',') if s.strip()]
                except Exception:
                    exist_sim = []
                remain_sim = [s for s in exist_sim if s not in set(selected_kvar_sim)]
                c['sn_SIM'] = ", ".join(remain_sim)
                kl[idx] = c
                write_klijenti(kl)

            # deaktiviraj u aktivnim evidencijama
            try:
                akt = read_aktivni_uredjaji()
            except Exception:
                akt = []
            try:
                sel_set = set(selected_kvar_dev or [])
                before = len(akt)
                akt = [d for d in akt if str(d.get('serijski')) not in sel_set]
                if len(akt) != before:
                    write_aktivni_uredjaji(akt)
            except Exception:
                pass

            try:
                sel_set_s = set(selected_kvar_sim or [])
                akt_s = read_aktivni_sim()
                before_s = len(akt_s)
                akt_s = [row for row in akt_s if str(row.get('serijski')) not in sel_set_s]
                if len(akt_s) != before_s:
                    write_aktivni_sim(akt_s)
            except Exception:
                pass

            # Vrati KVAR SIM-ove u sim.json pool (opcionalno)
            if selected_kvar_sim:
                pool_sim = read_sim()
                seen_sim = {str(d.get('serijski')) for d in pool_sim}
                for s in selected_kvar_sim:
                    if s not in seen_sim:
                        pool_sim.append({'provider': provider_by_sn.get(s,''), 'serijski': s, 'created_at': datetime.datetime.now().isoformat()})
                        seen_sim.add(s)
                write_sim(pool_sim)

            # Ukloni KVAR uređaje iz drugih JSON-a (zaduženi, skladište) i premjesti u servis.uredjaji.JSON
            try:
                # remove from available store
                pool = read_uredjaji()
                pool = [u for u in pool if str(u.get('serijski')) not in set(selected_kvar_dev)]
                write_uredjaji(pool)
            except Exception:
                pass
            try:
                zlist = read_zaduzene_uredjaje()
                zlist = [z for z in zlist if str(z.get('serijski')) not in set(selected_kvar_dev)]
                write_zaduzene_uredjaje(zlist)
            except Exception:
                pass
            try:
                serv_list = read_servis_uredjaji()
            except Exception:
                serv_list = []
            for s in selected_kvar_dev:
                serv_list.append({
                    'serijski': s,
                    'model': model_by_sn.get(s,''),
                    'client': client_name,
                    'oib': client_oib,
                    'rn': rn_str,
                    'moved_at': datetime.datetime.now().isoformat(),
                    'status': 'na servisu'
                })
            write_servis_uredjaji(serv_list)

            # Uvedi zamjenski uređaj/SIM kod klijenta (slično kao instalacija)
            # (1) Uređaj
            if rep_devs:
                try:
                    zlist = read_zaduzene_uredjaje()
                except Exception:
                    zlist = []
                # Build model map from current zaduženi + fallback skladište
                model_map = {d.get('serijski'): d.get('model') for d in zlist}
                skladiste = read_uredjaji()
                for rep_dev in rep_devs:
                    # remove from zaduženi
                    zlist = [z for z in zlist if str(z.get('serijski')) != str(rep_dev)]
                    # fallback to skladiste for model if missing
                    if rep_dev not in model_map:
                        for it in skladiste:
                            if it.get('serijski') == rep_dev:
                                model_map[rep_dev] = it.get('model','')
                                break
                write_zaduzene_uredjaje(zlist)
                akt = read_aktivni_uredjaji()
                now = datetime.datetime.now().isoformat()
                for rep_dev in dict.fromkeys(rep_devs):
                    akt.append({'serijski': rep_dev,
                                'model': model_map.get(rep_dev,''),
                                'client': client_name,
                                'oib': client_oib,
                                'active': True,
                                'assigned_at': now})
                write_aktivni_uredjaji(akt)
            # (2) SIM
            if rep_sims:
                try:
                    zlist_sim = read_zaduzene_sim()
                except Exception:
                    zlist_sim = []
                provider_map = {s.get('serijski'): s.get('provider') for s in zlist_sim}
                skladiste_sim = read_sim()
                # remove selected SIMs from zaduženi and complete provider map
                for rep_sim in rep_sims:
                    zlist_sim = [z for z in zlist_sim if str(z.get('serijski')) != str(rep_sim)]
                    if rep_sim not in provider_map:
                        for it in skladiste_sim:
                            if it.get('serijski') == rep_sim:
                                provider_map[rep_sim] = it.get('provider','')
                                break
                write_zaduzene_sim(zlist_sim)
                akt_sim = read_aktivni_sim()
                now = datetime.datetime.now().isoformat()
                for rep_sim in dict.fromkeys(rep_sims):
                    akt_sim.append({'serijski': rep_sim,
                                    'provider': provider_map.get(rep_sim,''),
                                    'client': client_name,
                                    'oib': client_oib,
                                    'active': True,
                                    'assigned_at': now})
                write_aktivni_sim(akt_sim)

            # Zatvori nalog (zakljuceni.nalozi + nalozi.json status + ukloni iz otvorenih i zaduženih)
            try:
                zk = read_zakljuceni_nalozi()
            except Exception:
                zk = []
            closed = dict(nalog)
            # Ne spremati Zahtjev u zatvoreni zapis
            try:
                if 'zahtjev_image' in closed:
                    closed['zahtjev_image'] = ''
            except Exception:
                pass
            closed['status'] = 'zaključen'
            closed['closed_at'] = datetime.datetime.now().isoformat()
            closed['zapisnik_image'] = _saved_rel
            zk.append(closed)
            write_zakljuceni_nalozi(zk)

            # ukloni iz otvorenih skupova
            try:
                oi = read_otvorene_instalacije()
                oi = [o for o in oi if str(o.get('rn')) != str(rn_str)]
                write_otvorene_instalacije(oi)
            except Exception:
                pass
            try:
                od = read_otvorene_deinstalacije()
                od = [o for o in od if str(o.get('rn')) != str(rn_str)]
                write_otvorene_deinstalacije(od)
            except Exception:
                pass
            try:
                osv = read_otvoreni_servisi()
                osv = [o for o in osv if str(o.get('rn')) != str(rn_str)]
                write_otvoreni_servisi(osv)
            except Exception:
                pass

            try:
                # ukloni i iz zaduženih
                zn2 = read_zaduzene_nalozi()
                zn2 = [z for z in zn2 if str(z.get('rn')) != str(rn_str)]
                write_zaduzene_nalozi(zn2)
            except Exception:
                pass

            # status u nalozi.json
            try:
                base = read_nalozi()
            except Exception:
                base = []
            found=False
            for it in base:
                if str(it.get('rn')) == str(rn_str):
                    it['status'] = 'zaključen'
                    it['closed_at'] = closed['closed_at']
                    it['zapisnik_image'] = _saved_rel
                    found=True
                    break
            if not found:
                base.append(closed)
            write_nalozi(base)

            flash(f"Servis RN {rn_str} zaključen.", "success")
            if tech:
                return redirect(url_for('operater_profil', username=tech))
            return redirect(url_for('zaduzeni_nalozi'))
# load assigned devices/SIMs for that technician
    zdev = read_zaduzene_uredjaje() if 'read_zaduzene_uredjaje' in globals() else []
    zsim = read_zaduzene_sim() if 'read_zaduzene_sim' in globals() else []
    dev_opts = [d for d in zdev if (d.get('assigned_to') == tech)]
    sim_opts = [s for s in zsim if (s.get('assigned_to') == tech)]
    # ensure unique by serial
    seen = set(); dev_opts_u=[]
    for d in dev_opts:
        s = d.get('serijski');
        if s and s not in seen: dev_opts_u.append(d); seen.add(s)
    seen = set(); sim_opts_u=[]
    for s in sim_opts:
        ss = s.get('serijski');
        if ss and ss not in seen: sim_opts_u.append(s); seen.add(ss)
    if request.method == 'POST':
        # collect device(s) and sim(s)
        uredjaji_sel = request.form.getlist('uredjaj')
        sim_sel = request.form.getlist('sim')
        # strip empties, uniq
        uredjaji_sel = [u for u in dict.fromkeys([x.strip() for x in uredjaji_sel if x.strip()])]
        sim_sel = [u for u in dict.fromkeys([x.strip() for x in sim_sel if x.strip()])]
        file = request.files.get('zapisnik_img')
        allow_empty = (request.form.get('allow_empty') == '1')
        # zapisnik je uvijek obavezan
        if not file or not file.filename:
            flash('Obavezno je priložiti sliku zapisnika.', 'danger')
            return render_template('zakljuci.nalog.html', title='Zaključi nalog', username=current_user.username, nalog=nalog, klijent=cli, uredjaji=dev_opts_u, sims=sim_opts_u)
        # ako nije dozvoljen "prazan" unos, mora biti odabran barem jedan uređaj ILI jedan SIM
        if not allow_empty and not (uredjaji_sel or sim_sel):
            flash('Odaberite barem jedan Uređaj ili barem jedan SIM (ili označite kućicu za zaključivanje bez stavki).', 'danger')
            return render_template('zakljuci.nalog.html', title='Zaključi nalog', username=current_user.username, nalog=nalog, klijent=cli, uredjaji=dev_opts_u, sims=sim_opts_u)
        # validate file type
        name = file.filename
        ext = os.path.splitext(name)[1].lower()
        if ext not in ('.jpg','.jpeg','.png'):
            flash('Dozvoljeni formati slike: JPG/PNG.', 'danger')
            return render_template('zakljuci.nalog.html', title='Zaključi nalog', username=current_user.username, nalog=nalog, klijent=cli, uredjaji=dev_opts_u, sims=sim_opts_u)
        # save image
        ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_rn = re.sub(r'[^A-Za-z0-9_-]+', '_', str(rn))
        fname = f'{safe_rn}_{ts}{ext}'
        # NOVO: spremi zapisnik u static/zapisnici/<Vrsta>
        _type = str(nalog.get('type','')).strip().lower()
        if _type == 'instalacija':
            _subdir = 'Instalacija'
        elif _type == 'deinstalacija':
            _subdir = 'Deinstalacija'
        elif _type == 'servis':
            _subdir = 'Servis'
        else:
            _subdir = 'Ostalo'
        _dest_dir = os.path.join(STATIC_DIR, 'zapisnici', _subdir)
        os.makedirs(_dest_dir, exist_ok=True)
        dest = os.path.join(_dest_dir, fname)
        file.save(dest)
        # move nalog to nalozi.json with status closed
        ## MOVE-ZN-TO-ZK: premjesti zatvoreni nalog u zakljuceni.nalozi.JSON
        try:
            zk = read_zakljuceni_nalozi()
        except Exception:
            zk = []
        try:
            closed = dict(nalog)
            # Ne spremati Zahtjev (PDF/sliku) prilikom zaključka — samo e-mail ide
            try:
                if 'zahtjev_image' in closed:
                    closed['zahtjev_image'] = ''
            except Exception:
                pass
            zk.append(closed)
            write_zakljuceni_nalozi(zk)
        except Exception:
            pass
        # ukloni RN iz otvorene_instalacije.JSON nakon implementacije u zakljuceni.nalozi.JSON
        try:
            oi = read_otvorene_instalacije()
            oi = [o for o in oi if str(o.get('rn')) != str(rn)]
            write_otvorene_instalacije(oi)
        except Exception:
            pass
        # ukloni i iz otvorene_deinstalacije.JSON
        try:
            od = read_otvorene_deinstalacije()
            od = [o for o in od if str(o.get('rn')) != str(rn)]
            write_otvorene_deinstalacije(od)
        except Exception:
            pass
        # ukloni i iz otvoreni_servisi.JSON
        try:
            osv = read_otvoreni_servisi()
            osv = [o for o in osv if str(o.get('rn')) != str(rn)]
            write_otvoreni_servisi(osv)
        except Exception:
            pass
        zn = [z for z in zn if str(z.get('rn')) != str(rn)]
        write_zaduzene_nalozi(zn)
        base = read_nalozi()
        nalog_closed = dict(nalog)
        # Ne spremati Zahtjev (PDF/sliku) u zatvorenu kopiju koja ide u nalozi.json (kako se ne bi dupliralo na profilu)
        try:
            if 'zahtjev_image' in nalog_closed:
                nalog_closed['zahtjev_image'] = ''
        except Exception:
            pass
        nalog_closed['status'] = 'zaključen'
        nalog_closed['closed_at'] = datetime.datetime.now().isoformat()
        nalog_closed['closed_by'] = tech
        nalog_closed['devices_used'] = uredjaji_sel
        nalog_closed['sims_used'] = sim_sel
        _type = str(nalog.get('type','')).strip().lower()
        if _type == 'instalacija':
            _subdir2 = 'Instalacija'
        elif _type == 'deinstalacija':
            _subdir2 = 'Deinstalacija'
        elif _type == 'servis':
            _subdir2 = 'Servis'
        else:
            _subdir2 = 'Ostalo'
        nalog_closed['zapisnik_image'] = os.path.join('static','zapisnici', _subdir2, fname)
        base.append(nalog_closed)
        write_nalozi(base)
        ## AKTIVNI-TRANSFER: premjesti iz zaduzeni.uredjaji.json u aktivni_uredjaji.JSON
        try:
            # ukloni uređaje iz zaduženih (profil tehničara)
            zlist = read_zaduzene_uredjaje()
            z_ser = set(uredjaji_sel)
            zlist = [z for z in zlist if str(z.get('serijski')) not in z_ser]
            write_zaduzene_uredjaje(zlist)
            # upiši u aktivne kod klijenta
            akt = read_aktivni_uredjaji()
            # map serijski->model iz ponuđenih uređaja
            model_map = {d.get('serijski'): d.get('model') for d in dev_opts_u}
            client_name = (cli.get('name') if cli else (nalog.get('client') or ''))
            client_oib = (cli.get('oib') if cli else (nalog.get('client_id') or nalog.get('clientId') or nalog.get('oib') or ''))
            now = datetime.datetime.now().isoformat()
            for srl in uredjaji_sel:
                akt.append({'serijski': srl,
                            'model': model_map.get(srl,''),
                            'client': client_name,
                            'oib': client_oib,
                            'active': True,
                            'assigned_at': now})
            write_aktivni_uredjaji(akt)
        except Exception as e:
            # ne prekidaj tok, samo zabilježi u flash
            try:
                flash(f'Napomena: nije bilo moguće kompletno ažurirati aktivne uređaje ({e}).', 'warning')
            except Exception:
                pass


        ## AKTIVNI-SIM-TRANSFER: premjesti iz zaduzeni.sim.json u aktivni_sim.JSON
        try:
            zlist_sim = read_zaduzene_sim()
            z_ser_sim = set(sim_sel)
            # ukloni SIM-ove iz zaduženih (profil tehničara)
            zlist_sim = [z for z in zlist_sim if str(z.get('serijski')) not in z_ser_sim]
            write_zaduzene_sim(zlist_sim)
            # upiši u aktivne SIM kod klijenta
            akt_sim = read_aktivni_sim()
            provider_map = {s.get('serijski'): s.get('provider') for s in sim_opts_u}
            client_name = (cli.get('name') if cli else (nalog.get('client') or ''))
            client_oib = (cli.get('oib') if cli else (nalog.get('client_id') or nalog.get('clientId') or nalog.get('oib') or ''))
            now = datetime.datetime.now().isoformat()
            for srl in sim_sel:
                akt_sim.append({
                    'serijski': srl,
                    'provider': provider_map.get(srl, ''),
                    'client': client_name,
                    'oib': client_oib,
                    'active': True,
                    'assigned_at': now
                })
            write_aktivni_sim(akt_sim)
        except Exception as e:
            try:
                flash(f'Napomena: nije bilo moguće kompletno ažurirati aktivne SIM kartice ({e}).', 'warning')
            except Exception:
                pass


        
        # === AUTO E-MAIL: zaključena instalacija ===
        try:
            if (str(nalog.get('type','')).strip().lower() == 'instalacija') and app.config.get('MAIL_SERVER'):
                from flask_mail import Message
                # Map SN->model / provider from currently offered options
                model_map = { (d.get('serijski') or ''): (d.get('model') or '') for d in (locals().get('dev_opts_u') or []) }
                provider_map = { (s.get('serijski') or ''): (s.get('provider') or '') for s in (locals().get('sim_opts_u') or []) }
                # Compose body lines
                _lines = []
                _op = getattr(current_user, 'username', '') or (tech or '')
                _lines.append(f'Operater {_op} je zaključio instalacijski nalog sa sljedećim stavkama:')
                _lines.append('')
                # Uređaji
                if uredjaji_sel:
                    _lines.append('Uređaji:')
                    for sn in uredjaji_sel:
                        _model = model_map.get(sn, '')
                        if _model:
                            _lines.append(f'{_model} - SN - {sn}')
                        else:
                            _lines.append(f'SN - {sn}')
                    _lines.append('')
                # SIM
                if sim_sel:
                    _lines.append('SIM:')
                    for sn in sim_sel:
                        _prov = provider_map.get(sn, '')
                        if _prov:
                            _lines.append(f'{_prov} - SN - {sn}')
                        else:
                            _lines.append(f'SN - {sn}')
                    _lines.append('')
                _lines.append('Ovo je automatska poruka, molim Vas da ne odgovarate na nju.')
                _lines.append('')
                _lines.append('Vaš Billy Pos d.o.o.')
                body = "\n".join(_lines)
                subject = f"Zaključen instalacijski nalog {nalog.get('rn') or rn}"
                msg = Message(subject=subject, sender='webtest806@gmail.com', recipients=['webtest806@gmail.com'], body=body)
                # Attach zapisnik (the image just saved)
                try:
                    _abs = dest if os.path.isabs(dest) else os.path.join(app.root_path, dest)
                    import mimetypes
                    _mtype = mimetypes.guess_type(_abs)[0] or 'application/octet-stream'
                    with open(_abs, 'rb') as fp:
                        msg.attach(os.path.basename(_abs), _mtype, fp.read())
                except Exception:
                    pass
                try:
                    mail.send(msg)
                except Exception as _e:
                    app.logger.warning(f"Slanje e-pošte (instalacija) nije uspjelo: {_e}")
        except Exception as _outer_e:
            app.logger.warning(f'Greška pri pripremi e-maila (instalacija): {_outer_e}')
        flash(f'Nalog RN {rn} zaključen.', 'success')
        # back to operator profile
        if tech:
            return redirect(url_for('operater_profil', username=tech))
        return redirect(url_for('zaduzeni_nalozi'))
    return render_template('zakljuci.nalog.html', title='Zaključi nalog', username=current_user.username, nalog=nalog, klijent=cli, uredjaji=dev_opts_u, sims=sim_opts_u)


# -------------------- RUTE (općenite) --------------------
@app.route('/')
@login_required
def home():
    return redirect(url_for('naslovna'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username','').strip()
        password = request.form.get('password','').strip()
        user = User.query.filter(db.func.lower(User.username) == username.lower()).first()
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('home'))
        flash("Neispravno korisničko ime ili lozinka.", "danger")
    return render_template('login.html', username=None, title="Prijava")

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))




# -------------------- NASLOVNA (početna nakon prijave) --------------------
@app.route('/naslovna')
@login_required
def naslovna():
    return render_template('naslovna.html', title="Naslovna", username=getattr(current_user, 'username', None))
@app.route('/nalozi')
@login_required
def nalozi():



    all_orders = read_nalozi()
    # Otvorene instalacije -> iz otvorene_instalacije.JSON
    instalacije = read_otvorene_instalacije()
    try:
        instalacije.sort(key=lambda x: x.get('created_at',''), reverse=True)
    except Exception:
        pass
    # Otvoreni servisi -> iz otvoreni_servisi.JSON
    try:
        servisi = read_otvoreni_servisi()
    except Exception:
        servisi = []
    try:
        servisi.sort(key=lambda x: x.get('created_at',''), reverse=True)
    except Exception:
        pass
    # Zaključeni nalozi
    zakljuceni = [n for n in all_orders if (n.get('status') or '').lower() == 'zaključen']
    def _ctime(n):
        return (n.get('closed_at') or n.get('created_at') or '')
    try:
        zakljuceni.sort(key=_ctime, reverse=True)
    except Exception:
        pass
    assigned_map = { str(z.get('rn') or z.get('RN') or z.get('broj') or z.get('id') or z.get('number')): (z.get('assigned_to') or '') for z in read_zaduzene_nalozi() }
    # Deinstalacije u tijeku — PRIKAŽI SVE: zadužene + nezadužene, ali isključi one koje su već zaključene
    try:
        _zn_all = read_zaduzene_nalozi()
    except Exception:
        _zn_all = []
    try:
        _od_all = read_otvorene_deinstalacije()
    except Exception:
        _od_all = []
    try:
        _closed_all = read_zakljuceni_nalozi()
    except Exception:
        _closed_all = []

    # helper: normaliziraj RN (npr. 0001/2025)
    def _pure_rn_local(item):
        r = str(item.get('rn') or '')
        rp = str(item.get('rn_plain') or '')
        m = re.search(r'(\d{4}/\d{4})', r)
        return m.group(1) if m else (rp or r)

    # skup RN-ova deinstalacija koje su već zaključene (ne smiju na listu u tijeku)
    _closed_deinst_rns = {
        _pure_rn_local(z) for z in _closed_all
        if str((z.get('type') or '')).lower().strip() == 'deinstalacija'
    }

    # Zadužene deinstalacije (iz zaduzene_nalozi)
    _assigned_deinst = [
        z for z in _zn_all
        if (str(z.get('type') or '').lower() == 'deinstalacija') and _pure_rn_local(z) not in _closed_deinst_rns
    ]

    # Nezadužene deinstalacije (iz otvorenih), bez duplikata po RN i bez onih koje su zatvorene
    assigned_rns = { _pure_rn_local(x) for x in _assigned_deinst if x.get('rn') or x.get('rn_plain') }
    _unassigned_deinst = []
    for x in _od_all:
        try:
            t = str(x.get('type') or '').lower()
            if t not in ('instalacija','deinstalacija','servis'):
                file_name = str(x.get('file') or '').lower()
                if 'deinst' in file_name or 'deinstal' in file_name:
                    t = 'deinstalacija'
            if t == 'deinstalacija':
                rn_val = _pure_rn_local(x)
                if rn_val and rn_val not in assigned_rns and rn_val not in _closed_deinst_rns:
                    row = dict(x)
                    row.setdefault('assigned_to', '')
                    _unassigned_deinst.append(row)
        except Exception:
            pass

    # Spojeno
    dein_u_tijeku = list(_assigned_deinst) + list(_unassigned_deinst)

    # Sort: zaduženi prije nezaduženih; zatim po vremenu pa RN
    def _sort_key_deinst(it):
        is_assigned = 1 if (it.get('assigned_to') or '').strip() else 0
        ts = it.get('assigned_at') or it.get('created_at') or ''
        return (is_assigned, ts, str(it.get('rn','')))

    try:
        dein_u_tijeku.sort(key=_sort_key_deinst, reverse=True)
    except Exception:
        pass





        ## BADGE-DEINST — pripremi prikaz s mini oznakom za zatvorene deinstalacije
    zakljuceni_view = []
    try:
        for n in zakljuceni:
            d = dict(n)
            t = (d.get('type') or '').lower()
            rn_val = str(d.get('rn') or '')
            if t == 'deinstalacija' and rn_val:
                # Prefiks "DEINST" uz RN (mini badge efekt u tekstu)
                d['rn'] = f"(DEINST) {rn_val}"
            zakljuceni_view.append(d)
        zakljuceni = zakljuceni_view
    except Exception:
        pass
    
    # === Tehničar filtriranje: vidi isključivo naloge zadužene na njega ===
    try:
        _is_teh = has_role('tehnicar','tehničar')
    except Exception:
        _is_teh = False
    try:
        _is_privileged = has_role('superadmin','admin','prodaja','podrska','podrška','voditelj','serviser')
    except Exception:
        _is_privileged = False
    if _is_teh and not _is_privileged:
        _me = getattr(current_user, 'username', None) or getattr(current_user, 'name', None)
        def _mine(seq):
            out = []
            for _x in (seq or []):
                try:
                    at = str((_x.get('assigned_to') or _x.get('zaduzen') or _x.get('zadužen') or '-')).strip()
                except Exception:
                    at = '-'
                if _me and at and at.lower() == _me.lower():
                    out.append(_x)
            return out
        try:
            instalacije = _mine(instalacije)
        except Exception:
            pass
        try:
            servisi = _mine(servisi)
        except Exception:
            pass
        try:
            deinstalacije = _mine(deinstalacije)
        except Exception:
            pass
        try:
            zakljuceni = _mine(zakljuceni)
        except Exception:
            pass
    return render_template('nalozi.html', title="Nalozi", username=current_user.username, instalacije=instalacije, servisi=servisi, zakljuceni=zakljuceni, assigned_map=assigned_map, deinstalacije_u_tijeku=dein_u_tijeku)

@app.route('/zaduzeni-nalozi')
@login_required
def zaduzeni_nalozi():
    items = read_zaduzene_nalozi()
    try:
        items = sorted(items, key=lambda x: (x.get('assigned_at',''), str(x.get('rn',''))), reverse=True)
    except Exception:
        pass
    return render_template('zaduzeni.nalozi.html', title="Zaduženi nalozi", username=current_user.username, nalozi=items)

@app.route('/aktivni-uredjaji')
@login_required
def aktivni_uredjaji():
    lst = read_aktivni_uredjaji()
    # normalize missing fields
    for d in lst:
        d.setdefault('active', True)
    # sort by client name then model then serial
    try:
        lst.sort(key=lambda x: (str(x.get('client','')), str(x.get('model','')), str(x.get('serijski',''))))
    except Exception:
        pass
    return render_template('aktivni.uredjaji.html',
                           title="Aktivni uređaji",
                           username=current_user.username,
                           uredjaji=lst)

@app.route('/datoteke')
@login_required
def datoteke():
    return render_template('datoteke.html', title="Datoteke", username=current_user.username)

@app.route('/info')
@login_required
def info():
    return render_template('info.html', title="Info", username=current_user.username)

# -------------------- OPERATERI --------------------
@app.route('/operateri')
@login_required
def operateri():
    ops = read_operateri()
    # Ako je korisnik Voditelj, filtriraj samo Tehničare
    try:
        if has_role('voditelj'):
            ops = [o for o in ops if str(o.get('role','')).strip().lower() in ('tehničar','tehnicar')]
    except Exception:
        pass
    ops.sort(key=lambda x: x.get('role',''))
    return render_template('operateri.html', title="Operateri", username=current_user.username, operators=ops)


@app.route('/kreiraj-profil', methods=['GET','POST'])
@login_required
def kreiraj_profil():
    if not current_user.is_superadmin:
        flash("Samo superadmin može kreirati profile.", "warning")
        return redirect(url_for('operateri'))
    if request.method == 'POST':
        first_name = request.form.get('first_name','').strip()
        last_name = request.form.get('last_name','').strip()
        username = request.form.get('username','').strip()
        password = request.form.get('password','').strip()
        email = request.form.get('email','').strip()
        phone = request.form.get('phone','').strip()
        role = request.form.get('role','').strip()
        if not all([first_name,last_name,username,password,email,phone,role]):
            flash("Sva polja su obavezna.", "danger")
            return render_template('kreiraj_profil.html', title="Kreiraj profil", username=current_user.username)
        ops = read_operateri()
        if any(o['username']==username or o['email']==email for o in ops):
            flash("Korisničko ime ili Email već postoji.", "danger")
            return render_template('kreiraj_profil.html', title="Kreiraj profil", username=current_user.username)
        new_op = dict(first_name=first_name,last_name=last_name,
                      username=username,email=email,phone=phone,role=role)
        ops.append(new_op)
        write_operateri(ops)

        # === DB upsert for new operator ===
        try:
            db.create_all()
        except Exception:
            pass
        try:
            role_db = _normalize_role_for_db(role)
            u = User.query.filter(db.func.lower(User.username) == username.lower()).first()
            if u is None:
                u = User(username=username, role=role_db or 'user')
                u.set_password(password)
                db.session.add(u)
            else:
                if role_db:
                    u.role = role_db
                if password:
                    u.set_password(password)
            db.session.commit()
        except Exception as _e:
            try:
                print('[kreiraj_profil] DB sync error:', _e)
            except Exception:
                pass
        flash("Uspješno kreiran", "success")
        return redirect(url_for('operateri'))
    return render_template('kreiraj_profil.html', title="Kreiraj profil", username=current_user.username)

@app.route('/operater/<username>')
@login_required
def operater_profil(username):
    ops = read_operateri()
    op = next((o for o in ops if o['username']==username), None)
    if not op: abort(404)
    zlist = read_zaduzene_uredjaje()
    my_devices = [z for z in zlist if z.get('assigned_to') == username]
    my_devices.sort(key=lambda x: (x.get('model',''), x.get('serijski','')))
    # --- SIM ---
    zsim = read_zaduzene_sim()
    my_sims = [z for z in zsim if z.get('assigned_to') == username]
    my_sims.sort(key=lambda x: (x.get('provider',''), x.get('serijski','')))
    # --- NALOZI ---
    zn = read_zaduzene_nalozi()
    my_nalozi = [z for z in zn if (str(z.get('assigned_to') or z.get('assignedTo') or z.get('tehnicar') or '')).lower() == username.lower()]
    my_nalozi.sort(key=lambda x: (str(x.get('rn','')), x.get('client','')))
    return render_template(
        'operater_profil.html',
        title="Profil operatera",
        username=current_user.username,
        operater=op,
        zaduzeni_uredjaji=my_devices,
        zaduzeni_sim=my_sims,
        zaduzeni_nalozi=my_nalozi,
    )

@app.route('/api/operater/<username>/avatar', methods=['POST'])
@login_required
def api_operater_avatar(username):
    # Dozvoljeno: superadmin/admin ili sam korisnik
    if not (getattr(current_user, 'is_superadmin', False) or current_user.username == username or has_role('admin')):
        return jsonify({'ok': False, 'error': 'Zabranjeno.'}), 403
    file = request.files.get('avatar')
    if not file or not getattr(file, 'filename', ''):
        return jsonify({'ok': False, 'error': 'Nije odabrana slika.'}), 400
    # Provjera ekstenzije
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ('.jpg', '.jpeg', '.png', '.gif', '.webp'):
        return jsonify({'ok': False, 'error': 'Dozvoljeni formati: JPG, PNG, GIF, WEBP.'}), 400
    # Sigurno ime datoteke -> username + ext
    uname = re.sub(r'[^A-Za-z0-9_.-]+', '_', str(username))
    fname = f"{uname}{ext}"
    try:
        os.makedirs(PROFILE_DIR, exist_ok=True)
    except Exception:
        pass
    dest = os.path.join(PROFILE_DIR, fname)
    try:
        file.save(dest)
    except Exception as e:
        return jsonify({'ok': False, 'error': f'Ne mogu spremiti sliku: {e}'}), 500
    # Relativna putanja za url_for('static', filename=...)
    rel = os.path.join('avatars', fname).replace('\\', '/')
    # Upis u operateri bazu (ako postoje ove helper funkcije)
    try:
        ops = read_operateri() if 'read_operateri' in globals() else []
        updated = False
        for o in ops:
            try:
                if str(o.get('username') or '').strip().lower() == str(username).strip().lower():
                    o['avatar'] = rel
                    updated = True
                    break
            except Exception:
                pass
        if updated and 'write_operateri' in globals():
            write_operateri(ops)
    except Exception:
        pass
    # Vratimo i apsolutni URL za prikaz (frontendu može trebati)
    try:
        abs_url = url_for('static', filename=rel)
    except Exception:
        abs_url = f"/static/{rel}"
    return jsonify({'ok': True, 'avatar': rel, 'url': abs_url}), 200


@app.route('/api/operater/<username>/signature', methods=['POST'])
@login_required
def api_operater_signature(username):
    # RBAC: superadmin/admin ili isti korisnik
    if not (getattr(current_user, 'is_superadmin', False) or has_role('admin') or getattr(current_user, 'username', None) == username):
        return jsonify({'ok': False, 'error': 'Zabranjeno.'}), 403
    file = request.files.get('signature')
    if not file or not getattr(file, 'filename', ''):
        return jsonify({'ok': False, 'error': 'Nije odabrana slika potpisa.'}), 400
    ext = os.path.splitext(file.filename)[1].lower()
    if ext != '.png':
        return jsonify({'ok': False, 'error': 'Potpis mora biti PNG (.png).'}), 400
    # safe name: username.png
    uname = re.sub(r'[^A-Za-z0-9_.-]+', '_', str(username))
    fname = f"{uname}.png"
    try:
        os.makedirs(SIGNATURE_DIR, exist_ok=True)
        dest = os.path.join(SIGNATURE_DIR, fname)
        file.save(dest)
    except Exception as e:
        return jsonify({'ok': False, 'error': f'Ne mogu spremiti potpis: {e}'}), 500
    rel = os.path.join('signatures', fname).replace('\\','/')
    # upiši u operateri.json
    try:
        ops = read_operateri()
        for o in ops:
            if str(o.get('username') or '').strip().lower() == str(username).strip().lower():
                o['signature'] = rel
                break
        write_operateri(ops)
    except Exception:
        pass
    try:
        abs_url = url_for('static', filename=rel)
    except Exception:
        abs_url = f"/static/{rel}"
    return jsonify({'ok': True, 'signature': rel, 'url': abs_url})



@app.route('/api/operater/<username>/update', methods=['POST'])
@login_required
def api_operater_update(username):
    # Dozvoljeno: superadmin/admin ili sam korisnik
    if not (getattr(current_user, 'is_superadmin', False) or has_role('admin') or current_user.username == username):
        return jsonify({'ok': False, 'error': 'Zabranjeno.'}), 403

    first_name = (request.form.get('first_name') or '').strip()
    last_name  = (request.form.get('last_name') or '').strip()
    email      = (request.form.get('email') or '').strip()
    password   = (request.form.get('password') or '').strip()

    ops = read_operateri()
    idx = next((i for i,o in enumerate(ops) if str(o.get('username')) == str(username)), None)
    if idx is None:
        return jsonify({'ok': False, 'error': 'Operater nije pronađen.'}), 404

    if first_name: ops[idx]['first_name'] = first_name
    if last_name:  ops[idx]['last_name']  = last_name
    if email:      ops[idx]['email']      = email

    try:
        write_operateri(ops)
    except Exception:
        return jsonify({'ok': False, 'error': 'Greška pri spremanju.'}), 500

    # Sinkroniziraj (ili kreiraj) login račun u User tablici
    try:
        db.create_all()
    except Exception:
        pass
    try:
        role_req = (request.form.get('role') or '').strip()
        role_db = _normalize_role_for_db(role_req) if role_req else None
        u = User.query.filter(db.func.lower(User.username) == username.lower()).first()
        if u is None:
            if password:
                u = User(username=username, role=role_db or 'user')
                u.set_password(password)
                db.session.add(u)
                db.session.commit()
        else:
            changed = False
            if password:
                u.set_password(password); changed = True
            if role_db and role_db != (u.role or ''):
                u.role = role_db; changed = True
            if changed:
                db.session.commit()
    except Exception as _e:
        try:
            print('[api_operater_update] DB sync error:', _e)
        except Exception:
            pass
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True}), 200






@app.route('/api/operateri/avatars', methods=['GET'])
def api_operateri_avatars():
    # Javni endpoint: vraća mapu username -> avatar (ako postoji) ili None
    ops = read_operateri()
    data = {}
    for o in ops:
        u = str(o.get('username', ''))
        av = o.get('avatar')
        if av:
            data[u] = url_for('static', filename=av)
        else:
            data[u] = None
    return jsonify({'ok': True, 'avatars': data}), 200

@app.route('/profile/<username>')
@login_required
def profile(username):
    return operater_profil(username)

# -------------------- KLIJENTI --------------------
@app.route('/klijenti')
@login_required
def klijenti():

    kl = read_klijenti()
    kl.sort(key=lambda x: x.get('name',''))
    # Izračun aktivnih klijenata (ima barem 1 uređaj ili SIM)
    try:
        akt_dev = read_aktivni_uredjaji() or []
    except Exception:
        akt_dev = []
    try:
        akt_sim = read_aktivni_sim() or []
    except Exception:
        akt_sim = []
    # skupovi po imenu i po OIB-u
    dev_names = {str(d.get('client','')) for d in akt_dev if d.get('active', True)}
    dev_oibs  = {str(d.get('oib','')) for d in akt_dev if d.get('active', True) and d.get('oib')}
    sim_names = {str(s.get('client','')) for s in akt_sim if s.get('active', True)}
    sim_oibs  = {str(s.get('oib','')) for s in akt_sim if s.get('active', True) and s.get('oib')}
    # --- Override status: ako su SVI uređaji klijenta u 'privremeno isključenim', klijent je NEAKTIVAN ---
    try:
        priv = read_privremeno_iskljuceni() or []
    except Exception:
        priv = []
    import re as _re2
    def _norm_sn(v): return _re2.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()
    priv_set = {_norm_sn(p.get('serijski')) for p in priv}

    aktivni_count = 0
    neaktivni_count = 0
    for c in kl:
        nm = str(c.get('name',''))
        oib = str(c.get('oib',''))
        is_active = (nm in dev_names) or (oib in dev_oibs) or (nm in sim_names) or (oib in sim_oibs)
    # provjeri uređaje klijenta i privremeno isključene
        dev_sns = [ _norm_sn(d.get('serijski')) for d in akt_dev if (str(d.get('client',''))==nm or (oib and str(d.get('oib',''))==oib)) ]
        if dev_sns:
            if all(sn in priv_set for sn in dev_sns):
                is_active = False
        c['active'] = bool(is_active)
        if c['active']:
            aktivni_count += 1
        else:
            neaktivni_count += 1

    return render_template('klijenti.html',
                           title="Klijenti",
                           username=current_user.username,
                           klijenti=kl,
                           aktivni_count=aktivni_count,
                           neaktivni_count=neaktivni_count)

@app.route('/klijent/<name>')
@login_required
def klijent_profil(name):
    klijenti = read_klijenti()
    k = next((c for c in klijenti if c['name'] == name), None)
    if not k: abort(404)
    orders = [n for n in read_nalozi() if n.get('client') == k['name']]
    orders.sort(key=lambda x: x.get('created_at',''), reverse=True)

    # UREĐAJI KOD KLIJENTA (iz aktivni_uredjaji.JSON)
    try:
        aktivni = read_aktivni_uredjaji()
    except Exception:
        aktivni = []
    client_name = k.get('name') if isinstance(k, dict) else (k.name if hasattr(k, 'name') else str(name))
    client_oib = (k.get('oib') if isinstance(k, dict) else None)
    uredjaji_klijent = [d for d in aktivni if (str(d.get('client','')) == str(client_name)) or (client_oib and str(d.get('oib','')) == str(client_oib))]
    # prikazuj samo aktivne (default True ako nema ključa)
    try:
        uredjaji_klijent = [d for d in uredjaji_klijent if d.get('active', True)]
        uredjaji_klijent.sort(key=lambda x: (str(x.get('model','')), str(x.get('serijski',''))))
    except Exception:
        pass


    
    # --- DODANO: Namjena uređaja (Kupnja/Najam) po SN iz static/uredjaji.JSON za prikaz u profilu klijenta ---
    try:
        _uredjaji_base = read_uredjaji()  # list from static/uredjaji.JSON
    except Exception:
        _uredjaji_base = []
    try:
        _namjena_by_sn = {str(x.get('serijski')): str(x.get('namjena', '')) for x in (_uredjaji_base or [])}
    except Exception:
        _namjena_by_sn = {}
    try:
        for _row in uredjaji_klijent or []:
            try:
                _sn = str(_row.get('serijski') or '')
                _row['namjena'] = _namjena_by_sn.get(_sn, '')
            except Exception:
                _row['namjena'] = ''
    except Exception:
        pass

    # --- STATUS ZA BOX 'Uređaji' NA PROFILU KLIJENTA (Aktivan/Privremeno isključen) ---
    try:
        _priv = read_privremeno_iskljuceni()
    except Exception:
        _priv = []
    import re as _re
    _pset = {_re.sub(r'[^A-Za-z0-9]', '', str(p.get('serijski') or '')).upper() for p in _priv}
    def _norm_sn(v):
        return _re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()
    try:
        for _row in uredjaji_klijent or []:
            try:
                if _norm_sn(_row.get('serijski')) in _pset:
                    _row['status_color'] = 'orange'
                    _row['status_label'] = 'Privremeno isključen'
                else:
                    _row['status_color'] = 'green'
                    _row['status_label'] = 'Aktivan kod klijenta'
            except Exception:
                pass
    except Exception:
        pass

    # SIM KOD KLIJENTA (iz aktivni_sim.JSON)
    try:
        aktivni_sim = read_aktivni_sim()
    except Exception:
        aktivni_sim = []
    sims_klijent = [s for s in aktivni_sim if (str(s.get('client','')) == str(client_name)) or (client_oib and str(s.get('oib','')) == str(client_oib))]
    try:
        sims_klijent = [s for s in sims_klijent if s.get('active', True)]
        sims_klijent.sort(key=lambda x: (str(x.get('provider','')), str(x.get('serijski',''))))
    except Exception:
        pass


    
    # --- STATUS kolona za BOX 'SIM' na profilu klijenta (iz sim.json) ---
    try:
        _all_sim_rows = read_sim()
    except Exception:
        _all_sim_rows = []
    try:
        _status_by_sn = {str(r.get('serijski')): str(r.get('status') or '') for r in (_all_sim_rows or [])}
    except Exception:
        _status_by_sn = {}
    def _norm_sim_status(v):
        v = str(v or '').strip().lower()
        return 'Privremeno isključen' if 'privremeno' in v else 'Aktivan'
    try:
        for _row in (sims_klijent or []):
            _st = _status_by_sn.get(str(_row.get('serijski') or ''), 'aktivan')
            _row['status_label'] = _norm_sim_status(_st)
    except Exception:
        pass

# --- ZAPISNICI (instalacija/deinstalacija/servis) za ovog klijenta ---
    try:
        all_orders_k = read_nalozi()
    except Exception:
        all_orders_k = []
    # Filtriraj: samo za ovog klijenta, status zaključen i postoji zapisnik slika
    def _rel_static(p):
        p = str(p or '')
        # Normalize slashes and remove leading 'static/' or 'static\'
        p = p.replace('\\', '/')
        if p.startswith('static/'):
            return p[7:]
        return p
    def _row(n):
        src_rel = _rel_static(n.get('zapisnik_image'))
        return {'src': src_rel, 'rn': n.get('rn','')}
    _client_cmp = str(k.get('name') if isinstance(k, dict) else name)
    closed_for_client = []
    for n in all_orders_k:
        try:
            if str(n.get('client','')) == _client_cmp and (str(n.get('status','')).lower() == 'zaključen') and n.get('zapisnik_image'):
                closed_for_client.append(n)
        except Exception:
            pass
    def _sort_key(n):
        return (n.get('closed_at') or n.get('created_at') or '')
    try:
        closed_for_client.sort(key=_sort_key, reverse=True)
    except Exception:
        pass
    # Robusno razvrstaj zapisnike po vrsti:
    # 1) primarno prema putanji slike (zapisnici/Instalacija|Deinstalacija|Servis)
    # 2) fallback prema polju 'type'
    inst_zapisnici = []
    deinst_zapisnici = []
    servis_zapisnici = []
    for _n in closed_for_client:
        _t = str(_n.get('type','')).lower()
        _p = str(_n.get('zapisnik_image') or '')
        _pn = _p.replace('\\','/').lower()
        _row_item = _row(_n)
        if '/deinstalacija/' in _pn:
            deinst_zapisnici.append(_row_item)
        elif '/instalacija/' in _pn:
            inst_zapisnici.append(_row_item)
        elif '/servis' in _pn:
            servis_zapisnici.append(_row_item)
        elif _t == 'deinstalacija':
            deinst_zapisnici.append(_row_item)
        elif _t == 'servis':
            servis_zapisnici.append(_row_item)
        else:
            inst_zapisnici.append(_row_item)


    # --- STATUS: aktivan ako ima barem 1 aktivan uređaj ili SIM ---
    try:
        client_active = bool(uredjaji_klijent) or bool(sims_klijent)
    except Exception:
        client_active = False
        # --- OVERRIDE: ako su SVI uređaji privremeno isključeni, klijent je NEAKTIVAN ---
    try:
        if isinstance(uredjaji_klijent, list) and len(uredjaji_klijent) > 0:
            _any_active_device = any(str(r.get('status_color')) == 'green' for r in uredjaji_klijent)
            if not _any_active_device:
                client_active = False
    except Exception:
        pass



    # --- LOG PROMJENA (aktivnosti) za ovog klijenta ---
    client_log = []
    _client_name = k.get('name') if isinstance(k, dict) else str(name)
    _client_oib = k.get('oib') if isinstance(k, dict) else None
    now_iso = datetime.datetime.now().isoformat()

    # Helper: push entry
    def _push(evt_type, desc, when=None, operator='-'):
        ts = when or now_iso
        try:
            # Normalize timestamp to first 19 chars if ISO-like
            if isinstance(ts, str) and len(ts) >= 19:
                ts_short = ts[:19]
            else:
                ts_short = str(ts)
        except Exception:
            ts_short = str(ts or '')
        client_log.append({'ts': ts_short, 'type': evt_type, 'desc': desc, 'operator': operator})

    # 1) Iz zaključenih/otvorenih naloga
    try:
        all_n = read_nalozi() or []
    except Exception:
        all_n = []
    try:
        all_zk = read_zakljuceni_nalozi() or []
    except Exception:
        all_zk = []
    try:
        # Kombiniraj, ali kasniji zapisi mogu imati dodatna polja (devices_used, closed_by...)
        # prioritiziraj kasnije zapise koristeći RN kao ključ
        combined = []
        seen = set()
        for lst in (all_n, all_zk):
            for it in lst:
                try:
                    c_ok = (str(it.get('client','')) == str(_client_name)) or (_client_oib and str(it.get('oib','')) == str(_client_oib))
                    if not c_ok:
                        continue
                    key = (str(it.get('type','')).lower(), str(it.get('rn','')))
                    if key in seen:
                        continue
                    seen.add(key)
                    combined.append(dict(it))
                except Exception:
                    pass
    except Exception:
        combined = []

    for it in combined:
        t = (it.get('type') or '').lower()
        when = it.get('closed_at') or it.get('created_at')
        oper = it.get('closed_by') or it.get('assigned_to') or '-'
        if t == 'instalacija':
            devs = it.get('devices_used') or []
            sims = it.get('sims_used') or []
            d_txt = (', '.join([str(x) for x in devs]) if devs else '-')
            s_txt = (', '.join([str(x) for x in sims]) if sims else '-')
            desc = f"Instalacija — uređaji: {d_txt}; SIM: {s_txt}; RN: {it.get('rn','')}"
            _push('instalacija', desc, when, oper)
        elif t == 'deinstalacija':
            # polja iz deinst naloga
            devs = it.get('deinstalirani sn uredjaja') or it.get('deinstalirani sn uređaja') or ''
            sims = it.get('deinstalirani sn SIM-a') or ''
            d_txt = str(devs) if devs else '-'
            s_txt = str(sims) if sims else '-'
            desc = f"Deinstalacija — vraćeni uređaji: {d_txt}; vraćeni SIM: {s_txt}; RN: {it.get('rn','')}"
            _push('deinstalacija', desc, when, oper)
        elif t == 'servis':
            devs = it.get('servisirani sn uredjaja') or ''
            sims = it.get('servisirani sn SIM-a') or ''
            d_txt = str(devs) if devs else '-'
            s_txt = str(sims) if sims else '-'
            desc = f"Servis — uređaji/SN: {d_txt}; SIM: {s_txt}; RN: {it.get('rn','')}"
            _push('servis', desc, when, oper)

    # 2) Iz aktivnih evidencija — assigned/unassigned trenuci
    try:
        all_akt = read_aktivni_uredjaji() or []
    except Exception:
        all_akt = []
    try:
        all_sim = read_aktivni_sim() or []
    except Exception:
        all_sim = []

    for d in all_akt:
        try:
            if (str(d.get('client','')) == str(_client_name)) or (_client_oib and str(d.get('oib','')) == str(_client_oib)):
                if d.get('assigned_at'):
                    _push('dodjela-uredjaja', f"Uređaj {d.get('model','') or ''} SN {d.get('serijski','')} dodijeljen", d.get('assigned_at'), '-')
                if d.get('unassigned_at'):
                    _push('razduzenje-uredjaja', f"Uređaj SN {d.get('serijski','')} razdužen", d.get('unassigned_at'), '-')
        except Exception:
            pass

    for s in all_sim:
        try:
            if (str(s.get('client','')) == str(_client_name)) or (_client_oib and str(s.get('oib','')) == str(_client_oib)):
                if s.get('assigned_at'):
                    _push('dodjela-sim', f"SIM {s.get('provider','') or ''} SN {s.get('serijski','')} dodijeljen", s.get('assigned_at'), '-')
                if s.get('unassigned_at'):
                    _push('razduzenje-sim', f"SIM SN {s.get('serijski','')} razdužen", s.get('unassigned_at'), '-')
        except Exception:
            pass

    # Sortiraj silazno po vremenu
    try:
        client_log.sort(key=lambda x: x.get('ts') or '', reverse=True)
    except Exception:
        pass
    # ZAHTJEVI (slike zahtjeva) – iz otvorenih i zaključenih naloga
    zahtjevi = []
    try:
        _client_cmp = str(k.get('name') if isinstance(k, dict) else name)
        _all_src = []
        try:
            _all_src += all_n
        except Exception:
            pass
        try:
            _all_src += all_zk
        except Exception:
            pass
        for _n in _all_src:
            try:
                if str(_n.get('client','')) == _client_cmp and _n.get('zahtjev_image'):
                    _src = str(_n.get('zahtjev_image')).replace('\\','/')
                    if _src.startswith('static/'):
                        _src = _src[7:]
                    _ts = _n.get('created_at') or _n.get('closed_at') or ''
                    kind = (str((_n.get('type') or '')).lower() or ('deinstalacija' if 'deinst' in str((_n.get('file') or '')).lower() else 'instalacija'))
                    zahtjevi.append({'src': _src, 'rn': _n.get('rn',''), 'ts': _ts, 'kind': kind})
            except Exception:
                pass
        try:
            zahtjevi.sort(key=lambda x: x.get('ts') or '', reverse=True)
        except Exception:
            pass
    except Exception:
        zahtjevi = []
    # === STATUS KLIJENTA (Aktivan/Neaktivan) prema pravilima ===
    # Aktivan ako postoji BAREM 1 aktivan uređaj ILI BAREM 1 aktivan SIM.
    # Privremeno isključen SIM ne računa se kao aktivan.
    try:
        # Uređaji: računaj samo one koji su stvarno aktivni kod klijenta (status_color 'green')
        has_active_device = any((str(getattr(d, 'get', lambda k, default=None: d[k])('status_color', '')) == 'green') for d in (uredjaji_klijent or []))
    except Exception:
        has_active_device = False
    try:
        # SIM: aktivan ako status_label (iz sim.json) mapira na 'Aktivan'
        has_active_sim = any((str((s.get('status_label') if isinstance(s, dict) else getattr(s, 'status_label', '')) or '').strip().lower() == 'aktivan') for s in (sims_klijent or []))
    except Exception:
        has_active_sim = False
    client_active = bool(has_active_device or has_active_sim)

    
    return render_template('klijent_profil.html',
        title=f"Profil klijenta - {k['name']}",
                               username=current_user.username, klijent_status=('Aktivan' if client_active else 'Neaktivan'), client_active=client_active, client_log=client_log,
                               klijent=k, nalozi=orders, aktivni_uredjaji=uredjaji_klijent, uredjaji_klijenta=uredjaji_klijent, aktivni_simovi=sims_klijent, sims_klijenta=sims_klijent, sims=sims_klijent, simovi=sims_klijent, simovi_klijenta=sims_klijent, inst_zapisnici=inst_zapisnici, deinst_zapisnici=deinst_zapisnici, servis_zapisnici=servis_zapisnici, zahtjevi=zahtjevi)

@app.route('/api/klijent/<name>/sim', methods=['GET'])
@login_required
def api_klijent_sim(name):
    try:
        klijenti = read_klijenti()
        k = next((c for c in klijenti if c.get('name') == name), None)
        if not k:
            abort(404)
        try:
            aktivni_sim = read_aktivni_sim()
        except Exception:
            aktivni_sim = []
        client_name = k.get('name') if isinstance(k, dict) else str(name)
        client_oib = k.get('oib') if isinstance(k, dict) else None
        sims_klijent = [s for s in aktivni_sim if (str(s.get('client','')) == str(client_name)) or (client_oib and str(s.get('oib','')) == str(client_oib))]
        try:
            sims_klijent = [s for s in sims_klijent if s.get('active', True)]
            sims_klijent.sort(key=lambda x: (str(x.get('provider','')), str(x.get('serijski',''))))
        except Exception:
            pass
        return jsonify({'sims': sims_klijent, 'count': len(sims_klijent)}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/klijent', methods=['POST'])
@login_required
def api_klijent_save():
    # RBAC: samo superadmin/admin/prodaja smiju dodavati klijenta
    if not has_role('superadmin','admin','prodaja'):
        return jsonify({'ok': False, 'error': 'Zabranjeno: nedovoljna prava.'}), 403
    data = request.get_json(silent=True) or {}
    required = ['name','oib','headquarters','email','phone']
    if not all(data.get(k) for k in required):
        return jsonify({'ok': False, 'error': 'Sva polja osim adrese isporuke su obavezna.'}), 400
    if not re.fullmatch(r'\d{11}', data.get('oib','')):
        return jsonify({'ok': False, 'error': 'OIB mora imati točno 11 znamenki.'}), 400
    kl = read_klijenti()
    if any(c.get('name') == data['name'] for c in kl):
        return jsonify({'ok': False, 'error': 'Klijent s tim nazivom već postoji.'}), 409
    data['created_at'] = datetime.datetime.now().isoformat()
    kl.append({
        'name': data['name'],
        'oib': data['oib'],
        'headquarters': data['headquarters'],
        'shipping': data.get('shipping',''),
        'email': data['email'],
        'phone': data['phone'],
        'created_at': data['created_at']
    })
    write_klijenti(kl)
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True}), 200

@app.route('/api/klijent/<name>', methods=['POST'])
@login_required
def api_klijent_update(name):
    # RBAC: samo superadmin/admin/prodaja smiju uređivati osnovne podatke
    if not has_role('superadmin','admin','prodaja'):
        return jsonify({'ok': False, 'error': 'Zabranjeno: nedovoljna prava.'}), 403
    data = request.get_json(silent=True) or {}
    required = ['name','oib','headquarters','email','phone']
    if not all(data.get(k) for k in required):
        return jsonify({'ok': False, 'error': 'Sva obavezna polja moraju biti popunjena.'}), 400
    # Validacija OIB-a
    if not re.fullmatch(r'\d{11}', data.get('oib','')):
        return jsonify({'ok': False, 'error': 'OIB mora imati točno 11 znamenki.'}), 400
    kl = read_klijenti()
    idx = next((i for i,c in enumerate(kl) if c.get('name') == name), None)
    if idx is None:
        return jsonify({'ok': False, 'error': 'Klijent nije pronađen.'}), 404
    # Ako mijenja naziv, provjeri da nema kolizije
    new_name = data['name']
    if new_name != name and any(c.get('name') == new_name for c in kl):
        return jsonify({'ok': False, 'error': 'Klijent s tim nazivom već postoji.'}), 409
    # Ažuriraj polja (ne diramo created_at)
    kl[idx]['name'] = new_name
    kl[idx]['oib'] = data['oib']
    kl[idx]['headquarters'] = data['headquarters']
    kl[idx]['shipping'] = data.get('shipping','')
    kl[idx]['email'] = data['email']
    kl[idx]['phone'] = data['phone']
    try:
        write_klijenti(kl)
    except Exception as e:
        return jsonify({'ok': False, 'error': 'Greška pri spremanju podataka.'}), 500
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True}), 200
# -------------------- NALOZI + dodatne stranice --------------------
@app.route('/kreiraj-nalog/<name>')
@login_required
def kreiraj_nalog(name):
    klijenti = read_klijenti()
    k = next((c for c in klijenti if c['name'] == name), None)
    if not k: abort(404)
    return render_template('kreiraj_nalog.html',
                           title=f'Kreiraj nalog — {k["name"]}',
                           username=current_user.username,
                           klijent=k)

@app.route('/instalacija/<name>', methods=['GET','POST'])
@login_required
def instalacija(name):
    klijenti = read_klijenti()
    k = next((c for c in klijenti if c['name'] == name), None)
    if not k: abort(404)

    # dozvoljeno više instalacija za istog klijenta
    if request.method == 'POST':
        # === Validacija forme za INSTALACIJU ===
        form = request.form

        # 1) Usluge: mora biti označena barem jedna (Billy S, Billy M, Billy Pay)
        allowed_usluge = {'Billy S', 'Billy M', 'Billy Pay'}
        usluge_sel = [u for u in form.getlist('usluge') if u in allowed_usluge]
        if not usluge_sel:
            flash("Odaberi barem jednu uslugu (Billy S, Billy M ili Billy Pay).", "danger")
            return render_template('instalacija.html',
                                       title=f'Instalacija — {k["name"]}',
                                       username=current_user.username,
                                       klijent=k)

        # 2) Način isporuke: mora biti odabran jedan (osnovni); 'Hitna instalacija' može biti dodatna
        nacin = form.get('nacin', '').strip()
        if not nacin:
            flash("Odaberi 'Način isporuke'.", "danger")
            return render_template('instalacija.html',
                                       title=f'Instalacija — {k["name"]}',
                                       username=current_user.username,
                                       klijent=k)

        # 3) Podopcije dostave: mora biti odabrana točno jedna
        podopcija = form.get('podopcija', '').strip()
        if not podopcija:
            flash("Odaberi točno jednu podopciju dostave.", "danger")
            return render_template('instalacija.html',
                                       title=f'Instalacija — {k["name"]}',
                                       username=current_user.username,
                                       klijent=k)

        # 4) Uređaji: ako je količina > 0, za SVAKI uređaj mora biti odabrana odgovarajuća vrsta usluge
        try:
            device_names = read_nazivi_uredjaja()
        except Exception:
            device_names = []
        errors = []
        for dname in device_names:
            try:
                qty = int(form.get(dname, '0') or '0')
            except Exception:
                qty = 0
            if qty > 0:
                try:
                    dslug = slugify(str(dname))
                except Exception:
                    dslug = re.sub(r'[^a-zA-Z0-9_-]+', '-', str(dname)).strip('-').lower()
                # Pronadi odabire vrste usluge za ovaj uređaj neovisno o sufiksu ('[]') i stilu slug-a

                import unicodedata as _ud, re as _re

                def _norm_slug(_s):

                    _s = str(_s or '')

                    _s = _ud.normalize('NFKD', _s)

                    _s = ''.join(ch for ch in _s if not _ud.combining(ch))

                    _s = _s.lower()

                    return _re.sub(r'[^a-z0-9]+', '', _s)

                base_slug = str(dslug)

                selected_types = []

                # 1) Pretraži sve ključeve koji počinju s 'vrsta_usluge__' i usporedi normalizirani sufiks

                for _k in request.form.keys():

                    if not _k.startswith('vrsta_usluge__'):

                        continue

                    _suffix = _k[len('vrsta_usluge__'):]

                    if _suffix.endswith('[]'):

                        _suffix = _suffix[:-2]

                    if _norm_slug(_suffix) == _norm_slug(base_slug):

                        _vals = form.getlist(_k)

                        selected_types = [s.strip() for s in _vals if s and s.strip()]

                        break

                # 2) Fallback: pokušaj par očitih varijanti

                if not selected_types:

                    _cands = [

                        f"vrsta_usluge__{base_slug}", f"vrsta_usluge__{base_slug}[]",

                        f"vrsta_usluge__{base_slug.replace('-', '_')}", f"vrsta_usluge__{base_slug.replace('-', '_')}[]",

                        f"vrsta_usluge__{base_slug.replace('-', '')}", f"vrsta_usluge__{base_slug.replace('-', '')}[]",

                    ]

                    for _ck in _cands:

                        _vals = form.getlist(_ck)

                        if _vals:

                            selected_types = [s.strip() for s in _vals if s and s.strip()]

                            break

                if len(selected_types) != qty:

                    errors.append(f"Za '{dname}' odabrano {len(selected_types)} vrsta usluge, a količina je {qty}.")

        if errors:
            flash(" ".join(errors) + " Za svaki uređaj s količinom > 0 mora biti odabrana jedna 'Vrsta usluge' po uređaju.", "danger")
            return render_template('instalacija.html',
                                   title=f'Instalacija — {k["name"]}',
                                   username=current_user.username,
                                   klijent=k)

        # Ako je sve u redu, generiraj nalog
        # 5) Zahtjev: obavezno priložiti sliku (JPG/PNG)
        zahtjev_file = request.files.get('zahtjev_img')
        if not zahtjev_file or not zahtjev_file.filename:
            flash("Obavezno je priložiti sliku 'Zahtjev' (JPG/PNG).", 'danger')
            return render_template('instalacija.html',
                                   title=f'Instalacija — {k["name"]}',
                                   username=current_user.username,
                                   klijent=k)
        ext = os.path.splitext(zahtjev_file.filename)[1].lower()
        if ext not in ('.jpg', '.jpeg', '.png', '.pdf'):
            flash('Dozvoljeni formati Zahtjeva: JPG/PNG/PDF.', 'danger')
            return render_template('instalacija.html',
                                   title=f'Instalacija — {k["name"]}',
                                   username=current_user.username,
                                   klijent=k)

        return generate_nalog_docx(k, request.form)
    return render_template('instalacija.html',
                           title=f'Instalacija — {k["name"]}',
                           username=current_user.username,
                           klijent=k)

@app.route('/deinstalacija/<name>', methods=['GET','POST'])
@login_required
def deinstalacija(name):
    klijenti = read_klijenti()
    k = next((c for c in klijenti if c['name'] == name), None)
    if not k: abort(404)

    ## KLIJENT-OPREMA-IZ-KLIJENTIJSON
    # Iz polja u klijenti.json: "sn_uredjaja", "model_uredjaja", "sn_SIM"
    dev_rows = []
    sims_rows = []
    try:
        sn_str = str(k.get('sn_uredjaja') or '').strip()
        model_str = str(k.get('model_uredjaja') or '').strip()
        sn_list = [s.strip() for s in sn_str.split(',') if s.strip()] if sn_str else []
        model_list = [s.strip() for s in model_str.split(',') if s.strip()] if model_str else []
        # Pokušaj parirati po indeksu, a ako nema dovoljno modela – popuni prazno
        max_len = max(len(sn_list), len(model_list))
        for i in range(max_len):
            sn = sn_list[i] if i < len(sn_list) else ''
            md = model_list[i] if i < len(model_list) else ''
            if sn or md:
                dev_rows.append({'serijski': sn, 'model': md})
    except Exception:
        pass
    try:
        sim_str = str(k.get('sn_SIM') or '').strip()
        sims_rows = [{'serijski': s} for s in [x.strip() for x in sim_str.split(',') if x.strip()]] if sim_str else []
    except Exception:
        sims_rows = []
    # --- FALLBACK_AKTIVNI: ako u klijenti.json nema popisa, pokušaj iz aktivnih evidencija ---
    try:
        client_name = k.get('name') if isinstance(k, dict) else str(name)
        client_oib = k.get('oib') if isinstance(k, dict) else None
        if not dev_rows:
            try:
                akt = read_aktivni_uredjaji()
            except Exception:
                akt = []
            for d in akt:
                if (str(d.get('client','')) == str(client_name)) or (client_oib and str(d.get('oib','')) == str(client_oib)):
                    if d.get('active', True):
                        dev_rows.append({'serijski': str(d.get('serijski','')), 'model': str(d.get('model',''))})
        if not sims_rows:
            try:
                akt_s = read_aktivni_sim()
            except Exception:
                akt_s = []
            for s in akt_s:
                if (str(s.get('client','')) == str(client_name)) or (client_oib and str(s.get('oib','')) == str(client_oib)):
                    if s.get('active', True):
                        sims_rows.append({'serijski': str(s.get('serijski',''))})
    except Exception:
        pass
    # --- EXCLUDE SNs already requested for deinstalation (open orders) ---
    try:
        _open_deinst = read_otvorene_deinstalacije()
    except Exception:
        _open_deinst = []
    _blocked_dev = set()
    _blocked_sim = set()
    try:
        _client_name_cmp = str(k.get('name') if isinstance(k, dict) else str(name))
    except Exception:
        _client_name_cmp = str(name)
    for _o in _open_deinst:
        try:
            if str(_o.get('client','')) == _client_name_cmp:
                # collect SNs from open deinst orders for this client
                for _s in str(_o.get('deinstalirani sn uredjaja','')).split(','):
                    _s = _s.strip()
                    if _s:
                        _blocked_dev.add(_s)
                for _s in str(_o.get('deinstalirani sn SIM-a','')).split(','):
                    _s = _s.strip()
                    if _s:
                        _blocked_sim.add(_s)
        except Exception:
            pass
    if _blocked_dev:
        dev_rows = [r for r in dev_rows if str(r.get('serijski')) not in _blocked_dev]
    if _blocked_sim:
        sims_rows = [r for r in sims_rows if str(r.get('serijski')) not in _blocked_sim]


    if request.method == 'POST':
        action = (request.form.get('action') or '').strip().lower()
        selected_devs = request.form.getlist('odabrani_uredjaji') or request.form.getlist('return_device')
        selected_sims = request.form.getlist('odabrani_simovi') or request.form.getlist('return_sim')
        if action == 'kreiraj_deinst_nalog':
            rn_str, out_path = generate_deinstalacija_docx(k, selected_devs, selected_sims, napomena=request.form.get('napomena',''))
            # INIT zahtjev_image_rel (bit će postavljen ako je uploadan)
            zahtjev_image_rel = ''
            try:
                zahtjev_image_rel = ''
                _zf = request.files.get('zahtjev_img')
                if _zf and _zf.filename:
                    _ext = os.path.splitext(_zf.filename)[1].lower()
                    if _ext in ('.jpg','.jpeg','.png','.pdf'):
                        _ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                        _safe_rn = re.sub(r'[^A-Za-z0-9_-]+', '_', str(rn_str))
                        _fname = f'{_safe_rn}_{_ts}{_ext}'
                        _dest_dir = os.path.join(STATIC_DIR, 'zapisnici', 'Zahtjev')
                        os.makedirs(_dest_dir, exist_ok=True)
                        _save_path = os.path.join(_dest_dir, _fname)
                        _zf.save(_save_path)
                        zahtjev_image_rel = os.path.join('zapisnici', 'Zahtjev', _fname).replace('\\','/')
                if app.config.get('MAIL_SERVER') and app.config.get('MAIL_USERNAME'):
                    subject = f"Deinstalacija RN {rn_str} za {k['name']}" if isinstance(k, dict) else f"Deinstalacija RN {rn_str}"
                    body = None
                    for _cand in [
                        os.path.join(BASE_DIR, 'deinstalacija e-mail.txt') if 'BASE_DIR' in globals() else None,
                        os.path.join(STATIC_DIR, 'datoteke', 'deinstalacija e-mail.txt')
                    ]:
                        if _cand and os.path.exists(_cand):
                            try:
                                with open(_cand, 'r', encoding='utf-8') as _f:
                                    body = _f.read()
                                    break
                            except Exception:
                                pass
                    if not body:
                        body = "Pozdrav,\n\nU prilogu šaljem odobren zahtjev za novu uslugu.\n\nSrdačno,"
                    msg = Message(subject=subject, recipients=['webtest806@gmail.com'], body=body)
                    # Privitak 1: RN .docx
                    try:
                        with app.open_resource(out_path) as fp2:
                            _mtype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            msg.attach(os.path.basename(out_path), _mtype, fp2.read())
                    except Exception:
                        pass
                    # Privitak 2: Zahtjev
                    try:
                        if zahtjev_image_rel:
                            _abs_zahtjev = os.path.join(STATIC_DIR, zahtjev_image_rel).replace('\\','/')
                            if os.path.exists(_abs_zahtjev):
                                import mimetypes
                                _mtype = mimetypes.guess_type(_abs_zahtjev)[0] or 'application/octet-stream'
                                with open(_abs_zahtjev, 'rb') as _fpz:
                                    msg.attach(os.path.basename(_abs_zahtjev), _mtype, _fpz.read())
                    except Exception:
                        pass
                    try:
                        mail.send(msg)
                    except Exception as _e:
                        app.logger.warning(f"Slanje e-pošte (deinstalacija) nije uspjelo: {_e}")
            except Exception as e:
                app.logger.warning(f"Greška pri slanju e-maila za deinstalaciju: {e}")
            flash(f'Deinstalacijski nalog {rn_str} kreiran i prebačen u Otvorene deinstalacije.', 'success')

            # --- ZAPIŠI OTVORENI DEINST NALOG u nalozi + otvorene_deinstalacije, s linkom na Zahtjev ---
            try:
                korisnik = str(k.get('name') if isinstance(k, dict) else name)
                # relativna putanja RN datoteke
                try:
                    file_rel = out_path.replace(STATIC_DIR + os.sep, '').replace('\\','/').replace('\\','/')
                except Exception:
                    file_rel = os.path.basename(out_path)
                order = {
                    'zahtjev_image': zahtjev_image_rel,
                    'type': 'deinstalacija',
                    'client': korisnik,
                    'file': file_rel,
                    'created_at': datetime.datetime.now().isoformat(),
                    'rn': rn_str,
                    'status': 'nezadužen',
                    'assigned_to': '-'
                }
                nalozi = read_nalozi()
                nalozi.append(order)
                write_nalozi(nalozi)
                try:
                    od = read_otvorene_deinstalacije()
                except Exception:
                    od = []
                od.append(dict(order))
                write_otvorene_deinstalacije(od)
            except Exception:
                pass


            return redirect(url_for('nalozi'))
        rn_str, out_path = generate_deinstalacija_docx(k, selected_devs, selected_sims, napomena=request.form.get('napomena',''))
        try:
            zahtjev_image_rel = ''
            _zf = request.files.get('zahtjev_img')
            if _zf and _zf.filename:
                _ext = os.path.splitext(_zf.filename)[1].lower()
                if _ext in ('.jpg','.jpeg','.png','.pdf'):
                    _ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                    _safe_rn = re.sub(r'[^A-Za-z0-9_-]+', '_', str(rn_str))
                    _fname = f'{_safe_rn}_{_ts}{_ext}'
                    _dest_dir = os.path.join(STATIC_DIR, 'zapisnici', 'Zahtjev')
                    os.makedirs(_dest_dir, exist_ok=True)
                    _save_path = os.path.join(_dest_dir, _fname)
                    _zf.save(_save_path)
                    zahtjev_image_rel = os.path.join('zapisnici', 'Zahtjev', _fname).replace('\\','/')
            if app.config.get('MAIL_SERVER') and app.config.get('MAIL_USERNAME'):
                subject = f"Deinstalacija RN {rn_str} za {k['name']}" if isinstance(k, dict) else f"Deinstalacija RN {rn_str}"
                body = None
                for _cand in [
                    os.path.join(BASE_DIR, 'deinstalacija e-mail.txt') if 'BASE_DIR' in globals() else None,
                    os.path.join(STATIC_DIR, 'datoteke', 'deinstalacija e-mail.txt')
                ]:
                    if _cand and os.path.exists(_cand):
                        try:
                            with open(_cand, 'r', encoding='utf-8') as _f:
                                body = _f.read()
                                break
                        except Exception:
                            pass
                if not body:
                    body = "Pozdrav,\n\nU prilogu šaljem odobren zahtjev za novu uslugu.\n\nSrdačno,"
                msg = Message(subject=subject, recipients=['webtest806@gmail.com'], body=body)
                # Privitak 1: RN .docx
                try:
                    with app.open_resource(out_path) as fp2:
                        _mtype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        msg.attach(os.path.basename(out_path), _mtype, fp2.read())
                except Exception:
                    pass
                # Privitak 2: Zahtjev
                try:
                    if zahtjev_image_rel:
                        _abs_zahtjev = os.path.join(STATIC_DIR, zahtjev_image_rel).replace('\\','/')
                        if os.path.exists(_abs_zahtjev):
                            import mimetypes
                            _mtype = mimetypes.guess_type(_abs_zahtjev)[0] or 'application/octet-stream'
                            with open(_abs_zahtjev, 'rb') as _fpz:
                                msg.attach(os.path.basename(_abs_zahtjev), _mtype, _fpz.read())
                except Exception:
                    pass
                try:
                    mail.send(msg)
                except Exception as _e:
                    app.logger.warning(f"Slanje e-pošte (deinstalacija) nije uspjelo: {_e}")
        except Exception as e:
            app.logger.warning(f"Greška pri slanju e-maila za deinstalaciju: {e}")
        flash(f'Deinstalacija {rn_str} kreirana.', 'success')
        return redirect(url_for('nalozi'))



    return render_template('deinstalacija.html',
                           title=f'Deinstalacija — {k["name"]}',
                           username=current_user.username,
                           klijent=k,
                           uredjaji_klijent=dev_rows, uredjaji_klijenta=dev_rows, aktivni_uredjaji=dev_rows,
                           sims_klijent=sims_rows, sims_klijenta=sims_rows, aktivni_simovi=sims_rows)

@app.route('/servis/<name>', methods=['GET','POST'])
@login_required
def servis(name):
    klijenti = read_klijenti()
    k = next((c for c in klijenti if c['name'] == name), None)
    if not k: abort(404)

    # Pripremi popise kao i na /deinstalacija/
    dev_rows = []
    sims_rows = []
    try:
        sn_str = str(k.get('sn_uredjaja') or '').strip()
        model_str = str(k.get('model_uredjaja') or '').strip()
        sn_list = [s.strip() for s in sn_str.split(',') if s.strip()] if sn_str else []
        model_list = [s.strip() for s in model_str.split(',') if s.strip()] if model_str else []
        max_len = max(len(sn_list), len(model_list))
        for i in range(max_len):
            sn = sn_list[i] if i < len(sn_list) else ''
            md = model_list[i] if i < len(model_list) else ''
            if sn or md:
                dev_rows.append({'serijski': sn, 'model': md})
    except Exception:
        pass
    try:
        sim_str = str(k.get('sn_SIM') or '').strip()
        sims_rows = [{'serijski': s} for s in [x.strip() for x in sim_str.split(',') if x.strip()]] if sim_str else []
    except Exception:
        sims_rows = []

    # FALLBACK na aktivne evidencije ako u klijenti.json nema popisa
    try:
        client_name = k.get('name') if isinstance(k, dict) else str(name)
        client_oib = k.get('oib') if isinstance(k, dict) else None
        if not dev_rows:
            try:
                akt = read_aktivni_uredjaji()
            except Exception:
                akt = []
            for d in akt:
                if (str(d.get('client','')) == str(client_name)) or (client_oib and str(d.get('oib','')) == str(client_oib)):
                    if d.get('active', True):
                        dev_rows.append({'serijski': str(d.get('serijski','')), 'model': str(d.get('model',''))})
        if not sims_rows:
            try:
                akt_s = read_aktivni_sim()
            except Exception:
                akt_s = []
            for s in akt_s:
                if (str(s.get('client','')) == str(client_name)) or (client_oib and str(s.get('oib','')) == str(client_oib)):
                    if s.get('active', True):
                        sims_rows.append({'serijski': str(s.get('serijski',''))})
    except Exception:
        pass

    if request.method == 'POST':
        selected_devs = request.form.getlist('odabrani_uredjaji') or request.form.getlist('return_device')
        selected_sims = request.form.getlist('odabrani_simovi') or request.form.getlist('return_sim')
        rn_str, out_path = generate_servis_docx(k, selected_devs, selected_sims, napomena=request.form.get('napomena',''))
        flash(f'Servisni nalog {rn_str} kreiran.', 'success')
        return redirect(url_for('nalozi'))

    return render_template('servis.html',
                           title=f'Servis — {k["name"]}',
                           username=current_user.username,
                           klijent=k,
                           uredjaji_klijent=dev_rows, uredjaji_klijenta=dev_rows, aktivni_uredjaji=dev_rows,
                           sims_klijent=sims_rows, sims_klijenta=sims_rows, aktivni_simovi=sims_rows)

# -------------------- GENERIRANJE NALOGA --------------------
def generate_nalog_docx(klijent, formdata):
    template_path = os.path.join(STATIC_DIR, 'datoteke', 'rn_template.docx')
    if not os.path.exists(template_path):
        abort(500, description="Nedostaje rn_template.docx u static/datoteke")

    doc = Document(template_path)

    # === Način isporuke i Podopcije (iz nacin.isporuke.json) ===
    try:
        _nacin_cfg = read_nacin_isporuke() or []
    except Exception:
        _nacin_cfg = []
    nacin_value = (formdata.get('nacin') or '').strip()
    podopcija_value = (formdata.get('podopcija') or '').strip()
    try:
        _hitna_labels = read_hitna_instalacija() or []
    except Exception:
        _hitna_labels = []
    emergency_label = (_hitna_labels[0] if _hitna_labels else 'Hitna instalacija (instalacija sljedeći radni dan)')
    # labela = vrijednost (JSON lista je string lista); ipak zadrži helper za buduće strukture
    def _label_for(items, value):
        vv = str(value or '').strip()
        for it in (items or []):
            # podrži i dict i plain string
            lab = it.get('label') if isinstance(it, dict) else str(it)
            val = it.get('value') if isinstance(it, dict) else str(it)
            if str(val).strip() == vv or str(lab).strip() == vv:
                return str(lab or val)
        return vv
    # Od Nacin isporuke (glavna opcija) i Podopciju dostave
    nacin_label = _label_for(_nacin_cfg, nacin_value)
    # Podopcije su fiksno: "Kurirskom službom" / "Kod Pružatelja usluga" (vidi JSON)
    podopc_label = _label_for(_nacin_cfg, podopcija_value)
    # Hitna: zaseban checkbox u formi
    hitna_checked = str(formdata.get('hitna') or '').lower() in ('1','on','true','da')

    # selections
    usluge = formdata.getlist('usluge')
    ADD_KEYS = ["SIM kartica","Kasa","SumUp čitač","Termo traka (58mm)","Termo traka (80mm)","Ladica za novac","Stylus olovka"]
    uredjaji_qty = {k:int(v) for k,v in formdata.items() if k not in ['usluge','nacin','podopcija'] and v.isdigit() and int(v)>=0 and k not in ADD_KEYS}
    dodatne_qty = {}
    ADD_KEYS = ["SIM kartica","Kasa","SumUp čitač","Termo traka (58mm)","Termo traka (80mm)","Ladica za novac","Stylus olovka"]
    for k,v in formdata.items():
        if k in ADD_KEYS and v.isdigit():
            dodatne_qty[k] = int(v)

    nacin = formdata.get('nacin','')
    podopc = formdata.getlist('podopcija')

    korisnik = klijent['name']
    adresa = klijent.get('shipping') or klijent.get('headquarters') or ''
    year = datetime.datetime.now().year

    # RN broj: 0001/YYYY (samo za instalacije; reset 1.1. nove godine)
    nalozi = read_nalozi()
    rn_num = next_rn_for_year(nalozi, year)
    rn_str = f"{rn_num:04d}/{year}"

    mapping = {
        '{{RN}}': rn_str,
        '{{KORISNIK}}': korisnik,


        '{{ADRESA_ISPORUKE}}': adresa,
        '{{OIB}}': klijent.get('oib',''),
        '{{KONTAKT}}': klijent.get('phone',''),
        '{{RN}}': rn_str

    }
# === PLACEHOLDERI: isporuka/podopcija/hitna ===
    # Zamjena u DOCX + e-mail
    mapping['{{Isporuka}}'] = nacin_label
    mapping['*  {{Isporuka}}'] = (f"☑  {nacin_label}" if nacin_label else '')
    mapping['{{podopcija}}'] = podopc_label
    mapping['*  {{podopcija}}'] = (f"☑  {podopc_label}" if podopc_label else '')
    mapping['{{hitna}}'] = (f"☑  {emergency_label}" if hitna_checked else '')
    mapping['*  {{hitna}}'] = (f"☑  {emergency_label}" if hitna_checked else '')

    has_sim = (dodatne_qty.get('SIM kartica',0) > 0) or ('SIM kartica' in usluge)
    has_additional = any(q>0 for k,q in dodatne_qty.items() if k!='SIM kartica')
    mapping['{{SIM}}'] = 'X' if has_sim else ''
    mapping['{{DODATNA_OPREMA}}'] = 'X' if has_additional else ''

    has_device_one = any(q==1 for q in uredjaji_qty.values())


    # === Dinamički prikaz "Usluge:" iz usluge.JSON i pay.opcije.JSON ===
    try:
        usluge_cfg = read_usluge_config()
    except Exception:
        usluge_cfg = {}
    try:
        pay_opcije_cfg = read_pay_opcije()
    except Exception:
        pay_opcije_cfg = []

    def _label_for(service_name):
        # podrži dict {"Billy S":"Billy S",...} ili listu
        if isinstance(usluge_cfg, dict):
            return str(usluge_cfg.get(service_name, service_name))
        if isinstance(usluge_cfg, list):
            # ako lista sadrži dict-ove s 'name'/'label', pokušaj naći po imenu
            for it in usluge_cfg:
                if isinstance(it, str) and it.strip().lower() == service_name.lower():
                    return it
                if isinstance(it, dict):
                    _nm = str(it.get('name') or it.get('label') or it.get('title') or '').strip()
                    if _nm and _nm.lower() == service_name.lower():
                        return _nm
            return service_name
        return service_name

    def _first_option_from(cfg):
        if isinstance(cfg, list) and cfg:
            it = cfg[0]
            if isinstance(it, str):
                return it
            if isinstance(it, dict):
                return str(it.get('label') or it.get('name') or it.get('title') or next((v for v in it.values() if isinstance(v,str)), ''))
        if isinstance(cfg, dict):
            # uzmi prvu vrijednost
            try:
                k = next(iter(cfg.values()))
                return str(k)
            except Exception:
                return ''
        return ''

    # dohvat odabrane opcije Billy Pay iz forme (robustan: traži ključ koji sadrži 'pay' i 'opc')
    sel_pay_opcija = None
    for _k in formdata.keys():
        _lk = _k.lower()
        if 'pay' in _lk and 'opc' in _lk:
            sel_pay_opcija = formdata.get(_k) or formdata.getlist(_k)[0] if formdata.getlist(_k) else None
            break
    if not sel_pay_opcija:
        # probaj nekoliko poznatih naziva
        for _cand in ['pay_opcija','billy_pay_opcija','opcija_billy_pay','payPodopcija','pay_podopcija']:
            if _cand in formdata:
                sel_pay_opcija = formdata.get(_cand)
                break

    # izgradi listu usluga prema pravilima
    sel = [u for u in (usluge or []) if u in {'Billy S','Billy M','Billy Pay'}]
    parts = []
    if 'Billy S' in sel:
        parts.append(_label_for('Billy S'))
    if 'Billy M' in sel:
        parts.append(_label_for('Billy M'))
        if 'Billy Pay' in sel:
            parts.append(_label_for('Billy Pay'))
    else:
        if 'Billy Pay' in sel:
            parts.append(_label_for('Billy Pay'))

    # pripremi tekst s odvojnikom za Billy Pay opciju (samo ako je Billy Pay selektiran)
    text_services = " / ".join([p for p in parts if p])
    pay_suffix = ""
    if 'Billy Pay' in sel:
        pay_label = str(sel_pay_opcija or "").strip()
        if not pay_label:
            pay_label = _first_option_from(pay_opcije_cfg)
        if pay_label:
            pay_suffix = f" - {pay_label}"

    full_services_text = f"☑  {text_services}{pay_suffix}".strip()

    # zamijeni originalni * {{Usluga1}}/{{Usluga2}}/{{Usluga3}} - {{Opcija Billy Pay}} s novim tekstom
    def _replace_services_text_in_doc(_doc):
        targets = [
            "*  {{Usluga1}}/{{Usluga2}}/{{Usluga3}} - {{Opcija Billy Pay}}",
            "* {{Usluga1}}/{{Usluga2}}/{{Usluga3}} - {{Opcija Billy Pay}}",
            "•  {{Usluga1}}/{{Usluga2}}/{{Usluga3}} - {{Opcija Billy Pay}}",
            "{{Usluga1}}/{{Usluga2}}/{{Usluga3}} - {{Opcija Billy Pay}}",
            "*  {{Usluga1}}/{{Usluga2}}/{{Usluga3}} – {{Opcija Billy Pay}}",
            "* {{Usluga1}}/{{Usluga2}}/{{Usluga3}} – {{Opcija Billy Pay}}",
        ]
        # paragrafi
        for p in _doc.paragraphs:
            for t in targets:
                if t in p.text:
                    p.text = p.text.replace(t, full_services_text)
        # tablice
        for tbl in _doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for t in targets:
                        if t in cell.text:
                            cell.text = cell.text.replace(t, full_services_text)
    _replace_services_text_in_doc(doc)
    replace_text_in_doc(doc, mapping)
    # --- Post obrada kućica ---
    try:
        for p in doc.paragraphs:
            txt = p.text
            if txt.lstrip().startswith('*'):
                prefix_idx = txt.find('*')
                lead = txt[:prefix_idx]
                rest = txt[prefix_idx+1:].lstrip()
                if rest:
                    p.text = lead + '☑  ' + rest
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    ctext = cell.text
                    lines = []
                    for line in ctext.splitlines():
                        if line.lstrip().startswith('*'):
                            idx = line.find('*')
                            lead = line[:idx]
                            rest = line[idx+1:].lstrip()
                            if rest:
                                lines.append(lead + '☑  ' + rest)
                            else:
                                lines.append('')
                        else:
                            lines.append(line)
                    cell.text = '\n'.join(lines)
    except Exception:
        pass

    table = find_table_with_header(doc, ['Naziv Opreme','TID','Serijski broj','OTP'])
    # Pripremi redove uređaja s vrstom usluge (Kupnja/Najam) iz forme
    device_rows = []
    try:
        # helper za normalizaciju slug-a kao u validaciji
        import unicodedata as _ud, re as _re
        def _norm_slug(_s):
            _s = str(_s or '')
            _s = _ud.normalize('NFKD', _s)
            _s = ''.join(ch for ch in _s if not _ud.combining(ch))
            _s = _s.lower()
            return _re.sub(r'[^a-z0-9]+', '', _s)
    except Exception:
        def _norm_slug(s): return str(s or '').lower().replace(' ', '')
    # Mapiraj uređaj -> odabrane vrste usluge (lista)
    selected_types_map = {}
    for dev_name, qty in uredjaji_qty.items():
        if not qty or qty <= 0:
            continue
        # potraži polja koja počinju s 'vrsta_usluge__'
        vals = []
        base = str(dev_name)
        base_slug = re.sub(r'[^a-zA-Z0-9_-]+', '-', base).strip('-').lower()
        for key, v in formdata.items(multi=True):
            if not key.startswith('vrsta_usluge__'):
                continue
            suf = key.split('vrsta_usluge__', 1)[1]
            if _norm_slug(suf) == _norm_slug(base_slug):
                # moguća je lista
                if isinstance(v, (list, tuple)):
                    vals.extend([str(x).strip() for x in v if str(x).strip()])
                else:
                    vals.append(str(v).strip())
        # fallback uobičajene varijante imena polja
        if not vals:
            for cand in [f"vrsta_usluge__{base_slug}", f"vrsta_usluge__{base_slug}[]",
                         f"vrsta_usluge__{base_slug.replace('-', '_')}", f"vrsta_usluge__{base_slug.replace('-', '_')}[]",
                         f"vrsta_usluge__{base_slug.replace('-', '')}", f"vrsta_usluge__{base_slug.replace('-', '')}[]"]:
                got = formdata.getlist(cand)
                if got:
                    vals = [str(x).strip() for x in got if str(x).strip()]
                    break
        # osiguraj duljinu == qty (ako manje, popuni praznim)
        while len(vals) < qty:
            vals.append('')
        # ako je više nego qty, skrati
        if len(vals) > qty:
            vals = vals[:qty]
        selected_types_map[dev_name] = vals
        for t in vals:
            device_rows.append({'Naziv Opreme': f"{dev_name} – {t}".strip(' –'),
                                'TID': '', 'Serijski broj': '', 'OTP': ''})
    # dodaj SIM redove kao prije
    sim_qty = dodatne_qty.get('SIM kartica', 0)
    for _ in range(sim_qty):
        device_rows.append({'Naziv Opreme': 'SIM kartica', 'TID': '', 'Serijski broj': '', 'OTP': ''})
    # dodatne usluge iz forme (bez SIM) — s prefiksom količine za 2+
    for _name in ["Kasa","SumUp čitač","Termo traka (58mm)","Termo traka (80mm)","Ladica za novac","Stylus olovka"]:
        try:
            _n = int((dodatne_qty or {}).get(_name, 0) or 0)
        except Exception:
            _n = 0
        if _n > 0:
            _label = f"{_n}x {_name}" if _n >= 2 else _name
            device_rows.append({'Naziv Opreme': _label, 'TID': '', 'Serijski broj': '', 'OTP': ''})
    # dodatne usluge (bez SIM) — prikaz s prefiksom količine za 2+
    for _name in ["Kasa","SumUp čitač","Termo traka (58mm)","Termo traka (80mm)","Ladica za novac","Stylus olovka"]:
        _n = int((dodatne_qty or {}).get(_name, 0) or 0)
        if _n > 0:
            _label = f"{_n}x {_name}" if _n >= 2 else _name
            device_rows.append({'Naziv Opreme': _label, 'TID': '', 'Serijski broj': '', 'OTP': ''})

    # deduplikacija redova (spriječi dupla umetanja istih naziva)
    if device_rows:
        _seen=set()
        _unique=[]
        for _it in device_rows:
            # Ne dedupliciraj SIM kartice — trebaju višestruki redovi za količinu
            if str(_it.get('Naziv Opreme','')) == 'SIM kartica':
                _unique.append(_it)
                continue
            _key=(
                str(_it.get('Naziv Opreme','')), str(_it.get('TID','')),
                str(_it.get('Serijski broj','')), str(_it.get('OTP',''))
            )
            if _key in _seen:
                continue
            _seen.add(_key)
            _unique.append(_it)
        device_rows=_unique

    # Upis u tablicu: prvi red (Rb.=1) već postoji u predlošku -> popuni ga,
    # a za preostale uređaje kreiraj nove redove (Rb. 2, 3, ...)
    if table and device_rows:
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        # pomoćna: dohvat indeksa kolone po nazivu (dopušta djelomično poklapanje)
        def _col_idx(name):
            name_l = name.lower()
            for i, c in enumerate(header_cells):
                if name_l in c.lower():
                    return i
            return None
        idx_rb  = _col_idx('Rb')
        idx_no  = _col_idx('Naziv')
        idx_tid = _col_idx('TID')
        idx_sn  = _col_idx('Serijski')
        idx_otp = _col_idx('OTP')
        # popuni prvi postojeći red (index 1)
        first = device_rows[0]
        try:
            cells = table.rows[1].cells
            if idx_rb is not None: cells[idx_rb].text  = '1'
            if idx_no is not None: cells[idx_no].text  = first['Naziv Opreme']
            if idx_tid is not None: cells[idx_tid].text = first['TID']
            if idx_sn is not None: cells[idx_sn].text  = first['Serijski broj']
            if idx_otp is not None: cells[idx_otp].text = first['OTP']
        except Exception:
            pass
        # ostali redovi
        # Umjesto table.add_row() (koji zna izgubiti margine/okvire), kloniramo prvi data-red (index 1) i dodajemo ga na kraj.
        from copy import deepcopy
        def _append_cloned_row(_table, _clone_idx=1):
            try:
                tmpl_tr = _table.rows[_clone_idx]._tr
                new_tr = deepcopy(tmpl_tr)
                _table._tbl.append(new_tr)
                return _table.rows[-1]
            except Exception:
                # Fallback ako nešto pođe po zlu
                return _table.add_row()

        for i, item in enumerate(device_rows[1:], start=2):
            _row_obj = _append_cloned_row(table, _clone_idx=1)
            row = _row_obj.cells if hasattr(_row_obj, 'cells') else _row_obj
            if idx_rb is not None and idx_rb < len(row): row[idx_rb].text = str(i)
            if idx_no is not None and idx_no < len(row): row[idx_no].text = item['Naziv Opreme']
            if idx_tid is not None and idx_tid < len(row): row[idx_tid].text = item['TID']
            if idx_sn is not None and idx_sn < len(row): row[idx_sn].text = item['Serijski broj']
            if idx_otp is not None and idx_otp < len(row): row[idx_otp].text = item['OTP']
# Save to static/nalozi/<client-slug>/INSTALL RN 0001-YYYY Klijent,Adresa.docx
    client_slug = slugify(korisnik)
    out_dir = os.path.join(STATIC_DIR, 'nalozi', client_slug)
    os.makedirs(out_dir, exist_ok=True)
    filename = f"INSTALL RN {rn_str.replace('/', '-') } {korisnik},{adresa}.docx"
    out_path = os.path.join(out_dir, filename)
    process_hitna_in_doc(doc, hitna_checked, emergency_label)

    doc.save(out_path)

    # Save record (status: nezadužen | zadužen | zaključen)
    # --- ZAHTJEV: spremi sliku u static/zapisnici/Zahtjev ---
    zahtjev_image_rel = ''
    try:
        _zf = request.files.get('zahtjev_img')
        if _zf and _zf.filename:
            _ext = os.path.splitext(_zf.filename)[1].lower()
            if _ext in ('.jpg','.jpeg','.png','.pdf'):
                _ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                _safe_rn = re.sub(r'[^A-Za-z0-9_-]+', '_', str(rn_str))
                _fname = f'{_safe_rn}_{_ts}{_ext}'
                _dest_dir = os.path.join(STATIC_DIR, 'zapisnici', 'Zahtjev')
                os.makedirs(_dest_dir, exist_ok=True)
                _dest = os.path.join(_dest_dir, _fname)
                _zf.save(_dest)
                zahtjev_image_rel = os.path.join('zapisnici', 'Zahtjev', _fname).replace('\\','/')
    except Exception:
        zahtjev_image_rel = ''




    order = {
        'zahtjev_image': zahtjev_image_rel,
        'type': 'instalacija',
        'client': korisnik,
        'file': f"nalozi/{client_slug}/{filename}",
        'created_at': datetime.datetime.now().isoformat(),
        'nacin': nacin,
        'podopc': podopc,
        'rn': rn_str,
        'status': 'nezadužen',
        'assigned_to': '-'  # korisničko ime operatera
    }
    nalozi.append(order)
    write_nalozi(nalozi)
    # upiši u odgovarajuće otvorene liste ovisno o tipu
    try:
        t = (order.get('type') or '').lower()
        if t == 'instalacija':
            oi = read_otvorene_instalacije()
            oi.append(dict(order))
            write_otvorene_instalacije(oi)
        elif t == 'deinstalacija':
            od = read_otvorene_deinstalacije()
            od.append(dict(order))
            write_otvorene_deinstalacije(od)
        elif t == 'servis':
            osv = read_otvoreni_servisi()
            osv.append(dict(order))
            write_otvoreni_servisi(osv)
    except Exception:
        pass

    try:
        if app.config.get('MAIL_SERVER') and app.config.get('MAIL_USERNAME'):
            with app.open_resource(out_path) as fp:
                                # --- SUBJECT per nova shema: "B0001_YYYY; Klijent.; OIB: 00000000000" ---
                try:
                    # Broj slijeda: ukupan broj instalacijskih naloga (slijed se NE resetira po godinama)
                    _all_orders = read_nalozi()
                except Exception:
                    _all_orders = []
                try:
                    _seq_global = sum(1 for _n in _all_orders if str((_n.get('type') or _n.get('Type') or '')).lower() == 'instalacija')
                except Exception:
                    _seq_global = 1
                # Ako je tek kreiran i još nije u _all_orders (edge-case), osiguraj da je barem 1
                if _seq_global <= 0:
                    _seq_global = 1
                _year_now = datetime.datetime.now().year
                # Dohvati OIB klijenta
                try:
                    _kl = read_klijenti()
                except Exception:
                    _kl = []
                try:
                    _oib = next((str(c.get('oib','')) for c in _kl if str(c.get('name','')) == str(korisnik)), '')
                except Exception:
                    _oib = ''
                _subject_billy = f"B{_seq_global:04d}_{_year_now}; {korisnik}.; OIB: {_oib}"
                msg = Message(subject=_subject_billy, recipients=['webtest806@gmail.com'])

                # Sastavi tijelo e-maila prema predlošku
                # "Pozdrav,\n\n U prilogu šaljem odobren zahtjev za novu uslugu:" + stavke + "Možete nastaviti.\n\nSrdačno,"
                try:
                    # 1) Usluge (checkbox lista)
                    _usluge_list = list(usluge) if usluge else []

                    # 2) Uređaji + vrste usluge po uređaju
                    _stavke = []
                    try:
                        from slugify import slugify as _slugify
                    except Exception:
                        def _slugify(s): return re.sub(r'[^a-zA-Z0-9_-]+','-',str(s)).strip('-').lower()

                    import unicodedata as _ud, re as _re
                    def _norm_slug(_s):
                        _s = str(_s or '')
                        _s = _ud.normalize('NFKD', _s)
                        _s = ''.join(ch for ch in _s if not _ud.combining(ch))
                        _s = _s.lower()
                        return _re.sub(r'[^a-z0-9]+', '', _s)

                    # uredjaji_qty: npr. {'Billy Easy': 2, ...}
                    for _dname, _qty in (uredjaji_qty or {}).items():
                        if int(_qty or 0) <= 0:
                            continue
                        _dslug = _slugify(str(_dname))
                        _base = _norm_slug(_dslug)
                        _selected_types = []
                        # pronađi key s 'vrsta_usluge__<slug>' (uz razne varijacije)
                        for _k in formdata.keys():
                            if not str(_k).startswith('vrsta_usluge__'):
                                continue
                            _kslug = str(_k).split('__',1)[1].rstrip('[]')
                            if _norm_slug(_kslug) == _base:
                                _vals = [v for v in formdata.getlist(_k) if v and v.strip()]
                                if _vals:
                                    _selected_types = [s.strip() for s in _vals]
                                    break
                        if _selected_types:
                            # ako su sve iste vrste, agregiraj; inače grupiraj po vrsti
                            _counts = {}
                            for _v in _selected_types:
                                _counts[_v] = _counts.get(_v, 0) + 1
                            for _v, _c in _counts.items():
                                _stavke.append(f"{_c}x {_dname} {_v}")
                        else:
                            _stavke.append(f"{_qty}x {_dname}")

                    # 3) Dodatne stavke (SIM, trake, itd.)
                    for _k, _v in (dodatne_qty or {}).items():
                        try:
                            _n = int(_v or 0)
                        except Exception:
                            _n = 0
                        if _k == 'SIM kartica':
                            # uključi SIM u e-mail (1x bez prefiksa, 2+ s prefiksom)
                            if _n > 0:
                                if _n >= 2:
                                    _stavke.append(f"{_n}x {_k}")
                                else:
                                    _stavke.append(str(_k))
                            continue
                        if _n > 0:
                            if _n >= 2:
                                _stavke.append(f"{_n}x {_k}")
                            else:
                                _stavke.append(str(_k))

                    # E-mail tekst
                    _lines = []
                    _lines.append("Pozdrav,")
                    _lines.append("")
                    _lines.append("U prilogu šaljem odobren zahtjev za novu uslugu:")
                    _lines.append("")
                    # Prvo eventualne specifične usluge (npr. Kratkoročni najam opreme, Parametriziranje)
                    for _u in _usluge_list:
                        _lines.append(str(_u))
                        _lines.append("")
                    # Zatim uređaji + usluge i dodatne stavke
                    for _s in _stavke:
                        _lines.append(str(_s))
                        _lines.append("")
                    # --- Način isporuke (dodano u e-mail, kao u RN predlošku) ---
                    try:
                        _delivery_lines = []
                        # Glavna opcija (nacin_label) i podopcija (podopc_label)
                        if str(nacin_label or '').strip():
                            _delivery_lines.append(f"* {nacin_label}")
                        if str(podopc_label or '').strip() and str(podopc_label).strip().lower() != str(nacin_label).strip().lower():
                            _delivery_lines.append(f"* {podopc_label}")
                        # Hitna instalacija (ako je označeno)
                        if hitna_checked:
                            _delivery_lines.append(str(emergency_label))
                        if _delivery_lines:
                            _lines.append("")
                            for __l in _delivery_lines:
                                _lines.append(__l)
                            _lines.append("")
                    except Exception:
                        # Ne prekidaj e-mail ako dođe do greške
                        pass

                    _lines.append("Možete nastaviti.")
                    _lines.append("")
                    _lines.append("")
                    _lines.append("Srdačno,")
                    msg.body = "\n".join(_lines)
                except Exception:
                    # fallback
                    msg.body = f"U prilogu je RN {rn_str} za klijenta {korisnik}."

                # Privitak 1: RN .docx
                with app.open_resource(out_path) as fp2:
                    msg.attach(filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', fp2.read())

                # Privitak 2: Zahtjev (JPG/PNG/PDF), ako postoji
                try:
                    if zahtjev_image_rel:
                        _abs_zahtjev = os.path.join(STATIC_DIR, zahtjev_image_rel).replace('\\\\','/')
                        if os.path.exists(_abs_zahtjev):
                            _mtype = mimetypes.guess_type(_abs_zahtjev)[0] or 'application/octet-stream'
                            with open(_abs_zahtjev, 'rb') as _fpz:
                                msg.attach(os.path.basename(_abs_zahtjev), _mtype, _fpz.read())
                except Exception:
                    pass

                mail.send(msg)
    except Exception as e:
        app.logger.warning(f"Slanje e-pošte nije uspjelo: {e}")

    # Nakon kreiranja — redirect na profil klijenta
    flash(f"Nalog {rn_str} uspješno kreiran.", "success")
    return redirect(url_for('klijent_profil', name=korisnik))

# -------------------- STATIC BANNERS --------------------
@app.route('/static/banners/<path:filename>')
def banners(filename):
    return send_from_directory(os.path.join(app.root_path, 'static', 'banners'), filename)

@app.route('/zapisnik/delete', methods=['POST'])
@login_required
def delete_zapisnik():
    # RBAC: samo superadmin/admin smiju brisati
    if not has_role('superadmin','admin'):
        flash('Nemate ovlasti za brisanje.', 'danger')
        client = request.form.get('client') or ''
        return redirect(url_for('klijent_profil', name=client))
    rn = str(request.form.get('rn') or '').strip()
    src_rel = str(request.form.get('src') or '').strip()
    client = request.form.get('client') or ''
    # Normalize path and ensure it's under STATIC_DIR
    src_rel = src_rel.replace('\\', '/')
    # In case full path mistakenly includes leading 'static/', strip it
    if src_rel.startswith('static/'):
        src_rel = src_rel[7:]
    abs_path = os.path.abspath(os.path.join(STATIC_DIR, src_rel))
    if not abs_path.startswith(os.path.abspath(STATIC_DIR)):
        flash('Neispravna putanja.', 'danger')
        return redirect(url_for('klijent_profil', name=client))
    # Obrisi datoteku ako postoji
    try:
        if os.path.exists(abs_path):
            os.remove(abs_path)
    except Exception:
        pass
    # Ukloni referencu na sliku iz nalozi zapisa
    try:
        base = read_nalozi()
    except Exception:
        base = []
    changed = False
    for n in base:
        try:
            if str(n.get('rn')) == rn and n.get('zapisnik_image'):
                n['zapisnik_image'] = ''
                changed = True
        except Exception:
            pass
    if changed:
        try:
            write_nalozi(base)
        except Exception:
            pass
    flash('Zapisnik obrisan.', 'success')
    return redirect(url_for('klijent_profil', name=client))





# -------------------- GOTOVI SERVISI (lista zaključenih servisa) --------------------
@app.route('/gotovi-servisi', endpoint='gotovi_servisi')
@login_required
def gotovi_servisi():
    """Lista svih zaključenih servisa iz zakljuceni.nalozi.JSON."""
    try:
        all_closed = read_zakljuceni_nalozi()
    except Exception:
        all_closed = []
    servisi = [n for n in all_closed if (n.get('type','') or '').lower() == 'servis']
    try:
        servisi.sort(key=lambda x: (x.get('closed_at') or x.get('created_at') or ''), reverse=True)
    except Exception:
        pass
    return render_template('gotovi_servisi.html', title="Gotovi servisi", username=current_user.username, servisi=servisi)

# (old __main__ removed)

# Ensure open deinstalacije/servisi JSON files exist
for p, init in [(OTVORENE_DEINSTALACIJE_JSON_PATH, []), (OTVORENI_SERVISI_JSON_PATH, [])]:
    try:
        if not os.path.exists(p):
            _write_json(p, init)
    except Exception:
        pass


# Ensure servisni.nalog.JSON exists
try:
    if not os.path.exists(SERVISNI_NALOG_JSON_PATH):
        _write_json(SERVISNI_NALOG_JSON_PATH, [])
except Exception:
    pass
# Ensure aktivni_sim.JSON exists
try:
    if not os.path.exists(AKTIVNI_SIM_JSON_PATH):
        _write_json(AKTIVNI_SIM_JSON_PATH, [])
except Exception:
    pass
# ==================== NOVO: ZAKLJUČI DEINSTALACIJU & GOTOVE DEINSTALACIJE ====================

@app.route('/zakljuci-deinstalaciju/<path:rn>', methods=['GET','POST'])
@login_required
def zakljuci_deinstalaciju(rn):
    '''
    Gumb "Zaključi deinstalaciju" s /nalozi vodi ovdje.
    GET: prikaži izbor uređaja i SIM-ova iz naloga (polja 'deinstalirani sn uredjaja' i 'deinstalirani sn SIM-a')
    POST: premjesti nalog iz otvorene_deinstalacije.JSON u zakljuceni.nalozi.JSON,
          vrati označene SN uređaja u uredjaji.json i označene SIM-ove u sim.json kako bi ih bilo moguće ponovno zadužiti.
    '''
    try:
        od = read_otvorene_deinstalacije()
    except Exception:
        od = []
    nalog = next((n for n in od if str(n.get('rn')) == str(rn)), None)
    if not nalog:
        abort(404)

    # Popisi SN-ova iz naloga
    dev_sns = [s.strip() for s in str(nalog.get('deinstalirani sn uredjaja','')).split(',') if s.strip()]
    sim_sns = [s.strip() for s in str(nalog.get('deinstalirani sn SIM-a','')).split(',') if s.strip()]

    # Mapiranja (model i provider) iz aktivnih evidencija ako postoje
    model_by_sn = {}
    provider_by_sn = {}
    try:
        for d in read_aktivni_uredjaji():
            model_by_sn[str(d.get('serijski'))] = d.get('model','')
    except Exception:
        pass
    try:
        for s in read_aktivni_sim():
            provider_by_sn[str(s.get('serijski'))] = s.get('provider','')
    except Exception:
        pass

    if request.method == 'POST':
        selected_dev = [s.strip() for s in request.form.getlist('return_device') if s.strip()]
        selected_sim = [s.strip() for s in request.form.getlist('return_sim') if s.strip()]


        # Obavezni upload zapisnika (JPG/PNG) — spremi u static/zapisnici/Deinstalacija
        file = request.files.get('zapisnik_img')
        if not file or not file.filename:
            flash('Obavezno je priložiti sliku zapisnika (JPG/PNG).', 'danger')
            # Rekonstruiraj prikazne redove i vrati formu
            dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in dev_sns]
            sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sim_sns]
            return render_template('zakljuci.deinstalaciju.html',
                                   title='Zaključi deinstalaciju',
                                   username=current_user.username,
                                   nalog=nalog, klijent={'name': nalog.get('client')},
                                   uredjaji_klijent=dev_rows,
                                   sims_klijent=sim_rows)
        name = file.filename
        ext = os.path.splitext(name)[1].lower()
        if ext not in ('.jpg','.jpeg','.png'):
            flash('Dozvoljeni formati slike: JPG/PNG.', 'danger')
            dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in dev_sns]
            sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sim_sns]
            return render_template('zakljuci.deinstalaciju.html',
                                   title='Zaključi deinstalaciju',
                                   username=current_user.username,
                                   nalog=nalog, klijent={'name': nalog.get('client')},
                                   uredjaji_klijent=dev_rows,
                                   sims_klijent=sim_rows)
        ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_rn = re.sub(r'[^A-Za-z0-9_-]+', '_', str(rn))
        fname = f'{safe_rn}_{ts}{ext}'
        _dest_dir = os.path.join(STATIC_DIR, 'zapisnici', 'Deinstalacija')
        os.makedirs(_dest_dir, exist_ok=True)
        dest = os.path.join(_dest_dir, fname)
        file.save(dest)
        _saved_rel = os.path.join('static','zapisnici','Deinstalacija', fname)
        # DEAKTIVIRAJ IZ AKTIVNIH EVIDENCIJA (uređaji i SIM-ovi kod klijenta)
        dev_set = set(selected_dev or [])
        sim_set = set(selected_sim or [])

        try:
            akt = read_aktivni_uredjaji()
        except Exception:
            akt = []
        try:
            changed = False
            now_iso = datetime.datetime.now().isoformat()
            for d in akt:
                try:
                    if str(d.get('serijski')) in dev_set:
                        d['active'] = False
                        d['unassigned_at'] = now_iso
                        changed = True
                except Exception:
                    pass
            if changed:
                write_aktivni_uredjaji(akt)
        except Exception:
            pass

        try:
            akt_s = read_aktivni_sim()
        except Exception:
            akt_s = []
        try:
            changed_s = False
            now_iso = datetime.datetime.now().isoformat()
            for s in akt_s:
                try:
                    if str(s.get('serijski')) in sim_set:
                        s['active'] = False
                        s['unassigned_at'] = datetime.datetime.now().isoformat()
                        changed_s = True
                except Exception:
                    pass
            if changed_s:
                write_aktivni_sim(akt_s)
        except Exception:
            pass

        # (1) Uređaji -> uredjaji.json

        if selected_dev:
            pool = read_uredjaji()
            seen = {str(d.get('serijski')) for d in pool}
            now = datetime.datetime.now().isoformat()
            for sn in selected_dev:
                if sn not in seen:
                    pool.append({'model': model_by_sn.get(sn,''), 'serijski': sn, 'created_at': now})
                    seen.add(sn)
            write_uredjaji(pool)

        # (2) SIM -> sim.json
        if selected_sim:
            pool_sim = read_sim()
            seen_sim = {str(d.get('serijski')) for d in pool_sim}
            now = datetime.datetime.now().isoformat()
            for sn in selected_sim:
                if sn not in seen_sim:
                    pool_sim.append({'provider': provider_by_sn.get(sn,''), 'serijski': sn, 'created_at': now})
                    seen_sim.add(sn)
            write_sim(pool_sim)

        # (3) Premjesti nalog u zakljuceni.nalozi.JSON
        zk = read_zakljuceni_nalozi()
        closed = dict(nalog)
        # Ne spremati Zahtjev u zatvoreni zapis
        try:
            if 'zahtjev_image' in closed:
                closed['zahtjev_image'] = ''
        except Exception:
            pass
        closed['status'] = 'zaključen'
        closed['closed_at'] = datetime.datetime.now().isoformat()
        closed['zapisnik_image'] = _saved_rel
        if selected_dev:
            closed['returned_devices'] = selected_dev
        if selected_sim:
            closed['returned_sims'] = selected_sim
        zk.append(closed)
        write_zakljuceni_nalozi(zk)

        # (4) U globalnim nalozi.json označi kao zaključen (ako je prisutan)
        try:
            base = read_nalozi()
        except Exception:
            base = []
        found = False
        for it in base:
            if str(it.get('rn')) == str(rn):
                it.update({'status':'zaključen','closed_at':closed['closed_at'],'zapisnik_image': _saved_rel})
                found = True
                break
        if not found:
            base.append(closed)
        write_nalozi(base)

        # (5) Ukloni iz otvorene_deinstalacije.JSON
        od = [x for x in od if str(x.get('rn')) != str(rn)]
        write_otvorene_deinstalacije(od)

        flash(f"Deinstalacija RN {rn} zaključena.", "success")
        return redirect(url_for('gotove_deinstalacije'))

    # GET – prikaži uređaje i SIM-ove iz naloga
    dev_rows = [{'serijski': s, 'model': model_by_sn.get(s,'')} for s in dev_sns]
    sim_rows = [{'serijski': s, 'provider': provider_by_sn.get(s,'')} for s in sim_sns]
    return render_template(
        'zakljuci.deinstalaciju.html',
        title='Zaključi deinstalaciju',
        username=current_user.username,
        nalog=nalog,
        klijent={'name': nalog.get('client')},
        uredjaji_klijent=dev_rows,
        sims_klijent=sim_rows
    )

@app.route('/gotove-deinstalacije')
@login_required
def gotove_deinstalacije():
    """Lista svih zaključenih deinstalacija iz zakljuceni.nalozi.JSON."""
    try:
        all_closed = read_zakljuceni_nalozi()
    except Exception:
        all_closed = []
    deinstalacije = [n for n in all_closed if (n.get('type','') or '').lower() == 'deinstalacija']
    try:
        deinstalacije.sort(key=lambda x: (x.get('closed_at') or x.get('created_at') or ''), reverse=True)
    except Exception:
        pass
    return render_template('gotove_deinstalacije.html', title="Gotove deinstalacije", username=current_user.username, deinstalacije=deinstalacije)
# =================================================================================================

# ==================== API: SERVIS UREĐAJ -> POVRATAK U SKLADIŠTE / OTPIS ====================
@app.route('/api/servis/<serijski>/return', methods=['POST'])
@login_required
def api_servis_return(serijski):
    """Zelena kvačica: premjesti uređaj iz servis.uredjaji.JSON natrag u uredjaji.json i makni sa /servis-uredjaja."""
    try:
        serv_list = read_servis_uredjaji()
    except Exception:
        serv_list = []
    item = next((x for x in serv_list if str(x.get('serijski')) == str(serijski)), None)
    if not item:
        return jsonify({'ok': False, 'error': 'Uređaj nije pronađen u servisnoj listi.'}), 404
    # remove from servis list
    serv_list = [x for x in serv_list if str(x.get('serijski')) != str(serijski)]
    write_servis_uredjaji(serv_list)
    # add to skladiste uredjaji.json (ako već ne postoji)
    try:
        sklad = read_uredjaji()
    except Exception:
        sklad = []
    if not any(str(u.get('serijski')) == str(serijski) for u in sklad):
        sklad.append({
            'model': item.get('model',''),
            'serijski': item.get('serijski',''),
            'created_at': datetime.datetime.now().isoformat()
        })
        write_uredjaji(sklad)
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True})

@app.route('/api/servis/<serijski>/writeoff', methods=['POST'])
@login_required
def api_servis_writeoff(serijski):
    """Crveni X: premjesti uređaj iz servis.uredjaji.JSON u otpisani.uredjaji.JSON i prikaži u boxu 'Otpisani uređaji'."""
    try:
        serv_list = read_servis_uredjaji()
    except Exception:
        serv_list = []
    item = next((x for x in serv_list if str(x.get('serijski')) == str(serijski)), None)
    if not item:
        return jsonify({'ok': False, 'error': 'Uređaj nije pronađen u servisnoj listi.'}), 404
    # remove from servis
    serv_list = [x for x in serv_list if str(x.get('serijski')) != str(serijski)]
    write_servis_uredjaji(serv_list)
    # append to otpisani
    try:
        otp = read_otpisani_uredjaji()
    except Exception:
        otp = []
    new_item = dict(item)
    new_item['status'] = 'otpisan'
    new_item['otpisan_at'] = datetime.datetime.now().isoformat()
    otp.append(new_item)
    write_otpisani_uredjaji(otp)
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass
    return jsonify({'ok': True})

# --- Patch write_aktivni_uredjaji to include 'namjena' field ---
_original_write_aktivni_uredjaji = write_aktivni_uredjaji
def patched_write_aktivni_uredjaji(lst):
    try:
        for it in lst:
            if 'serijski' in it and 'namjena' not in it:
                it['namjena'] = get_namjena_by_serijski(it.get('serijski'))
    except Exception:
        pass
    return _original_write_aktivni_uredjaji(lst)
write_aktivni_uredjaji = patched_write_aktivni_uredjaji

@app.after_request
def inject_servis_link(response):
    # Uklonjeno: prestani dodavati gumb "Servis uređaja" na /uredjaji.
    # Dodatno: ako je isti gumb već u HTML-u odgovora, ukloni ga.
    try:
        if request.path == '/uredjaji' and response.content_type and response.content_type.startswith('text/html'):
            body = response.get_data(as_text=True)
            # Ukloni eventualno ubrizgan gumb (prethodna implementacija)
            try:
                injected_btn = '<div style="text-align:center;margin:24px 0;"><a class="btn btn-outline-secondary" href="%s">Servis uređaja</a></div>' % (url_for('servis_uredjaja'))
                body = body.replace(injected_btn, '')
            except Exception:
                pass
            response.set_data(body)
    except Exception:
        pass
    return response

@app.route('/zahtjev/delete', methods=['POST'])
@login_required
def delete_zahtjev():
    # RBAC: samo superadmin/admin smiju brisati
    if not has_role('superadmin','admin'):
        flash('Nemate ovlasti za brisanje.', 'danger')
        client = request.form.get('client') or ''
        return redirect(url_for('klijent_profil', name=client))
    rn = str(request.form.get('rn') or '').strip()
    src_rel = str(request.form.get('src') or '').strip()
    client = request.form.get('client') or ''
    src_rel = src_rel.replace('\\', '/')
    if src_rel.startswith('static/'):
        src_rel = src_rel[7:]
    abs_path = os.path.abspath(os.path.join(STATIC_DIR, src_rel))
    if not abs_path.startswith(os.path.abspath(STATIC_DIR)):
        flash('Neispravna putanja.', 'danger')
        return redirect(url_for('klijent_profil', name=client))
    try:
        if os.path.exists(abs_path):
            os.remove(abs_path)
    except Exception:
        pass
    # Ukloni referencu na Zahtjev iz naloga
    try:
        base = read_nalozi()
    except Exception:
        base = []
    changed = False
    for n in base:
        try:
            if str(n.get('rn')) == rn and n.get('zahtjev_image'):
                n['zahtjev_image'] = ''
                changed = True
        except Exception:
            pass
    if changed:
        try:
            write_nalozi(base)
        except Exception:
            pass
    flash('Zahtjev obrisan.', 'success')
    return redirect(url_for('klijent_profil', name=client))
# (removed stray app.run)



# -------------------- API: Pretraga klijenata (naziv ili OIB) --------------------
@app.route('/api/klijenti/search')
@login_required
def api_search_klijenti():
    try:
        q = (request.args.get('q') or '').strip()
    except Exception:
        q = ''
    try:
        data = read_klijenti() or []
    except Exception:
        data = []
    if not q:
        return jsonify({'ok': True, 'results': []})
    # normaliziraj tražilicu i očisti brojeve za OIB usporedbu
    q_norm = _norm(q)
    q_digits = re.sub(r'\D+', '', q)
    results = []
    for c in data:
        try:
            name = str(c.get('name',''))
            oib  = str(c.get('oib',''))
            name_ok = q_norm in _norm(name)
            oib_ok  = bool(q_digits) and (q_digits in re.sub(r'\D+','', oib or ''))
            if name_ok or oib_ok:
                results.append({'name': name, 'oib': oib})
                if len(results) >= 20:  # safety cap
                    break
        except Exception:
            pass
    return jsonify({'ok': True, 'results': results})


# -------------------- CHAT (desni sidebox) --------------------
try:
    STATIC_DIR
except NameError:
    STATIC_DIR = os.path.join(app.root_path, 'static')
# -------------------- NAPOMENE KLIJENTI (JSONL) --------------------
NAPOMENE_JSONL_PATH = os.path.join(STATIC_DIR, 'napomene_klijenti.JSONL')

def _ensure_jsonl(path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    if not os.path.exists(path):
        with open(path, 'w', encoding='utf-8') as _f:
            _f.write("")

def append_client_note(client_name: str, operator: str, text: str, ts: datetime.datetime = None):
    """Upiši jednu napomenu za klijenta u JSONL (jedan JSON po retku)."""
    _ensure_jsonl(NAPOMENE_JSONL_PATH)
    ts = ts or datetime.datetime.now()
    row = {
        "client": str(client_name or ""),
        "operator": str(operator or ""),
        "text": str(text or ""),
        "ts_iso": ts.isoformat(),
    }
    with open(NAPOMENE_JSONL_PATH, 'a', encoding='utf-8') as f:
        f.write(json.dumps(row, ensure_ascii=False) + "\n")
    return row

def read_client_notes(client_name: str):
    """Vrati sve napomene za zadani klijent iz JSONL, sortirane po vremenu (ASC)."""
    items = []
    try:
        _ensure_jsonl(NAPOMENE_JSONL_PATH)
        with open(NAPOMENE_JSONL_PATH, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    row = json.loads(line)
                except Exception:
                    continue
                if str(row.get("client","")) == str(client_name):
                    items.append(row)
    except Exception:
        items = []
    # sort po ts_iso
    try:
        items.sort(key=lambda r: r.get("ts_iso",""))
    except Exception:
        pass
    return items

@app.route('/api/klijent/<path:name>/napomene', methods=['GET','POST'])
@login_required
def api_klijent_napomene_new(name):
    """GET: vrati povijest napomena za klijenta.
       POST: dodaj novu napomenu (JSON: {"text": "..."}) i vrati ažuriranu listu.
    """
    if request.method == 'GET':
        notes = read_client_notes(name)
        return jsonify({"ok": True, "items": notes})

    data = {}
    try:
        data = request.get_json(silent=True) or {}
    except Exception:
        data = {}
    text = str(data.get('text') or '').strip()
    if not text:
        return jsonify({"ok": False, "error": "Tekst napomene je obavezan."}), 400
    op = getattr(current_user, 'username', '') or 'unknown'
    row = append_client_note(name, op, text)
    notes = read_client_notes(name)
    return jsonify({"ok": True, "created": row, "items": notes})

@app.route('/api/klijent/<path:name>/napomene/save', methods=['POST'])
@login_required
def api_klijent_napomene_save(name):
    """Bulk spremanje svih prikazanih napomena u JSONL (append).
       Očekuje JSON: {"items":[{"operator": "...", "text":"...", "ts_iso":"..."}]}
       Duplikati se preskaču na temelju (client, operator, ts_iso, text).
    """
    try:
        payload = request.get_json(silent=True) or {}
        items = payload.get('items') or []
        if not isinstance(items, list):
            return jsonify({"ok": False, "error": "Neispravan payload."}), 400

        # Učitaj postojeće zapise u set za deduplikaciju
        existing_keys = set()
        _ensure_jsonl(NAPOMENE_JSONL_PATH)
        try:
            with open(NAPOMENE_JSONL_PATH, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line: 
                        continue
                    try:
                        row = json.loads(line)
                        key = (str(row.get("client","")), str(row.get("operator","")), str(row.get("ts_iso","")), str(row.get("text","")))
                        existing_keys.add(key)
                    except Exception:
                        pass
        except Exception:
            pass

        appended = 0
        for it in items:
            text = str(it.get('text') or '').strip()
            operator = str(it.get('operator') or getattr(current_user, 'username', '') or 'unknown')
            ts_iso = (it.get('ts_iso') or '').strip()
            ts = None
            if ts_iso:
                try:
                    ts = datetime.datetime.fromisoformat(ts_iso[:26])
                except Exception:
                    ts = None
            key = (str(name), operator, ts_iso, text)
            if text and key not in existing_keys:
                append_client_note(name, operator, text, ts=ts)
                existing_keys.add(key)
                appended += 1
        return jsonify({"ok": True, "appended": appended})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


CHAT_JSON_PATH = os.path.join(STATIC_DIR, 'chat.json')
CHAT_LOG_PATH = os.path.join(STATIC_DIR, 'chat.log.jsonl')

def read_chat():
    try:
        with open(CHAT_JSON_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return []

def write_chat(data):
    try:
        os.makedirs(os.path.dirname(CHAT_JSON_PATH), exist_ok=True)
        with open(CHAT_JSON_PATH, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def append_chatlog(entry: dict):
    try:
        os.makedirs(os.path.dirname(CHAT_LOG_PATH), exist_ok=True)
        with open(CHAT_LOG_PATH, 'a', encoding='utf-8') as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\\n")
    except Exception:
        pass

@app.route('/api/chat', methods=['GET', 'POST'])
@login_required
def api_chat():
    if request.method == 'GET':
        msgs = read_chat()
        return jsonify({'ok': True, 'messages': msgs[-200:]}), 200

    # POST
    try:
        ops = read_operateri()
    except Exception:
        ops = []
    is_operator = any(str(o.get('username','')).strip().lower() == str(current_user.username or '').strip().lower() for o in (ops or []))
    if not is_operator:
        return jsonify({'ok': False, 'error': 'Samo operateri mogu pisati poruke.'}), 403

    data = request.get_json(silent=True) or {}
    text = (data.get('text') or '').strip()
    if not text:
        return jsonify({'ok': False, 'error': 'Prazna poruka.'}), 400

    now_iso = datetime.datetime.now().isoformat()
    msgs = read_chat()
    entry = {'username': current_user.username, 'text': text, 'ts': now_iso}
    msgs.append(entry)
    write_chat(msgs)
    append_chatlog(entry)
    return jsonify({'ok': True, 'count': len(msgs)}), 200

@app.route('/api/chat/log', methods=['GET'])
@login_required
def api_chat_log():
    try:
        return send_file(CHAT_LOG_PATH, mimetype='text/plain', as_attachment=True, download_name='chat.log.jsonl')
    except Exception:
        return jsonify({'ok': False, 'error': 'Log je prazan ili nedostupan.'}), 404

# -------------------- LOZINKA: ZABORAVLJENA --------------------

    # Pošalji e-mail s linkom za reset
    try:
        token = generate_reset_token(username)
        reset_url = url_for('reset_password', token=token, _external=True)
        subject = 'Zahtjev za promjenu lozinke'
        body = f"""Poštovani/na {op.get('first_name','') or ''} {op.get('last_name','') or ''},

Zaprimili smo zahtjev za promjenu lozinke za profil: {username}.
Za postavljanje nove lozinke otvorite poveznicu:
{reset_url}

Ako niste zatražili promjenu lozinke, slobodno zanemarite ovaj e-mail.

Srdačan pozdrav,
Billy sustav
"""
        html = f"""<p>Poštovani/na {op.get('first_name','') or ''} {op.get('last_name','') or ''},</p>
<p>Zaprimili smo zahtjev za promjenu lozinke za profil: <strong>{username}</strong>.</p>
<p><a href="{reset_url}" style="display:inline-block;padding:10px 16px;border-radius:6px;background:#355cfc;color:#fff;text-decoration:none;">Postavi novu lozinku</a></p>
<p>Ako gumb ne radi, kopirajte i zalijepite ovu adresu u preglednik:<br><code>{reset_url}</code></p>
<p>Ako niste zatražili promjenu lozinke, slobodno zanemarite ovaj e-mail.</p>
<p>Srdačan pozdrav,<br>Billy sustav</p>"""
        msg = Message(subject=subject, recipients=[email], body=body)
        try:
            msg.html = html
        except Exception:
            pass
        mail.send(msg)
        try:
            flash('Poslali smo poruku s poveznicom za promjenu lozinke.', 'info')
        except Exception:
            pass
    except Exception as e:
        try:
            flash(f'Greška pri slanju e-maila: {e}', 'danger')
        except Exception:
            pass
    return redirect(url_for('login'))


    # Potvrdi da operater postoji u sustavu
    try:
        ops = read_operateri()
    except Exception:
        ops = []
    op = None
    try:
        for o in ops:
            if str(o.get('email','')).strip().lower() == email.lower():
                op = o
                break
    except Exception:
        op = None
    if not op:
        try:
            flash('Operater s navedenim e-mailom ne postoji u sustavu.', 'warning')
        except Exception:
            pass
        return redirect(url_for('login'))

    # Pronađi korisnički račun
    username = str(op.get('username') or '').strip()
    try:
        db.create_all()
    except Exception:
        pass
    user = None
    if username:
        try:
            user = User.query.filter(db.func.lower(User.username) == username.lower()).first()
        except Exception:
            user = None
    if not user:
        try:
            flash('Nije pronađen korisnički račun za navedenog operatera.', 'warning')
        except Exception:
            pass
        return redirect(url_for('login'))

    # Pošalji e-mail
    try:
        subject = 'Zahtjev za promjenu lozinke'
        body = f"""Poštovani/na {op.get('first_name','') or ''} {op.get('last_name','') or ''},

Zaprimili smo zahtjev za promjenu lozinke za profil: {username}.
Ako ste Vi inicirali zahtjev, obratite se administratoru sustava kako bi zabilježio promjenu,
ili odgovorite na ovu poruku kako bismo potvrdili Vaš identitet i pokrenuli promjenu lozinke.

Ako niste zatražili promjenu lozinke, slobodno zanemarite ovaj e-mail.

Srdačan pozdrav,
Billy sustav
"""
        msg = Message(subject=subject, recipients=[email], body=body)
        mail.send(msg)
        try:
            flash('E-mail uspješno poslan!', 'success')
        except Exception:
            pass
    except Exception as e:
        try:
            flash(f'Greška pri slanju e-maila: {e}', 'danger')
        except Exception:
            pass
    return redirect(url_for('login'))


# -------------------- LOZINKA: RESET --------------------
@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    # Vrati username iz tokena (None ako je istekao/neispravan)
    username = verify_reset_token(token)
    if not username:
        flash('Poveznica za promjenu lozinke je nevažeća ili je istekla.', 'danger')
        return redirect(url_for('login'))
    user = User.query.filter(db.func.lower(User.username) == str(username).lower()).first()
    if not user:
        flash('Korisnički račun nije pronađen.', 'warning')
        return redirect(url_for('login'))
    if request.method == 'POST':
        new_pass = (request.form.get('password') or '').strip()
        new_pass2 = (request.form.get('password2') or '').strip()
        if not new_pass or not new_pass2:
            flash('Unesite novu lozinku u oba polja.', 'danger')
            return render_template('reset_password.html', title='Nova lozinka', token=token, username=username)
        if new_pass != new_pass2:
            flash('Lozinke se ne podudaraju.', 'danger')
            return render_template('reset_password.html', title='Nova lozinka', token=token, username=username)
        if len(new_pass) < 6:
            flash('Lozinka mora imati najmanje 6 znakova.', 'danger')
            return render_template('reset_password.html', title='Nova lozinka', token=token, username=username)
        try:
            user.set_password(new_pass)
            db.session.commit()
            flash('Lozinka je uspješno promijenjena. Prijavite se s novom lozinkom.', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            db.session.rollback()
            flash(f'Greška pri spremanju lozinke: {e}', 'danger')
            return render_template('reset_password.html', title='Nova lozinka', token=token, username=username)
    # GET
    return render_template('reset_password.html', title='Nova lozinka', token=token, username=username)


# -------------------- LOZINKA: ZABORAVLJENA --------------------
@app.route('/forgot-password', methods=['POST'])
def forgot_password():
    try:
        email = (request.form.get('email') or '').strip()
    except Exception:
        email = ''
    if not email:
        try:
            flash('Unesite e-mail adresu.', 'danger')
        except Exception:
            pass
        return redirect(url_for('login'))

    # Potvrdi da operater postoji u sustavu
    try:
        ops = read_operateri()
    except Exception:
        ops = []
    op = None
    try:
        for o in ops:
            if str(o.get('email','')).strip().lower() == email.lower():
                op = o
                break
    except Exception:
        op = None
    if not op:
        try:
            flash('Operater s navedenim e-mailom ne postoji u sustavu.', 'warning')
        except Exception:
            pass
        return redirect(url_for('login'))

    # Pronađi korisnički račun
    username = str(op.get('username') or '').strip()
    try:
        db.create_all()
    except Exception:
        pass
    user = None
    if username:
        try:
            user = User.query.filter(db.func.lower(User.username) == username.lower()).first()
        except Exception:
            user = None
    if not user:
        try:
            flash('Nije pronađen korisnički račun za navedenog operatera.', 'warning')
        except Exception:
            pass
        return redirect(url_for('login'))

    # Pošalji e-mail s linkom za reset
    try:
        token = generate_reset_token(username)
        reset_url = url_for('reset_password', token=token, _external=True)
        subject = 'Zahtjev za promjenu lozinke'
        body = f"""Poštovani/na {op.get('first_name','') or ''} {op.get('last_name','') or ''},

Zaprimili smo zahtjev za promjenu lozinke za profil: {username}.
Za postavljanje nove lozinke otvorite poveznicu:
{reset_url}

Ako niste zatražili promjenu lozinke, slobodno zanemarite ovaj e-mail.

Srdačan pozdrav,
Billy sustav
"""
        html = f"""<p>Poštovani/na {op.get('first_name','') or ''} {op.get('last_name','') or ''},</p>
<p>Zaprimili smo zahtjev za promjenu lozinke za profil: <strong>{username}</strong>.</p>
<p><a href="{reset_url}" style="display:inline-block;padding:10px 16px;border-radius:6px;background:#355cfc;color:#fff;text-decoration:none;">Postavi novu lozinku</a></p>
<p>Ako gumb ne radi, kopirajte i zalijepite ovu adresu u preglednik:<br><code>{reset_url}</code></p>
<p>Ako niste zatražili promjenu lozinke, slobodno zanemarite ovaj e-mail.</p>
<p>Srdačan pozdrav,<br>Billy sustav</p>"""
        msg = Message(subject=subject, recipients=[email], body=body)
        try:
            msg.html = html
        except Exception:
            pass
        mail.send(msg)
        try:
            flash('Poslali smo poruku s poveznicom za promjenu lozinke.', 'info')
        except Exception:
            pass
    except Exception as e:
        try:
            flash(f'Greška pri slanju e-maila: {e}', 'danger')
        except Exception:
            pass
    return redirect(url_for('login'))

    def _to_str(v):
        if v is None: return ''
        # Avoid 1.0 for integers
        if isinstance(v, float) and v.is_integer():
            return str(int(v))
        return str(v).strip()

    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    # Header detection
    try:
        header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        header = [ _to_str(x) for x in header_row ]
    except Exception:
        header = []

    H = [h.lower() for h in header]
    def find_idx(names):
        for n in names:
            try:
                i = H.index(n.lower())
                return i
            except ValueError:
                continue
        return -1

    idx_provider = find_idx(['Provider','Operater','Pružatelj','Provider/Operater'])
    idx_sn       = find_idx(['Serijski broj','Serijski','SN','Serijski_broj','Serijski broj SIM'])

    header_mode = 'explicit'
    if idx_provider == -1 or idx_sn == -1:
        # Fallback: assume first two columns
        idx_provider, idx_sn = 0, 1
        header_mode = 'fallback-first-two'

    # Load base
    try:
        base = read_sim()
    except Exception:
        base = []
    if not isinstance(base, list):
        base = []

    existing = {_norm_sn((x.get('serijski') if isinstance(x, dict) else '')) for x in base}

    total_rows = 0
    added = 0
    skipped_existing = 0
    skipped_invalid = 0
    skipped_dup_in_file = 0
    seen_in_file = set()

    # Process rows
    for row in ws.iter_rows(min_row=2, values_only=True):
        total_rows += 1
        provider = _to_str(row[idx_provider] if (row and idx_provider < len(row)) else '')
        sn_raw   = _to_str(row[idx_sn]       if (row and idx_sn       < len(row)) else '')

        # Require both present
        if not provider or not sn_raw:
            skipped_invalid += 1
            continue

        # Accept almost any visible char; we only normalize for duplicate check
        sn_norm = _norm_sn(sn_raw)
        if not sn_norm:
            skipped_invalid += 1
            continue

        if sn_norm in seen_in_file:
            skipped_dup_in_file += 1
            continue
        seen_in_file.add(sn_norm)

        if sn_norm in existing:
            skipped_existing += 1
            continue

        base.append({'provider': provider, 'serijski': sn_raw, 'created_at': datetime.datetime.now().isoformat()})
        existing.add(sn_norm)
        added += 1

    # Save
    try:
        write_sim(base)
    except Exception as e:
        try:
            flash(f'Greška pri spremanju: {e}', 'danger')
        except Exception:
            pass
        try:
            print('[SIM IMPORT] save error:', e)
        except Exception:
            pass
        return redirect(url_for('sim'))

    summary = f'Uvoz: header={header_mode}, redaka={total_rows}, dodano={added}, već postoji={skipped_existing}, duplikati u datoteci={skipped_dup_in_file}, neispravno={skipped_invalid}.'
    try:
        flash(summary, 'success' if added else 'warning')
    except Exception:
            pass
    try:
        print('[SIM IMPORT]', summary)
    except Exception:
            pass

    return redirect(url_for('sim'))

# -------------------- IMPORT SIM (XLSX) --------------------


    def _to_str(v):
        if v is None: return ''
        if isinstance(v, float) and v.is_integer(): return str(int(v))
        return str(v).strip()

    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    try:
        header_cells = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        header = [_to_str(x) for x in header_cells]
    except Exception:
        header = []

    H = [h.lower() for h in header]
    def find_idx(names):
        for n in names:
            try:
                return H.index(n.lower())
            except ValueError: continue
        return -1

    idx_provider = find_idx(['Provider','Operater','Pružatelj','Provider/Operater'])
    idx_sn       = find_idx(['Serijski broj','Serijski','SN','Serijski_broj','Serijski broj SIM'])

    header_mode = 'explicit'
    if idx_provider == -1 or idx_sn == -1:
        idx_provider, idx_sn = 0, 1
        header_mode = 'fallback-first-two'

    try:
        base = read_sim()
    except Exception: base = []
    if not isinstance(base, list): base = []
    existing = {_norm_sn((x.get('serijski') if isinstance(x, dict) else '')) for x in base}

    total_rows = 0
    added = 0
    skipped_existing = 0
    skipped_invalid = 0
    skipped_dup_in_file = 0
    seen_in_file = set()

    for row in ws.iter_rows(min_row=2, values_only=True):
        total_rows += 1
        provider = _to_str(row[idx_provider] if (row and idx_provider < len(row)) else '')
        sn_raw   = _to_str(row[idx_sn]       if (row and idx_sn       < len(row)) else '')

        if not provider or not sn_raw:
            skipped_invalid += 1
            continue

        sn_norm = _norm_sn(sn_raw)
        if not sn_norm:
            skipped_invalid += 1
            continue

        if sn_norm in seen_in_file:
            skipped_dup_in_file += 1
            continue
        seen_in_file.add(sn_norm)

        if sn_norm in existing:
            skipped_existing += 1
            continue

        base.append({'provider': provider, 'serijski': sn_raw, 'created_at': datetime.datetime.now().isoformat()})
        existing.add(sn_norm)
        added += 1

    try:
        write_sim(base)
    except Exception as e:
        try:
            flash(f'Greška pri spremanju: {e}', 'danger')
        except Exception:
            pass
        try:
            print('[SIM IMPORT] save error:', e)
        except Exception:
            pass
        return redirect(url_for('sim'))

    summary = f'Uvoz: header={header_mode}, redaka={total_rows}, dodano={added}, već postoji={skipped_existing}, duplikati u datoteci={skipped_dup_in_file}, neispravno={skipped_invalid}.'
    try:
        flash(summary, 'success' if added else 'warning')
    except Exception:
            pass
    try:
        print('[SIM IMPORT]', summary)
    except Exception:
            pass

    return redirect(url_for('sim'))

# -------------------- IMPORT SIM (XLSX) --------------------
@app.route('/sim/import', methods=['POST'], endpoint='import_sim_xlsx')
@login_required
def import_sim_xlsx():
    if not has_role('superadmin','admin','prodaja'):
        try:
            flash('Nemate ovlasti za uvoz SIM kartica.', 'danger')
        except Exception:
            pass
        return redirect(url_for('sim'))

    # Lokalni import da izbjegnemo NameError
    try:
        from openpyxl import load_workbook as _load_workbook
    except Exception:
        try:
            import openpyxl as _oxl
            _load_workbook = _oxl.load_workbook
        except Exception as e:
            try:
                flash(f'Nedostaje openpyxl: {e}', 'danger')
            except Exception:
                pass
            pass
            return redirect(url_for('sim'))

    file = request.files.get('file')
    if not file or not file.filename.lower().endswith('.xlsx'):
        try:
            flash('Odaberite .xlsx datoteku.', 'danger')
        except Exception:
            pass
        return redirect(url_for('sim'))

    try:
        wb = _load_workbook(file, read_only=True, data_only=True)
        ws = wb.active
    except Exception as e:
        try:
            flash(f'Ne mogu pročitati XLSX: {e}', 'danger')
        except Exception:
            pass
        return redirect(url_for('sim'))

    def _to_str(v):
        if v is None: return ''
        if isinstance(v, float) and v.is_integer(): return str(int(v))
        return str(v).strip()

    def _norm_sn(v):
        return re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    try:
        header_cells = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        header = [_to_str(x) for x in header_cells]
    except Exception:
        header = []

    H = [h.lower() for h in header]
    def find_idx(names):
        for n in names:
            try:
                return H.index(n.lower())
            except ValueError: continue
        return -1

    idx_provider = find_idx(['Provider','Operater','Pružatelj','Provider/Operater'])
    idx_sn       = find_idx(['Serijski broj','Serijski','SN','Serijski_broj','Serijski broj SIM'])

    header_mode = 'explicit'
    if idx_provider == -1 or idx_sn == -1:
        idx_provider, idx_sn = 0, 1
        header_mode = 'fallback-first-two'

    try:
        base = read_sim()
    except Exception: base = []
    if not isinstance(base, list): base = []
    existing = {_norm_sn((x.get('serijski') if isinstance(x, dict) else '')) for x in base}

    total_rows = 0
    added = 0
    skipped_existing = 0
    skipped_invalid = 0
    skipped_dup_in_file = 0
    seen_in_file = set()

    for row in ws.iter_rows(min_row=2, values_only=True):
        total_rows += 1
        provider = _to_str(row[idx_provider] if (row and idx_provider < len(row)) else '')
        sn_raw   = _to_str(row[idx_sn]       if (row and idx_sn       < len(row)) else '')

        if not provider or not sn_raw:
            skipped_invalid += 1
            continue

        sn_norm = _norm_sn(sn_raw)
        if not sn_norm:
            skipped_invalid += 1
            continue

        if sn_norm in seen_in_file:
            skipped_dup_in_file += 1
            continue
        seen_in_file.add(sn_norm)

        if sn_norm in existing:
            skipped_existing += 1
            continue

        base.append({'provider': provider, 'serijski': sn_raw, 'created_at': datetime.datetime.now().isoformat()})
        existing.add(sn_norm)
        added += 1

    try:
        write_sim(base)
    except Exception as e:
        try:
            flash(f'Greška pri spremanju: {e}', 'danger')
        except Exception:
            pass
        try:
            print('[SIM IMPORT] save error:', e)
        except Exception:
            pass
        return redirect(url_for('sim'))

    summary = f'Uvoz: header={header_mode}, redaka={total_rows}, dodano={added}, već postoji={skipped_existing}, duplikati u datoteci={skipped_dup_in_file}, neispravno={skipped_invalid}.'
    try:
        flash(summary, 'success' if added else 'warning')
    except Exception:
            pass
    try:
        print('[SIM IMPORT]', summary)
    except Exception:
            pass

    return redirect(url_for('sim'))



# -------------------- EXPORT SIM (XLSX) — early registered --------------------
@app.route('/sim/export.xlsx', methods=['GET'], endpoint='export_sim_xlsx')
@login_required
def export_sim_xlsx():
    """Preuzmi XLSX popis SIM (Provider, Serijski broj)."""
    # Lokalni importi da ne diramo globalne import sekcije
    try:
        from openpyxl import Workbook as _WB
    except Exception as e:
        try:
            flash(f'Nedostaje openpyxl: {e}', 'danger')
        except Exception:
            pass
        return redirect(url_for('sim'))
    try:
        from flask import send_file as _send_file
    except Exception:
        # Ako send_file iz nekog razloga nije dostupan, fallback poruka
        try:
            flash('Nedostaje send_file iz Flask-a.', 'danger')
        except Exception:
            pass
        return redirect(url_for('sim'))
    import io as _io

    try:
        items = read_sim()
    except Exception:
        items = []

    wb = _WB()
    ws = wb.active
    ws.title = "SIM"
    ws.append(["Provider", "Serijski broj"])
    for it in items:
        provider = str(it.get('provider', '') if isinstance(it, dict) else '')
        serijski = str(it.get('serijski', '') if isinstance(it, dict) else '')
        ws.append([provider, serijski])

    buf = _io.BytesIO()
    wb.save(buf); buf.seek(0)
    return _send_file(
        buf,
        as_attachment=True,
        download_name="sim.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# -------------------- UREĐAJ - DETALJI --------------------
@app.route('/uredjaji/<path:serijski>', endpoint='uredjaj_detalj')
@login_required
def uredjaj_detalj(serijski):
    # Učitaj sve uređaje
    items = read_uredjaji() or []
    dev = next((d for d in items if str(d.get('serijski')) == str(serijski)), None)
    if not dev:
        abort(404)

    # --- Status & popratne info (isti princip kao lista /uredjaji) ---
    import re as _re
    def _norm_sn(v):
        return _re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    try:
        aktivni = read_aktivni_uredjaji()
    except Exception:
        aktivni = []
    active_serials = {_norm_sn(d.get('serijski')) for d in aktivni if d.get('active', True)}
    active_client_by_sn = {}
    for d in aktivni:
        try:
            if d.get('active', True) and d.get('client'):
                active_client_by_sn[_norm_sn(d.get('serijski'))] = str(d.get('client'))
        except Exception:
            pass

    try:
        zaduzeni = read_zaduzene_uredjaje()
    except Exception:
        zaduzeni = []
    tech_by_serial = {_norm_sn(z.get('serijski')): str(z.get('assigned_to') or '') for z in zaduzeni}

    try:
        klijenti = read_klijenti()
    except Exception:
        klijenti = []
    client_by_serial = {}
    for k in klijenti:
        cname = str(k.get('name') or '').strip()
        try:
            for sn in str(k.get('sn_uredjaja') or '').split(','):
                snn = _norm_sn(sn)
                if snn:
                    client_by_serial[snn] = cname
        except Exception:
            pass
    for snn, cname in active_client_by_sn.items():
        client_by_serial.setdefault(snn, cname)

    sn_norm = _norm_sn(serijski)
    assigned_to = tech_by_serial.get(sn_norm, '')
    is_assigned = bool(assigned_to)
    # Privremeno isključen override (orange)
    try:
        _priv = read_privremeno_iskljuceni()
    except Exception:
        _priv = []
    import re as _re
    _pset = {_re.sub(r'[^A-Za-z0-9]', '', str(p.get('serijski') or '')).upper() for p in _priv}
    if sn_norm in _pset:
        status_color = 'orange'
        status_label = 'Privremeno isključen'
    elif sn_norm in active_serials or sn_norm in client_by_serial:
        status_color = 'green'
        status_label = 'Aktivan kod klijenta'
    elif (not is_assigned) and (sn_norm not in active_serials) and (sn_norm not in client_by_serial):
        status_color = 'red'
        status_label = 'Neaktivan'
    else:
        status_color = 'yellow'
        status_label = 'Zadužen tehničaru'

    dev_view = dict(dev)
    dev_view['assigned_to'] = assigned_to
    dev_view['client_name'] = client_by_serial.get(sn_norm, '')
    dev_view['status_color'] = status_color
    dev_view['status_label'] = status_label

    logs = read_device_log_for(serijski)
    return render_template('uredjaj_detalji.html',
                           title=f"Uređaj {serijski}",
                           username=getattr(current_user, 'username', None),
                           uredjaj=dev_view,
                           logs=logs)
# -------------------- PRIVREMENO ISKLJUČI --------------------
@app.route('/uredjaji/<path:serijski>/status', methods=['POST'])
@login_required
def uredjaj_set_status(serijski):
    items = read_uredjaji() or []
    dev = next((d for d in items if str(d.get('serijski')) == str(serijski)), None)
    if not dev:
        abort(404)
    status = request.form.get('status')
    if status not in ('Aktivan','Privremeno isključen'):
        abort(400)
    # Do not allow Privremeno isključen for Najam
    if status == 'Privremeno isključen' and str(dev.get('namjena')).lower() == 'najam':
        flash('Uređaji s namjenom Najam ne mogu biti privremeno isključeni.','error')
        return redirect(url_for('uredjaj_detalj', serijski=serijski))
    # Save new field to JSON store if applicable
    dev['custom_status'] = status
    # Update JSON file on disk
    try:
        import json
        path = os.path.join(app.static_folder,'uredjaji.JSON')
        with open(path,'r',encoding='utf-8') as fr:
            data = json.load(fr)
        for d in data:
            if str(d.get('serijski'))==str(serijski):
                d['custom_status']=status
        with open(path,'w',encoding='utf-8') as fw:
            json.dump(data,fw,ensure_ascii=False,indent=2)
    except Exception as e:
        print('Greška spremanja statusa:',e)
    flash(f'Status uređaja promijenjen u {status}.','success')
    return redirect(url_for('uredjaj_detalj', serijski=serijski))



# -------------------- API: POST promjena statusa uređaja (aktivan/privremeno isključen) --------------------
@app.route('/api/uredjaj/<path:serijski>/status', methods=['POST'])
@login_required
def api_set_uredjaj_status(serijski):
    status = (request.form.get('status') or '').strip().lower()
    # Dozvoli samo 'aktivan' ili 'privremeno'
    if status not in ('aktivan', 'privremeno'):
        return jsonify({'ok': False, 'error': 'Nepoznat status.'}), 400

    # Namjena: samo 'Kupnja' smije biti privremeno isključen
    base = read_uredjaji() or []
    dev = next((d for d in base if str(d.get('serijski')) == str(serijski)), None)
    namjena = (dev or {}).get('namjena') or ''
    if status == 'privremeno' and str(namjena) != 'Kupnja':
        return jsonify({'ok': False, 'error': "Samo uređaji s namjenom 'Kupnja' mogu biti privremeno isključeni."}), 400

    # Ažuriraj listu u JSON-u
    lst = read_privremeno_iskljuceni()
    sn = str(serijski)
    exists = next((x for x in lst if str(x.get('serijski')) == sn), None)
    if status == 'privremeno':
        if not exists:
            lst.append({'serijski': sn, 'since': datetime.datetime.now().isoformat()})
            write_privremeno_iskljuceni(lst)
    else:  # aktivan -> ukloni iz liste
        lst = [x for x in lst if str(x.get('serijski')) != sn]
        write_privremeno_iskljuceni(lst)
    try:
        _act = 'privremeno isključen' if status == 'privremeno' else 'aktiviran'
        append_device_log(serijski, _act)
    except Exception:
        pass

    return jsonify({'ok': True})


# -------------------- DEVICE CHANGE LOG (JSONL in static/device_log.jsonl) --------------------
def _device_log_path():
    try:
        return os.path.join(app.static_folder, 'device_log.jsonl')
    except Exception:
        # fallback to local static folder relative to this file
        return os.path.join(os.path.dirname(__file__), 'static', 'device_log.jsonl')

def append_device_log(serijski: str, action: str, who: str = None, details: dict = None):
    """Append a single JSON line describing a device change."""
    try:
        os.makedirs(app.static_folder, exist_ok=True)
    except Exception:
        pass
    entry = {
        "ts": datetime.datetime.now().isoformat(),
        "serijski": str(serijski),
        "action": str(action),
        "who": who or getattr(current_user, 'username', None),
        "details": details or {}
    }
    try:
        path = _device_log_path()
        with open(path, "a", encoding="utf-8") as fw:
            fw.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception as e:
        try:
            app.logger.warning(f"Ne mogu upisati device_log.jsonl: {e}")
        except Exception:
            pass

def read_device_log_for(serijski: str, limit: int = 200):
    """Return recent log entries (newest last) for given serial."""
    out = []
    path = _device_log_path()
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as fr:
                for line in fr:
                    try:
                        item = json.loads(line.strip())
                        if str(item.get("serijski")) == str(serijski):
                            out.append(item)
                    except Exception:
                        continue
    except Exception:
        return []
    if limit and len(out) > limit:
        out = out[-limit:]
    return out


# -------------------- API: POST promjena namjene uređaja (Kupnja/Najam) --------------------
@app.route('/api/uredjaj/<path:serijski>/namjena', methods=['POST'])
@login_required
def api_set_uredjaj_namjena(serijski):
    # RBAC: samo Superadmin/Admin/Serviser/Podrška smiju mijenjati namjenu
    if not has_role('superadmin','admin','serviser','podrška'):
        return jsonify({'ok': False, 'error': 'Zabranjeno: nedovoljna prava.'}), 403
    new_val = (request.form.get('namjena') or '').strip()
    if new_val not in ('Kupnja','Najam'):
        return jsonify({'ok': False, 'error': "Namjena mora biti 'Kupnja' ili 'Najam'."}), 400
    # Ažuriraj u static/uredjaji.JSON
    base = read_uredjaji() or []
    idx = None
    for i, d in enumerate(base):
        if str(d.get('serijski')) == str(serijski):
            idx = i
            break
    if idx is None:
        return jsonify({'ok': False, 'error': 'Uređaj nije pronađen.'}), 404
    base[idx]['namjena'] = new_val
    try:
        write_uredjaji(base)
    except Exception as e:
        return jsonify({'ok': False, 'error': f'Greška pri zapisu: {e}'}), 500
    try:
        append_device_log(serijski, f"namjena promijenjena na {new_val}")
    except Exception:
        pass
    return jsonify({'ok': True, 'namjena': new_val}), 200
@app.route('/sim/<path:serijski>', endpoint='sim_detail')
@login_required
def sim_detail(serijski):
    # Učitaj sve SIM-ove iz sim.json
    items = read_sim() or []

    def _norm_sn(v: str) -> str:
        import re as _re
        return _re.sub(r'[^A-Za-z0-9]', '', str(v or '')).upper()

    # Mapiranja za status/klijenta/tehničara isto kao na /sim
    try:
        aktivni_sim = read_aktivni_sim()
    except Exception:
        aktivni_sim = []
    active_serials = {_norm_sn(s.get('serijski')) for s in aktivni_sim if s.get('active', True)}
    active_client_by_sn = {}

    # Privremeno isključeni SIM-ovi
    try:
        _priv_sim = read_privremeno_iskljuceni_sim()
    except Exception:
        _priv_sim = []
    priv_set = { _norm_sn(x.get('serijski')) for x in _priv_sim if x.get('serijski') }
    for s in aktivni_sim:
        try:
            if s.get('active', True) and s.get('client'):
                active_client_by_sn[_norm_sn(s.get('serijski'))] = str(s.get('client'))
        except Exception:
            pass

    try:
        zaduzeni_sim = read_zaduzene_sim()
    except Exception:
        zaduzeni_sim = []
    tech_by_serial = {_norm_sn(z.get('serijski')): str(z.get('assigned_to') or '') for z in zaduzeni_sim}

    try:
        klijenti = read_klijenti()
    except Exception:
        klijenti = []
    client_by_serial = {}
    for k in klijenti:
        cname = str(k.get('name') or '').strip()
        try:
            for sn in str(k.get('sn_SIM') or '').split(','):
                snn = _norm_sn(sn)
                if snn:
                    client_by_serial[snn] = cname
        except Exception:
            pass
    for snn, cname in active_client_by_sn.items():
        client_by_serial.setdefault(snn, cname)

    # Nađi traženi SIM (case-insensitive na originalnom serijskom)
    req_norm = _norm_sn(serijski)
    sim_item = None
    for it in items:
        if _norm_sn(it.get('serijski')) == req_norm:
            sim_item = dict(it)
            break
    if not sim_item:
        abort(404, description="SIM nije pronađen.")

    sn = _norm_sn(sim_item.get('serijski'))
    assigned_to = tech_by_serial.get(sn, '')
    is_assigned = bool(assigned_to)

    # Boja statusa — jednako kao na listi /sim
    if sn in priv_set:
        status_color = 'orange'
    elif sn in active_serials or sn in client_by_serial:
        status_color = 'green'
    elif (not is_assigned) and (sn not in active_serials) and (sn not in client_by_serial):
        status_color = 'red'
    else:
        status_color = 'yellow'

    sim_item['status_color'] = status_color
    sim_item['status'] = ('privremeno isključen' if status_color=='orange' else ('aktivan' if (sn in active_serials or sn in client_by_serial) else ('neaktivan' if (not is_assigned and (sn not in active_serials) and (sn not in client_by_serial)) else 'zadužen')))
    sim_item['assigned_to'] = assigned_to
    sim_item['client_name'] = client_by_serial.get(sn, '')

    return render_template(
        'sim_detail.html',
        title=f"SIM {sim_item.get('serijski','')}",
        username=getattr(current_user, 'username', None),
        sim=sim_item
    )

# === AUTO-GENERATED: POČETNA ANALYTICS ===
from flask import render_template
import datetime as dt

def _parse_dt(v):
    if not v:
        return None
    try:
        return dt.datetime.fromisoformat(str(v)[:19])
    except Exception:
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d.%m.%Y %H:%M", "%d.%m.%Y"):
            try:
                return dt.datetime.strptime(str(v), fmt)
            except Exception:
                pass
    return None

def _count_in_range(items, date_keys, start, end):
    cnt = 0
    for it in items or []:
        dt = None
        for k in (date_keys or []):
            if k in it and it.get(k):
                dt = _parse_dt(it.get(k))
                if dt: break
        if not dt:
            continue
        if (start is None or dt >= start) and (end is None or dt < end):
            cnt += 1
    return cnt

def _time_points(items, date_keys, start, end, step_days=1):
    out = []
    if start is None or end is None:
        return out
    cur = start
    while cur < end:
        nxt = cur + dt.timedelta(days=step_days)
        label = cur.strftime("%d.%m.")
        out.append({"label": label, "value": _count_in_range(items, date_keys, cur, nxt)})
        cur = nxt
    return out

def _month_buckets(items, date_keys, months_back=12):
    now = dt.datetime.now()
    labels, vals = [], []
    for i in range(months_back, 0, -1):
        yy = now.year
        mm = now.month - (i - 1)
        while mm <= 0:
            mm += 12
            yy -= 1
        start = dt.datetime(yy, mm, 1)
        if mm == 12:
            end = dt.datetime(yy+1, 1, 1)
        else:
            end = dt.datetime(yy, mm+1, 1)
        labels.append(start.strftime("%b %Y"))
        vals.append(_count_in_range(items, date_keys, start, end))
    return labels, vals

@app.route('/pocetna')
@login_required
def pocetna():
    try: klijenti = read_klijenti()
    except Exception: klijenti = []
    try: uredjaji = read_uredjaji()
    except Exception: uredjaji = []
    try: aktivni_uredjaji = read_aktivni_uredjaji()
    except Exception: aktivni_uredjaji = []
    try: zaduzeni_uredjaji = read_zaduzene_uredjaje()
    except Exception: zaduzeni_uredjaji = []
    try: simovi = read_sim()
    except Exception: simovi = []
    try: aktivni_sim = read_aktivni_sim()
    except Exception: aktivni_sim = []
    try: zaduzeni_sim = read_zaduzene_sim()
    except Exception: zaduzeni_sim = []
    try: nalozi = read_nalozi()
    except Exception: nalozi = []

    now = dt.datetime.now()
    ranges = {
        "7d": (now - dt.timedelta(days=7), now, 1),
        "30d": (now - dt.timedelta(days=30), now, 1),
        "6m": (now - dt.timedelta(days=182), now, 7),
        "1y": (now - dt.timedelta(days=365), now, 7),
        "all": (None, now, 30),
    }

    def build_series(items, date_keys):
        out = {}
        for key,(start,end,step) in ranges.items():
            curr = _count_in_range(items, date_keys, start, end)
            if start is not None:
                prev_start = start - (end - start)
                prev_end = start
                prev = _count_in_range(items, date_keys, prev_start, prev_end)
                points = _time_points(items, date_keys, start, end, step_days=step)
            else:
                prev = 0
                year_ago = now - dt.timedelta(days=365)
                points = _time_points(items, date_keys, year_ago, now, step_days=step)
            out[key] = {"curr": curr, "prev": prev, "points": points}
        return out

    cards = {
        "new_clients": {"title":"Novi klijenti","subtitle":"Broj novih klijenata","series": build_series(klijenti, ["created_at","createdAt","datum","date"])},
        "new_devices": {"title":"Novi uređaji","subtitle":"Dodani uređaji","series": build_series(uredjaji, ["created_at","createdAt","datum","date"])},
        "active_devices": {"title":"Aktivni uređaji (trenutno)","subtitle":"Broj aktivnih (status)","series": build_series([d for d in aktivni_uredjaji if d.get("active", True)], ["created_at","assigned_at","createdAt","date"])},
        "assigned_devices": {"title":"Zaduženi uređaji","subtitle":"Zaduženja tehničarima","series": build_series(zaduzeni_uredjaji, ["assigned_at","created_at","date"])},
        "new_sims": {"title":"Nove SIM kartice","subtitle":"Dodane SIM kartice","series": build_series(simovi, ["created_at","createdAt","date"])},
        "active_sims": {"title":"Aktivni SIM (trenutno)","subtitle":"Broj aktivnih SIM","series": build_series([s for s in aktivni_sim if s.get("active", True)], ["created_at","assigned_at","createdAt","date"])},
        "assigned_sims": {"title":"Zaduženi SIM","subtitle":"Zaduženja SIM-ova","series": build_series(zaduzeni_sim, ["assigned_at","created_at","date"])},
        "orders_all": {"title":"Nalozi (svi)","subtitle":"Kreirani nalozi","series": build_series(nalozi, ["created_at","date"])},
        "orders_install": {"title":"Instalacije","subtitle":"Nalozi tipa instalacija","series": build_series([n for n in nalozi if str(n.get('type','')).lower()=='instalacija'], ["created_at","date"])},
        "orders_deinstall": {"title":"Deinstalacije","subtitle":"Nalozi tipa deinstalacija","series": build_series([n for n in nalozi if str(n.get('type','')).lower()=='deinstalacija'], ["created_at","date"])},
        "orders_service": {"title":"Servisi","subtitle":"Nalozi tipa servis","series": build_series([n for n in nalozi if str(n.get('type','')).lower()=='servis'], ["created_at","date"])}
    }

    selected_range = "30d"
    for k,v in cards.items():
        v["value"] = v["series"][selected_range]["curr"]

    all_items_for_heat = (klijenti or []) + (uredjaji or []) + (nalozi or [])
    labels, values = _month_buckets(all_items_for_heat, ["created_at","createdAt","date"], months_back=12)
    analytics = {"cards": cards, "monthly": {"labels": labels, "values": values}}
    return render_template('pocetna.html', analytics_json=json.dumps(analytics, ensure_ascii=False), selected_range=selected_range, cards=cards)

# === AUTO-GENERATED: CONTEXT PROCESSOR HAS_POCETNA ===
@app.context_processor
def inject_has_pocetna():
    try:
        return {"has_pocetna": "pocetna" in app.view_functions}
    except Exception:
        return {"has_pocetna": False}

# === AUTO-INJECT DELETE BUTTONS AND ACTION ===
@app.after_request
def inject_delete_buttons(response):
    try:
        ctype = response.headers.get('Content-Type', '')
        if 'text/html' not in ctype.lower():
            return response
        html = response.get_data(as_text=True)
        if '<!--DEL-INJECT-->' in html:
            return response

        # CSS for the small grey delete button
        css = r"""<!--DEL-INJECT-->
<style>
.btn-delete{display:inline-flex;align-items:center;justify-content:center;gap:6px;padding:4px 8px;font-size:12px;line-height:1;border:1px solid #bdbdbd;border-radius:6px;background:#e0e0e0;color:#333;cursor:pointer;transition:background .15s ease,transform .05s ease}
.btn-delete:hover{background:#d5d5d5}
.btn-delete:active{transform:translateY(1px)}
.btn-delete img.icon-delete{width:14px;height:14px;object-fit:contain;display:inline-block}
</style>
"""
        try:
            icon_url = url_for('static', filename='icons/delete.png')
        except Exception:
            icon_url = '/static/icons/delete.png'

        js = r"""<script>
(function(){
  function r(fn){ if(document.readyState==='loading'){document.addEventListener('DOMContentLoaded',fn,{once:true});} else {fn();} }
  r(function(){
    var DELETE_ICON = '%ICON_URL%';
    function makeButton(el, href){
      var btn=document.createElement('button');
      btn.type='button'; btn.className='btn-delete'; btn.title='Obriši';
      btn.innerHTML='<img class="icon-delete" alt="Obriši" src="'+DELETE_ICON+'">';
      // Extract type/id from href or dataset
      var m = href && href.match(/\/(?:api\/)?(sim|uredjaj|nalog)\/([^\/]+)\/delete/i);
      var type = (m && m[1]||'').toLowerCase();
      var id = (m && m[2]) ? decodeURIComponent(m[2]) : (el.getAttribute('data-id')||'');
      btn.setAttribute('data-type', type);
      btn.setAttribute('data-id', id);
      var tr = el.closest && el.closest('tr'); if(tr && id && !tr.getAttribute('data-row-id')) tr.setAttribute('data-row-id', id);
      el.parentNode.replaceChild(btn, el);
    }
    function normalize(){
      var sels = 'a[href*="/delete"],[data-action="delete"],[data-delete],.delete,.icon-delete-btn,i[class*="trash"],i[class*="delete"],img[src*="trash"],img[src*="delete"]';
      var list = document.querySelectorAll(sels);
      for(var i=0;i<list.length;i++){
        var el=list[i]; if(el.__norm) continue; el.__norm=1;
        var href = el.getAttribute && (el.getAttribute('href')||el.getAttribute('formaction')||'');
        if(!href){
          var a = el.closest && el.closest('a[href]'); if(a) href = a.getAttribute('href');
        }
        if(!href){ continue; }
        // Must contain "/delete" to avoid false positives
        if(href.indexOf('/delete') === -1){ continue; }
        makeButton(el, href);
      }
    }
    async function doDelete(type,id){
      if(!type||!id){ alert('Nedostaju podaci za brisanje.'); return; }
      var url = '/api/'+type+'/'+encodeURIComponent(id)+'/delete';
      if(!confirm('Potvrdi brisanje?')) return;
      try{
        var res = await fetch(url, {method:'POST', credentials:'same-origin'});
        var data = {}; try{ data = await res.json(); }catch(e){}
        if(!res.ok || (data && data.ok===false)){
          alert((data && data.error) ? data.error : ('Greška ('+res.status+').')); return;
        }
        var row = document.querySelector('tr[data-row-id="'+CSS.escape(id)+'"]');
        if(row && row.parentNode) row.parentNode.removeChild(row); else location.reload();
      }catch(e){ alert('Greška pri brisanju.'); }
    }
    document.addEventListener('click', function(ev){
      var b = ev.target.closest && ev.target.closest('.btn-delete');
      if(!b) return; ev.preventDefault();
      var type = (b.getAttribute('data-type')||'').toLowerCase();
      var id = b.getAttribute('data-id')||'';
      if(!id){
        var tr=b.closest && b.closest('tr[data-row-id]'); if(tr) id = tr.getAttribute('data-row-id')||'';
      }
      if(!type){
        // Last-resort inference by surrounding section
        var container = b.closest('[data-entity]'); if(container){ type = (container.getAttribute('data-entity')||'').toLowerCase(); }
      }
      doDelete(type,id);
    }, {passive:false});
    normalize();
    new MutationObserver(normalize).observe(document.documentElement, {childList:true, subtree:true});
  });
})();
</script>
"""
        js = js.replace('%ICON_URL%', icon_url)

        # Inject CSS in <head> and JS before </body>
        if '</head>' in html:
            html = html.replace('</head>', css + '\n</head>')
        else:
            html = css + html
        if '</body>' in html:
            html = html.replace('</body>', js + '\n</body>')
        else:
            html = html + js
        response.set_data(html)
    except Exception as _e:
        # Swallow errors – never break the main response
        pass
    return response

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000, debug=True, use_reloader=False)


@app.route('/api/klijenti/statuses')
@login_required
def api_klijenti_statuses():
    out = _compute_all_clients_statuses()
    return jsonify({'ok': True, 'data': out}), 200


@app.route('/api/klijent/<path:name>/status', methods=['POST'])
@login_required
def api_set_klijent_status(name):
    try:
        payload = request.get_json(silent=True, force=True) or {}
    except Exception:
        payload = {}
    val = payload.get('active', None)
    if isinstance(val, str):
        val = val.strip().lower() in ('1','true','yes','da','aktivno','aktivan')
    if val is None:
        return jsonify({'ok': False, 'error': 'Nedostaje polje "active" (true/false).'}), 400
    try:
        kl = read_klijenti() or []
    except Exception:
        kl = []
    updated = False
    for k in kl:
        try:
            if str(k.get('name') or k.get('naziv') or '') == str(name):
                k['active'] = bool(val)
                updated = True
                break
        except Exception:
            pass
    if not updated:
        return jsonify({'ok': False, 'error': 'Klijent nije pronađen.'}), 404
    try:
        write_klijenti(kl)
    except Exception:
        pass
    # Broadcast change
    try:
        _broadcast_client_status(str(name), bool(val))
    except Exception:
        pass
    return jsonify({'ok': True, 'name': str(name), 'active': bool(val)}), 200

@app.route('/sim/export.xlsx')
def export_sim_xlsx():
    """Export SIM data to an .xlsx file."""
    # Expect a helper function get_sim_export_bytes() that returns bytes,
    # otherwise fall back to generating a simple empty workbook.
    try:
        content = get_sim_export_bytes()  # user-defined elsewhere
    except Exception:
        # Fallback: create a tiny XLSX in-memory so route always works
        import io
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.append(["SIM", "MSISDN", "IMEI", "Status"])
            buf = io.BytesIO()
            wb.save(buf)
            content = buf.getvalue()
        except Exception:
            content = b''
    from flask import Response
    return Response(
        content,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename=sim_export.xlsx'}
    )


@app.route('/sim/import', methods=['POST'])
def import_sim_xlsx():
    """Handle uploaded .xlsx to import SIM data."""
    from flask import request, redirect, url_for, flash
    file = request.files.get('file')
    if not file or not file.filename.lower().endswith('.xlsx'):
        flash('Molimo odaberite .xlsx datoteku.', 'danger')
        return redirect(request.referrer or url_for('sim'))
    try:
        data = file.read()
        # Expect a helper function process_sim_import_bytes() to parse & persist
        try:
            process_sim_import_bytes(data)  # user-defined elsewhere
            flash('Uvoz uspješno završen.', 'success')
        except Exception as e:
            flash(f'Greška pri uvozu: {e}', 'danger')
    except Exception:
        flash('Greška pri čitanju datoteke.', 'danger')
    return redirect(request.referrer or url_for('sim'))


# ================== XLSX EXPORT/IMPORT ZA KLIJENTE ==================
@app.route('/klijenti/export.xlsx')
@login_required
def export_klijenti_xlsx():
    """Excel export svih klijenata (sva polja kao kolone). Dozvoljeno admin/superadmin."""
    try:
        from flask import send_file, abort
        from openpyxl import Workbook
        import io, json
    except Exception:
        # Ako nema potrebnih paketa, vrati 500
        abort(500)

    # Dozvoljeno samo superadmin/admin
    if not has_role('superadmin','admin'):
        abort(403)

    try:
        data = read_klijenti() or []
    except Exception:
        data = []

    # Skupi sve kolone iz svih zapisa
    columns = []
    for row in data:
        if isinstance(row, dict):
            for k in row.keys():
                if k not in columns:
                    columns.append(k)
    # Ako nema zapisa, barem osnovne kolone
    if not columns:
        columns = ['name', 'oib', 'headquarters', 'shipping', 'email', 'phone', 'created_at']

    # Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Klijenti"
    ws.append(columns)
    for row in data:
        vals = []
        if isinstance(row, dict):
            for c in columns:
                v = row.get(c, '')
                if isinstance(v, (list, dict)):
                    v = json.dumps(v, ensure_ascii=False)
                vals.append(str(v) if v is not None else '')
        else:
            vals = [''] * len(columns)
        ws.append(vals)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="klijenti.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.route('/klijenti/import', methods=['POST'])
@login_required
def import_klijenti_xlsx():
    """Excel import klijenata. Sprečava duplikate po OIB-u. Dozvoljeno admin/superadmin."""
    from flask import request, redirect, url_for, flash
    import re, json, datetime
    try:
        from openpyxl import load_workbook
    except Exception as e:
        flash('Nedostaje paket openpyxl: %s' % e, 'danger')
        return redirect(url_for('klijenti'))

    # Dozvoljeno samo superadmin/admin
    if not has_role('superadmin','admin'):
        flash('Nemate ovlasti za uvoz klijenata.', 'danger')
        return redirect(url_for('klijenti'))

    file = request.files.get('file')
    if not file or not file.filename.lower().endswith('.xlsx'):
        flash('Odaberite .xlsx datoteku.', 'danger')
        return redirect(url_for('klijenti'))

    try:
        wb = load_workbook(file, read_only=True, data_only=True)
        ws = wb.active
    except Exception as e:
        flash(f'Ne mogu pročitati XLSX: {e}', 'danger')
        return redirect(url_for('klijenti'))

    # Header
    try:
        header_cells = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        header = [str(x or '').strip() for x in header_cells]
    except Exception:
        header = []

    def _idx(name):
        ln = str(name or '').strip().lower()
        for i, h in enumerate(header):
            if str(h).strip().lower() == ln:
                return i
        return -1

    idx_oib = _idx('oib')
    idx_name = _idx('name')
    if idx_name == -1:
        idx_name = _idx('naziv')

    if idx_oib == -1 or idx_name == -1:
        flash("Zaglavlje mora sadržavati barem stupce 'OIB' i 'Name' (ili 'Naziv').", 'danger')
        return redirect(url_for('klijenti'))

    # Postojeći OIB-ovi (normalizirani: samo znamenke)
    def _norm_oib(v):
        return re.sub(r'[^0-9]', '', str(v or ''))

    try:
        existing = read_klijenti() or []
    except Exception:
        existing = []
    existing_oib = {_norm_oib(x.get('oib')) for x in existing if isinstance(x, dict)}

    added = 0
    skipped_dup = 0
    skipped_invalid = 0

    # Skupljamo sve kolone iz headera za fleksibilan import
    import_columns = header[:]

    for r in ws.iter_rows(min_row=2, values_only=True):
        try:
            oib_raw = r[idx_oib] if idx_oib < len(r) else ''
            name_raw = r[idx_name] if idx_name < len(r) else ''
            oib = _norm_oib(oib_raw)
            name = str(name_raw or '').strip()
            if not oib or not name:
                skipped_invalid += 1
                continue
            if oib in existing_oib:
                skipped_dup += 1
                continue
            # Mapiraj cijeli red u dict prema headerima
            row = {}
            for i, col in enumerate(import_columns):
                key = str(col or '').strip()
                if not key:
                    continue
                val = r[i] if i < len(r) else ''
                if isinstance(val, (datetime.datetime, datetime.date)):
                    val = val.isoformat()
                row[key] = val if val is not None else ''
            # Normaliziraj ključna polja
            row.setdefault('name', name)
            row['oib'] = oib_raw if isinstance(oib_raw, str) else str(oib_raw or '')
            row.setdefault('created_at', datetime.datetime.now().isoformat())
            existing.append(row)
            existing_oib.add(oib)
            added += 1
        except Exception:
            skipped_invalid += 1

    if added:
        try:
            write_klijenti(existing)
        except Exception as e:
            flash(f'Greška pri spremanju: {e}', 'danger')
            return redirect(url_for('klijenti'))

    flash(f'Uvoz dovršen. Dodano: {added}, duplikati (po OIB-u): {skipped_dup}, neispravni redovi: {skipped_invalid}.', 'success' if added else 'warning')
    return redirect(url_for('klijenti'))
# ================== /XLSX EXPORT/IMPORT ==================

